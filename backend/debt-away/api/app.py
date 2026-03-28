import boto3
import json
import stripe
import uuid
import time
from datetime import datetime
from config import STRIPE_SECRET_KEY, STRIPE_SECRET_KEY_AMBITOLOGY, OPENAI_APIKEY
import openai
import requests
from bs4 import BeautifulSoup
from boto3.dynamodb.conditions import Key
import PyPDF2
import io
from aws_lambda_powertools import Logger
from decimal import Decimal
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, Request, Query, HTTPException, File, UploadFile, Form, BackgroundTasks
from fastapi.responses import FileResponse, Response, JSONResponse
from mangum import Mangum  # type: ignore
import asyncio
from pydantic import BaseModel
from typing import List, Optional, Dict, Any, Tuple
import subprocess
import os
import shutil
try:
    from docx import Document
except ImportError:
    Document = None


logger = Logger()
s3 = boto3.client("s3", region_name="us-east-1")
stripe.api_key = STRIPE_SECRET_KEY

# Initialize OpenAI client
client = openai.OpenAI(api_key=OPENAI_APIKEY)

# Pydantic models for structured output
class AspectScore(BaseModel):
    score: int  # 1-10
    strengths: List[str]
    improvements: List[str]  # 2-5 improvement advice

class BackgroundInfo(BaseModel):
    full_name: Optional[str] = None
    email_address: Optional[str] = None
    home_address: Optional[str] = None
    personal_website: Optional[str] = None

class EducationInfo(BaseModel):
    college_name: Optional[str] = None
    degree: Optional[str] = None
    major: Optional[str] = None
    location: Optional[str] = None
    duration: Optional[str] = None
    gpa: Optional[str] = None
    coursework: List[str] = []

class ProfessionalInfo(BaseModel):
    company_name: Optional[str] = None
    job_title: Optional[str] = None
    location: Optional[str] = None
    duration: Optional[str] = None
    technologies_used: List[str] = []

class TechnicalSkills(BaseModel):
    domain_knowledge: List[str] = []
    technologies: List[str] = []
    tools: List[str] = []
    software: List[str] = []

class TeamworkInfo(BaseModel):
    teamwork_experience: List[str] = []
    leadership_experience: List[str] = []

class ATSReview(BaseModel):
    formatting_issues: List[str] = []
    syntax_issues: List[str] = []
    ats_compatibility_score: int  # 0-100

class ResumeParsedData(BaseModel):
    background: BackgroundInfo
    education: List[EducationInfo]
    professional: List[ProfessionalInfo]
    technical_skills: TechnicalSkills
    teamwork: TeamworkInfo
    ats_review: ATSReview

class ResumeAnalysis(BaseModel):
    background_score: int  # 1-10
    education_score: int  # 1-10
    professional_score: int  # 1-10
    technical_skills_score: int  # 1-10
    teamwork_score: int  # 1-10
    ats_score: int  # 1-10
    overall_score: int  # 1-10
    background_improvements: List[str] = []
    education_improvements: List[str] = []
    professional_improvements: List[str] = []
    technical_skills_improvements: List[str] = []
    teamwork_improvements: List[str] = []
    ats_improvements: List[str] = []
    general_improvements: List[str] = []

class JobExtraction(BaseModel):
    standardized_title: str
    technical_skills: List[str]
    soft_skills: List[str]
    industry: str
    experience_level: str
    key_responsibilities: List[str]
    salary_range: Optional[str]
    company_name: Optional[str]

class JobUrlExtraction(BaseModel):
    target_job_title: str
    target_job_company: str
    target_job_description: str
    target_job_skill_keywords: List[str]

# Pydantic models for craft_resume_from_knowledge_base
class CraftResumeEducation(BaseModel):
    college_name: str
    degree: str
    major: Optional[str] = None
    coursework: Optional[List[str]] = []
    location: Optional[str] = None
    start_date: str
    end_date: str

class CraftResumeProfessionalHistory(BaseModel):
    company_name: str
    job_title: str
    location: Optional[str] = None
    start_date: str
    end_date: str

class CraftResumeProfessionalAchievement(BaseModel):
    type: str
    value: str

class CraftResumeProjectTechContent(BaseModel):
    content: str  # Limited to 250 characters

class CraftResumeProject(BaseModel):
    project_name: str
    location: Optional[str] = None
    start_date: str
    end_date: str
    overview_content: str  # Limited to 250 characters
    tech_content: List[CraftResumeProjectTechContent]  # Up to 2 items, each 250 chars
    achievement_content: str  # Limited to 250 characters
    technologies: List[str]  # Technical keywords

class CraftResumeProfessionalProject(BaseModel):
    project_name: str
    work_experience: str
    location: Optional[str] = None
    overview_content: str  # Limited to 250 characters
    tech_content: List[CraftResumeProjectTechContent]  # Up to 2 items, each 250 chars
    achievement_content: str  # Limited to 250 characters
    technologies: List[str]  # Technical keywords

class CraftResumeTechnicalSkills(BaseModel):
    Languages: List[str] = []
    Frameworks: List[str] = []
    Tools: List[str] = []
    # Additional topics as needed

class CraftResumeResponse(BaseModel):
    full_name: str
    email_address: str
    phone_number: Optional[str] = None
    home_address: Optional[str] = None
    links: Optional[List[str]] = []
    education_history: List[CraftResumeEducation]
    professional_history: List[CraftResumeProfessionalHistory]
    professional_achievement: Optional[List[CraftResumeProfessionalAchievement]] = []
    personal_projects: List[CraftResumeProject]
    professional_projects: List[CraftResumeProfessionalProject]
    technical_skills: CraftResumeTechnicalSkills

class JobMatchAnalysis(BaseModel):
    standardized_title: str
    technical_skills: List[str]
    soft_skills: List[str]
    industry: str
    experience_level: str
    key_responsibilities: List[str]
    salary_range: Optional[str]
    company_name: Optional[str]

    # User comparison and scoring
    background_score: AspectScore
    education_score: AspectScore
    professional_score: AspectScore
    tech_skills_score: AspectScore
    teamwork_score: AspectScore
    job_match_score: AspectScore

    overall_score: int  # 1-100
    analysis_confidence: str  # high/medium/low

# Resume PDF Generation Models
class ResumeContact(BaseModel):
    label: str  # e.g., "Email", "Phone", "LinkedIn", "GitHub"
    value: str

class ResumeJobTitle(BaseModel):
    id: str
    title: str
    date: str
    bullets: List[str]
    projectTechnologies: Optional[Dict[str, Any]] = None

class ResumeProfessionalExperience(BaseModel):
    id: str
    company: str
    jobTitles: List[ResumeJobTitle]

class ResumeDegree(BaseModel):
    id: str
    degree: str
    description: str

class ResumeEducation(BaseModel):
    id: str
    university: str
    date: str
    degrees: List[ResumeDegree]

class ResumeProject(BaseModel):
    id: str
    name: str
    date: str
    description: str
    bullets: List[str]
    technologies: Optional[List[str]] = None

class ResumeSkill(BaseModel):
    id: str
    topic: str  # e.g., "Languages", "Frameworks", "Tools"
    keywords: str  # comma-separated

class ResumeAchievement(BaseModel):
    id: str
    type: str  # e.g., "Award", "Certification", "Publication"
    value: str

class GenerateResumePDFRequest(BaseModel):
    user_id: str
    name: str
    contact: List[ResumeContact]
    professional_experiences: List[ResumeProfessionalExperience]
    education: List[ResumeEducation]
    projects: List[ResumeProject]
    skills: List[ResumeSkill]
    achievements: List[ResumeAchievement]

# Pydantic models for auto-fill from resume
class AutoFillDegree(BaseModel):
    degree: str
    major: str
    startMonth: str
    startYear: str
    endMonth: str
    endYear: str
    coursework: str

class AutoFillCollege(BaseModel):
    collegeName: str
    location: str
    degrees: List[AutoFillDegree]

class AutoFillCompany(BaseModel):
    companyName: str
    jobTitle: str
    startMonth: str
    startYear: str
    endMonth: str
    endYear: str
    isPresent: bool
    location: str

class AutoFillAchievement(BaseModel):
    type: str
    value: str

class AutoFillProjectDescription(BaseModel):
    overview: str
    techAndTeamwork: str
    achievement: str

class AutoFillPersonalProject(BaseModel):
    projectName: str
    projectDescription: AutoFillProjectDescription
    projectStartMonth: str
    projectStartYear: str
    projectEndMonth: str
    projectEndYear: str
    location: str
    selectedTechnologies: List[str]
    selectedFrameworks: List[str]
    customTechnologies: List[str]
    customFrameworks: List[str]

class AutoFillProfessionalProject(BaseModel):
    projectName: str
    projectDescription: AutoFillProjectDescription
    selectedWorkExperience: str
    projectStartMonth: str
    projectStartYear: str
    projectEndMonth: str
    projectEndYear: str
    selectedTechnologies: List[str]
    selectedFrameworks: List[str]
    customTechnologies: List[str]
    customFrameworks: List[str]

class AutoFillResumeResult(BaseModel):
    career_focus: str
    firstName: str
    middleName: str
    lastName: str
    email: str
    phone: str
    addressStreet: str
    addressState: str
    addressZip: str
    personalWebsite: str
    linkedin: str
    education: List[AutoFillCollege]
    companies: List[AutoFillCompany]
    achievements: List[AutoFillAchievement]
    personal_projects: List[AutoFillPersonalProject]
    professional_projects: List[AutoFillProfessionalProject]
    technical_skill_items: List[str]
    custom_technical_items: List[str]

# Pydantic models for overall analysis
class PersonalCapabilityAnalysis(BaseModel):
    background_score: int  # 1-10
    education_score: int
    professional_score: int
    tech_skills_score: int
    teamwork_score: int
    job_match_score: int  # Average of above 5
    background_advice: List[str]  # 4+ items, 15-30 words each
    education_advice: List[str]
    professional_advice: List[str]
    tech_skills_advice: List[str]
    teamwork_advice: List[str]

class ResumePowerAnalysis(BaseModel):
    background_score: int  # 1-10
    education_score: int
    professional_score: int
    tech_skills_score: int
    teamwork_score: int
    job_match_score: int  # Average of above 5
    ats_compatibility_score: int  # 0-100
    background_advice: List[str]
    education_advice: List[str]
    professional_advice: List[str]
    tech_skills_advice: List[str]
    teamwork_advice: List[str]
    ats_issues: List[str]  # Formatting, syntax, missing info

class ProjectBulletPoints(BaseModel):
    overview: str          # ~13 words: project purpose
    showcase1: str         # ~13 words: key tools/architecture used
    showcase2: Optional[str] = None  # ~13 words: second feature (omit if only 3 bullets)
    achievement: str       # ~13 words: quantitative outcome

from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse as _JSONResponse

app = FastAPI()
handler = Mangum(app, lifespan="off")

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    logger.error(f"422 Validation Error on {request.method} {request.url}: {exc.errors()}")
    return _JSONResponse(status_code=422, content={"detail": exc.errors()})

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    # allow_origins=["*"],
    allow_origins=["http://localhost:3000",
                   "https://ambitology.com",
                   "https://www.ambitology.com",
                   "https://main.d3a19hn400g7xw.amplifyapp.com"],
    allow_credentials=True,
    allow_methods=["OPTIONS", "POST", "GET"],  # Ensure OPTIONS is included
    # Include necessary headers
    allow_headers=["Content-Type", "Authorization"],
    # allow_methods=["*"],
    # allow_headers=["*"],
)


'''
API Definition
'''

@app.middleware("http")
async def log_requests(request, call_next):
    logger.info(f"Incoming request: {request.method} {request.url}")
    response = await call_next(request)
    return response



@app.post("/resume-analysis-lab")
async def resume_analysis_lab(file: UploadFile = File(...), form_data: str = Form(...)):
    try:
        logger.info(f"Received resume analysis lab request")
        
        # Parse form data
        form_data_json = json.loads(form_data)
        email = form_data_json.get('email', '')
        full_name = form_data_json.get('fullName', '')
        current_role = form_data_json.get('currentRole', '')
        target_role = form_data_json.get('targetRole', '')
        career_objectives = form_data_json.get('careerObjectives', '')
        
        logger.info(f"Processing resume analysis for {full_name} ({email})")
        
        # Save file to S3 for later processing (non-blocking)
        bucket_name = "career-landing-group"
        current_date = datetime.now().strftime("%Y-%m-%d")
        base_submission_id = f"{email}-{current_date}"
        client_uuid = str(uuid.uuid4())
        submission_id = f"{base_submission_id}-{client_uuid}"
        file_content = await file.read()
        file_name = file.filename
        object_key = f"resume-analysis-lab/{submission_id}/{file_name}"
        async def upload_file_to_s3():
            s3.put_object(
                Bucket=bucket_name, 
                Key=object_key,
                Body=file_content, 
                ContentType=file.content_type
            )
            logger.info(f"Resume file saved to S3: {object_key}")
        asyncio.create_task(upload_file_to_s3())
        # Save client info JSON to S3 (non-blocking)
        info_key = f"resume-analysis-lab/{submission_id}/client_info.json"
        client_info = {
            "client_uuid": client_uuid,
            "email": email,
            "full_name": full_name,
            "submission_id": submission_id,
            "current_role": current_role,
            "target_role": target_role,
            "career_objectives": career_objectives,
            "submission_date": current_date,
            "file_name": file_name,
            "resume_link": f"https://career-landing-group.s3.us-east-1.amazonaws.com/{object_key}"
        }
        client_info_json = json.dumps(client_info, indent=2)
        async def upload_client_info_to_s3():
            s3.put_object(
                Bucket=bucket_name,
                Key=info_key,
                Body=client_info_json,
                ContentType="application/json"
            )
            logger.info(f"Client information saved to S3: {info_key}")
        asyncio.create_task(upload_client_info_to_s3())
        service_path = "/resume-analysis-lab"
        # Create a Stripe checkout session with custom fields
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=[
                'card',
                'afterpay_clearpay',
                'klarna',
                'cashapp'
                ],
            customer_email=email,
            client_reference_id=client_uuid,
            metadata={
                'service_name': 'Resume Analysis Lab',
                'submission_id': submission_id,
                'client_uuid': client_uuid,
                'email': email,
                'full_name': full_name,
                'target_role': target_role,
                's3_client_info_link': f"https://career-landing-group.s3.us-east-1.amazonaws.com/{info_key}"
            },
            line_items=[
                {
                    'price': 'price_1RGVGwCNRzh0Tnp2WgcmYivd',
                    'quantity': 1,
                },
            ],
            mode='payment',
            allow_promotion_codes=True,
            success_url=f"https://www.careerlandinggroup.com/payment-confirmation?session_id={{CHECKOUT_SESSION_ID}}&email={email}&service_path={service_path}",
            cancel_url="https://www.careerlandinggroup.com/resume-design/resume-analysis-lab/"
        )
        logger.info(f"Stripe checkout session created: {checkout_session.id}")
        return {
            "status": "success",
            "checkout_session_id": checkout_session.id,
            "payment_url": checkout_session.url,
            "direct_payment_link": checkout_session.url,
            "submission_id": submission_id,
            "client_uuid": client_uuid
        }
    except Exception as e:
        logger.error(f"Error in resume analysis lab: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")



@app.post("/instant-mock-interview")
async def instant_mock_interview(
    file: Optional[UploadFile] = File(None),
    form_data: str = Form(...)
):
    try:
        logger.info(f"Received instant mock interview request")
        
        # Parse form data
        form_data_json = json.loads(form_data)
        email = form_data_json.get('email', '')
        full_name = form_data_json.get('fullName', '')
        current_role = form_data_json.get('currentRole', '')
        target_role = form_data_json.get('targetRole', '')
        target_job_link = form_data_json.get('targetJobLink', '')
        interview_question_type = form_data_json.get('interviewQuestionType', '')
        
        logger.info(f"Processing instant mock interview for {full_name} ({email})")
        
        # S3 bucket and date
        bucket_name = "career-landing-group"
        current_date = datetime.now().strftime("%Y-%m-%d")
        base_submission_id = f"{email}-{current_date}"
        client_uuid = str(uuid.uuid4())
        submission_id = f"{base_submission_id}-{client_uuid}"
        logger.info(f"Generated unique submission ID: {submission_id}")
        logger.info(f"UUID component: {client_uuid}")
        
        # Non-blocking S3 file upload
        async def upload_file_to_s3():
            if file is not None:
                file_content = await file.read()
                file_name = file.filename
                object_key = f"instant-mock-interview/{submission_id}/{file_name}"
                s3.put_object(
                    Bucket=bucket_name,
                    Key=object_key,
                    Body=file_content,
                    ContentType=file.content_type
                )
                logger.info(f"Resume file saved to S3: {object_key}")
        if file is not None:
            asyncio.create_task(upload_file_to_s3())
        
        # Store client info as JSON (non-blocking)
        client_info = {
            "client_uuid": client_uuid,
            "email": email,
            "full_name": full_name,
            "current_role": current_role,
            "target_role": target_role,
            "target_job_link": target_job_link,
            "interview_question_type": interview_question_type,
            "submission_date": current_date,
            "submission_id": submission_id,
        }
        async def upload_client_info_to_s3():
            client_info_json = json.dumps(client_info, indent=2)
            info_key = f"instant-mock-interview/{submission_id}/client_info.json"
            s3.put_object(
                Bucket=bucket_name,
                Key=info_key,
                Body=client_info_json,
                ContentType="application/json"
            )
            logger.info(f"Client information saved to S3: {info_key}")
        asyncio.create_task(upload_client_info_to_s3())
        
        service_path = "/instant-mock-interview"
        info_key = f"instant-mock-interview/{submission_id}/client_info.json"

        # Create a Stripe checkout session
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=[
                'card',
                'afterpay_clearpay',
                'klarna',
                'cashapp'
            ],
            customer_email=email,
            client_reference_id=client_uuid,
            metadata={
                'service_name': 'Instant Mock Interview',
                'submission_id': submission_id,
                'client_uuid': client_uuid,
                'email': email,
                'full_name': full_name,
                'target_role': target_role,
                'interview_question_type': interview_question_type,
                's3_client_info_link': f"https://career-landing-group.s3.us-east-1.amazonaws.com/{info_key}"
            },
            line_items=[
                {
                    #'price': 'price_1RLXQuCNRzh0Tnp24yLckBAN', # test price
                    'price': 'price_1RHWPoCNRzh0Tnp2WKqgRJXb',  # live price
                    'quantity': 1,
                },
            ],
            mode='payment',
            allow_promotion_codes=True,
            success_url=f"https://www.careerlandinggroup.com/payment-confirmation?session_id={{CHECKOUT_SESSION_ID}}&email={email}&service_path={service_path}",
            cancel_url="https://www.careerlandinggroup.com/interview-prep/instant-mock-interview/",
            payment_intent_data={
                'setup_future_usage': None,
                'capture_method': 'automatic',
            }
        )

        logger.info(f"Stripe checkout session created: {checkout_session.id}")
        
        return {
            "status": "success",
            "sessionId": checkout_session.id, # Official stripe checkout session id
            "payment_url": checkout_session.url,
            "direct_payment_link": checkout_session.url,
            "client_uuid": client_uuid
        }
    except Exception as e:
        logger.error(f"Error in instant mock interview: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing instant mock interview: {str(e)}")



@app.post("/flash-chat")
async def flash_chat(request: Request):
    try:
        # Parse JSON data from the request
        data = await request.json()
        email = data.get('email', '')
        full_name = data.get('fullName', '')
        current_role = data.get('currentRole', '')
        target_role = data.get('targetRole', '')
        message = data.get('message', '')
        selected_questions = data.get('selectedQuestions', [])
        subscription_type = data.get('subscriptionType', 'monthly')
        # Determine price_id based on subscription_type
        if subscription_type == 'biweekly':
            price_id = 'price_1RUrC2CNRzh0Tnp2ikTobjoM'
        else:
            price_id = 'price_1RJ4LnCNRzh0Tnp2aOj91XNX'
        
        logger.info(f"Received flash chat request from {full_name} ({email})")
        
        # Create a human-readable date string
        current_date = datetime.now().strftime("%Y-%m-%d")
        
        # Create a unique identifier for this chat
        base_chat_id = f"{email}-{current_date}"
        
        # Generate a UUID based on the chat_id
        client_uuid = str(uuid.uuid4())
        
        # Final chat ID combines the readable ID and UUID
        chat_id = f"{base_chat_id}-{client_uuid}"
        
        logger.info(f"Generated unique chat ID: {chat_id}")
        logger.info(f"UUID component: {client_uuid}")
        
        # Create and store client information as JSON (non-blocking)
        client_info = {
            "client_uuid": client_uuid,
            "email": email,
            "full_name": full_name,
            "current_role": current_role,
            "target_role": target_role,
            "message": message,
            "selected_questions": selected_questions,
            "submission_date": current_date,
            'subscription_type': subscription_type,
            "chat_id": chat_id
        }
        info_key = f"flash-chat/{chat_id}/client_info.json"
        client_info_json = json.dumps(client_info, indent=2)
        async def upload_client_info_to_s3():
            s3.put_object(
                Bucket="career-landing-group",
                Key=info_key,
                Body=client_info_json,
                ContentType="application/json"
            )
            logger.info(f"Client chat information saved to S3: {info_key}")
        asyncio.create_task(upload_client_info_to_s3())
        service_path = "/flash-chat"
        # Create a Stripe checkout session with custom fields for monthly subscription
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=[ 'card' ],
            customer_email=email,
            client_reference_id=client_uuid,
            metadata={
                'service_name': 'Flash Chat',
                'chat_id': chat_id,
                'client_uuid': client_uuid,
                'email': email,
                'full_name': full_name,
                'target_role': target_role,
                's3_client_info_link': f"https://career-landing-group.s3.us-east-1.amazonaws.com/{info_key}",
                'subscription_type': subscription_type,
            },
            line_items=[
                {
                    'price': price_id,
                    'quantity': 1,
                },
            ],
            mode='subscription',
            allow_promotion_codes=True,
            subscription_data={
                'trial_period_days': 2,
            },
            success_url=f"https://www.careerlandinggroup.com/payment-confirmation?session_id={{CHECKOUT_SESSION_ID}}&email={email}&service_path={service_path}",
            cancel_url="https://www.careerlandinggroup.com/career-cruise/flash-chat-2/"
        )
        logger.info(f"Stripe subscription checkout session created: {checkout_session.id}")
        return {
            "status": "success",
            "checkout_session_id": checkout_session.id,
            "payment_url": checkout_session.url,
            "direct_payment_link": checkout_session.url,
            "chat_id": chat_id,
            "client_uuid": client_uuid
        }
    except Exception as e:
        logger.error(f"Error in flash chat: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing chat request: {str(e)}")


# ambitology
@app.post("/subscription_stripe_checkout_page_handler")
async def subscription_stripe_checkout_page_handler(request: Request):
    try:
        data = await request.json()
        cognito_sub = data.get('cognito_sub', '')
        email = data.get('email', '')
        selected_plan = data.get('selected_plan', '')

        plan_price_map = {
            '2weeks': 'price_1SzqA3Fixq4WY15eJGcxWv7l',
            '1month': 'price_1Szq9NFixq4WY15eMQzhzm3k',
            '3months': 'price_1Szq83Fixq4WY15eDVms6mWR',
        }

        price_id = plan_price_map.get(selected_plan)
        if not price_id:
            raise HTTPException(status_code=400, detail=f"Invalid plan: {selected_plan}")

        logger.info(f"Creating subscription checkout for {email}, plan: {selected_plan}")

        # Determine success/cancel URLs based on environment
        is_local = os.environ.get("AWS_SAM_LOCAL") == "true"
        if is_local:
            base_url = "http://localhost:3000"
        else:
            base_url = "https://ambitology.com"

        success_url = f"{base_url}/dashboard?session_id={{CHECKOUT_SESSION_ID}}"
        cancel_url = f"{base_url}/dashboard"

        # Temporarily use the Ambitology Stripe key
        original_key = stripe.api_key
        stripe.api_key = STRIPE_SECRET_KEY_AMBITOLOGY

        try:
            checkout_session = stripe.checkout.Session.create(
                payment_method_types=['card'],
                customer_email=email,
                metadata={
                    'service_name': 'Ambitology Subscription',
                    'cognito_sub': cognito_sub,
                    'email': email,
                    'selected_plan': selected_plan,
                },
                line_items=[
                    {
                        'price': price_id,
                        'quantity': 1,
                    },
                ],
                mode='subscription',
                allow_promotion_codes=True,
                success_url=success_url,
                cancel_url=cancel_url,
                branding_settings={
                    'background_color': '#f5f2eb',
                    'button_color': '#9b6a10',
                    'border_style': 'rounded',
                    'font_family': 'inter',
                    'display_name': 'Ambitology',
                },
            )
        finally:
            stripe.api_key = original_key

        logger.info(f"Stripe subscription checkout session created: {checkout_session.id}")

        return {
            "status": "success",
            "payment_url": checkout_session.url,
            "checkout_session_id": checkout_session.id,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in subscription checkout: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error creating checkout session: {str(e)}")


# ambitology
@app.post("/user_authentication")
async def user_authentication(request: Request):
    """
    Authenticate and register user in DynamoDB after Cognito login.
    Creates user record if not exists, returns success if already exists.
    
    Single-Table Design:
    - PK: Cognito subID (user identifier)
    - SK: Section key (METADATA, PROFILE#*, KNOWLEDGE#*, RESUME#*, ANALYSIS#*)
    - GSI1PK: EMAIL#<email> for email-based lookups
    - GSI1SK: User ID for GSI sorting
    """
    try:
        logger.info("Received user authentication request")
        
        # Parse request body
        body = await request.json()
        cognito_sub = body.get('cognito_sub', '').strip()
        email = body.get('email', '').strip()
        
        if not cognito_sub:
            raise HTTPException(status_code=400, detail="cognito_sub is required")
        
        if not email:
            raise HTTPException(status_code=400, detail="email is required")
        
        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        logger.info("Initializing DynamoDB connection for user authentication")
        
        table_name = 'ambit-dashboard-application-data'
        
        # Check if table exists, create if not
        try:
            table = dynamodb.Table(table_name)
            table.load()
            logger.info(f"Table {table_name} exists")
        except dynamodb.meta.client.exceptions.ResourceNotFoundException:
            logger.info(f"Table {table_name} does not exist, creating it...")
            
            # Create table with PK/SK schema and GSI1 for email lookups
            table = dynamodb.create_table(
                TableName=table_name,
                KeySchema=[
                    {'AttributeName': 'PK', 'KeyType': 'HASH'},
                    {'AttributeName': 'SK', 'KeyType': 'RANGE'}
                ],
                AttributeDefinitions=[
                    {'AttributeName': 'PK', 'AttributeType': 'S'},
                    {'AttributeName': 'SK', 'AttributeType': 'S'},
                    {'AttributeName': 'GSI1PK', 'AttributeType': 'S'},
                    {'AttributeName': 'GSI1SK', 'AttributeType': 'S'}
                ],
                GlobalSecondaryIndexes=[
                    {
                        'IndexName': 'GSI1',
                        'KeySchema': [
                            {'AttributeName': 'GSI1PK', 'KeyType': 'HASH'},
                            {'AttributeName': 'GSI1SK', 'KeyType': 'RANGE'}
                        ],
                        'Projection': {'ProjectionType': 'ALL'}
                    }
                ],
                BillingMode='PAY_PER_REQUEST'
            )
            
            # Wait for table to be created
            table.wait_until_exists()
            logger.info(f"Table {table_name} created successfully with GSI1")
        except Exception as table_check_error:
            logger.warning(f"Could not check table status: {str(table_check_error)}")
            table = dynamodb.Table(table_name)
        
        # Check if user already exists
        timestamp = datetime.now().isoformat()
        
        try:
            # Try to get existing user metadata
            response = table.get_item(
                Key={
                    'PK': cognito_sub,
                    'SK': 'METADATA'
                }
            )
            
            if 'Item' in response:
                # User already exists, update last login
                logger.info(f"User {cognito_sub} already exists, updating last login")
                table.update_item(
                    Key={
                        'PK': cognito_sub,
                        'SK': 'METADATA'
                    },
                    UpdateExpression='SET updatedAt = :updatedAt, lastLoginAt = :lastLoginAt',
                    ExpressionAttributeValues={
                        ':updatedAt': timestamp,
                        ':lastLoginAt': timestamp
                    }
                )
                
                return {
                    "status": "success",
                    "message": "User authenticated successfully",
                    "user_id": cognito_sub,
                    "is_new_user": False
                }
            
        except Exception as get_error:
            logger.info(f"User lookup returned no item or error: {str(get_error)}")
        
        # User doesn't exist, create new user record
        logger.info(f"Creating new user record for {cognito_sub}")
        
        # Create user metadata item
        user_item = {
            'PK': cognito_sub,
            'SK': 'METADATA',
            'GSI1PK': f'EMAIL#{email}',
            'GSI1SK': cognito_sub,
            'email': email,
            'createdAt': timestamp,
            'updatedAt': timestamp,
            'lastLoginAt': timestamp,
            'data': {
                'accountStatus': 'active',
                'registrationSource': 'cognito'
            }
        }
        
        # Use conditional put to avoid overwriting existing users (race condition protection)
        try:
            table.put_item(
                Item=user_item,
                ConditionExpression='attribute_not_exists(PK) AND attribute_not_exists(SK)'
            )
            logger.info(f"Successfully created user record for {cognito_sub}")

            # Create SUBSCRIPTION item
            subscription_item = {
                'PK': cognito_sub,
                'SK': 'SUBSCRIPTION',
                'plan': 'free',
                'SUB_ID': '',
                'createdAt': timestamp,
                'updatedAt': timestamp
            }
            table.put_item(Item=subscription_item)
            logger.info(f"Successfully created subscription record for {cognito_sub}")

            # Create USAGE item
            usage_item = {
                'PK': cognito_sub,
                'SK': 'USAGE',
                'craft_count': 0,
                'analysis_count': 0,
                'download_count': 0,
                'createdAt': timestamp,
                'updatedAt': timestamp
            }
            table.put_item(Item=usage_item)
            logger.info(f"Successfully created usage record for {cognito_sub}")

            return {
                "status": "success",
                "message": "User registered successfully",
                "user_id": cognito_sub,
                "is_new_user": True
            }
            
        except dynamodb.meta.client.exceptions.ConditionalCheckFailedException:
            # User was created by another request, treat as existing user
            logger.info(f"User {cognito_sub} was created by concurrent request")
            return {
                "status": "success",
                "message": "User authenticated successfully",
                "user_id": cognito_sub,
                "is_new_user": False
            }
        
    except HTTPException as http_err:
        logger.error(f"HTTP error in user_authentication: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in user_authentication: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@app.get("/get_subscription/{cognito_sub}")
async def get_subscription(cognito_sub: str):
    dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
    table = dynamodb.Table('ambit-dashboard-application-data')
    response = table.get_item(Key={'PK': cognito_sub, 'SK': 'SUBSCRIPTION'})
    item = response.get('Item', {})
    plan = item.get('plan', 'free')
    return {"plan": plan}


# ambitology
@app.get("/get_profile/{cognito_sub}")
async def get_profile(cognito_sub: str):
    """
    Retrieve user profile data from DynamoDB.
    
    Single-Table Design:
    - PK: Cognito subID (user identifier)
    - SK: PROFILE#MAIN
    
    Returns profile data structure:
    - careerFocus: string
    - basicInfo: { firstName, middleName, lastName, email, phone, 
                   addressStreet, addressState, addressZip,
                   personalWebsite, linkedin, links[] }
    - education: [                          # List of college objects
        { id, collegeName, degrees: [
            { id, degree, major, startMonth, startYear, endMonth, endYear, isPresent }
        ]}
      ]
    - professional: {
        companies: [                        # List of company objects
            { id, companyName, jobTitle, startMonth, startYear, endMonth, endYear, isPresent }
        ],
        achievements: [                     # List of achievement objects
            { id, type, value }
        ]
      }
    """
    try:
        logger.info(f"Received get profile request for user: {cognito_sub}")
        
        if not cognito_sub or not cognito_sub.strip():
            raise HTTPException(status_code=400, detail="cognito_sub is required")
        
        cognito_sub = cognito_sub.strip()
        
        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        logger.info("Initializing DynamoDB connection for get profile")
        
        table_name = 'ambit-dashboard-application-data'
        table = dynamodb.Table(table_name)
        
        # Get profile from DynamoDB
        response = table.get_item(
            Key={
                'PK': cognito_sub,
                'SK': 'PROFILE#MAIN'
            }
        )
        
        if 'Item' not in response:
            # No profile found, return empty profile structure
            logger.info(f"No profile found for user {cognito_sub}")
            return {
                "status": "success",
                "profile_exists": False,
                "data": None
            }
        
        # Profile found, extract and return data
        item = response['Item']
        profile_data = item.get('data', {})
        
        logger.info(f"Successfully retrieved profile for user {cognito_sub}")
        
        return {
            "status": "success",
            "profile_exists": True,
            "data": {
                "careerFocus": profile_data.get('careerFocus', ''),
                "basicInfo": profile_data.get('basicInfo', {}),
                "education": profile_data.get('education', []),
                "professional": profile_data.get('professional', {})
            },
            "createdAt": item.get('createdAt'),
            "updatedAt": item.get('updatedAt')
        }
        
    except HTTPException as http_err:
        logger.error(f"HTTP error in get_profile: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in get_profile: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


# ambitology
@app.post("/profile_update")
async def profile_update(request: Request):
    """
    Update user profile data in DynamoDB.
    
    Single-Table Design:
    - PK: Cognito subID (user identifier)
    - SK: PROFILE#MAIN
    
    Data Structure:
    - careerFocus: string
    - basicInfo: { firstName, middleName, lastName, email, phone, 
                   addressStreet, addressState, addressZip,
                   personalWebsite, linkedin, links[] }
    - education: [                          # List of college objects
        { id, collegeName, location, degrees: [
            { id, degree, major, startMonth, startYear, endMonth, endYear, coursework }
        ]}
      ]
    - professional: {
        companies: [                        # List of company objects
            { id, companyName, jobTitle, startMonth, startYear, endMonth, endYear, isPresent, location }
        ],
        achievements: [                     # List of achievement objects
            { id, type, value }
        ]
      }
    """
    try:
        logger.info("Received profile update request")
        
        # Parse request body
        body = await request.json()
        cognito_sub = body.get('cognito_sub', '').strip()
        
        if not cognito_sub:
            raise HTTPException(status_code=400, detail="cognito_sub is required")
        
        # Extract profile data
        career_focus = body.get('careerFocus', '')
        basic_info = body.get('basicInfo', {})
        education = body.get('education', [])
        professional = body.get('professional', {})
        
        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        logger.info("Initializing DynamoDB connection for profile update")
        
        table_name = 'ambit-dashboard-application-data'
        table = dynamodb.Table(table_name)
        
        # Create timestamp
        timestamp = datetime.now().isoformat()
        
        # Check if profile already exists
        try:
            response = table.get_item(
                Key={
                    'PK': cognito_sub,
                    'SK': 'PROFILE#MAIN'
                }
            )
            
            if 'Item' in response:
                # Profile exists, update it
                logger.info(f"Updating existing profile for user {cognito_sub}")
                table.update_item(
                    Key={
                        'PK': cognito_sub,
                        'SK': 'PROFILE#MAIN'
                    },
                    UpdateExpression='SET #data = :data, updatedAt = :updatedAt',
                    ExpressionAttributeNames={
                        '#data': 'data'
                    },
                    ExpressionAttributeValues={
                        ':data': {
                            'careerFocus': career_focus,
                            'basicInfo': basic_info,
                            'education': education,
                            'professional': professional
                        },
                        ':updatedAt': timestamp
                    }
                )
                
                return {
                    "status": "success",
                    "message": "Profile updated successfully",
                    "user_id": cognito_sub,
                    "is_new_profile": False
                }
            
        except Exception as get_error:
            logger.info(f"Profile lookup returned no item or error: {str(get_error)}")
        
        # Profile doesn't exist, create new profile record
        logger.info(f"Creating new profile record for {cognito_sub}")
        
        profile_item = {
            'PK': cognito_sub,
            'SK': 'PROFILE#MAIN',
            'data': {
                'careerFocus': career_focus,
                'basicInfo': basic_info,
                'education': education,
                'professional': professional
            },
            'createdAt': timestamp,
            'updatedAt': timestamp
        }
        
        table.put_item(Item=profile_item)
        logger.info(f"Successfully created profile record for {cognito_sub}")
        
        return {
            "status": "success",
            "message": "Profile created successfully",
            "user_id": cognito_sub,
            "is_new_profile": True
        }
        
    except HTTPException as http_err:
        logger.error(f"HTTP error in profile_update: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in profile_update: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


# ambitology
@app.post("/established_knowledge_update")
async def established_knowledge_update(request: Request):
    """
    Update user knowledge data in DynamoDB.

    Single-Table Design:
    - PK: Cognito subID (user identifier)
    - SK: KNOWLEDGE#ESTABLISHED#<career_focus>  (one item per career focus)

    Data Structure (stored flat per item, fully isolated by career focus):
    - personal_project: Array of personal project objects with selectedTechnologies[], selectedFrameworks[]
    - professional_project: Array of professional project objects with selectedTechnologies[], selectedFrameworks[]
    - technical_skills: { selectedSkills[], customKeywords: Record<str, str[]>, customLayers: [...] }
    - custom_keywords: Record<sectionName, str[]>  — custom technology keyword options
    - custom_framework_keywords: Record<sectionName, str[]>  — custom framework keyword options
    """
    try:
        logger.info("Received knowledge update request")

        # Parse request body
        body = await request.json()
        cognito_sub = body.get('cognito_sub', '').strip()

        if not cognito_sub:
            raise HTTPException(status_code=400, detail="cognito_sub is required")

        # Extract knowledge data
        career_focus = body.get('career_focus', 'software-engineering').strip() or 'software-engineering'
        personal_project = body.get('personal_project', [])
        professional_project = body.get('professional_project', [])
        technical_skills = body.get('technical_skills', {})
        custom_keywords = body.get('custom_keywords', {})
        custom_framework_keywords = body.get('custom_framework_keywords', {})

        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        logger.info("Initializing DynamoDB connection for knowledge update")

        table_name = 'ambit-dashboard-application-data'
        table = dynamodb.Table(table_name)

        # Each career focus gets its own DynamoDB item (separate sort key)
        sk = f'KNOWLEDGE#ESTABLISHED#{career_focus}'
        timestamp = datetime.now().isoformat()

        # Flat data payload — no career-focus nesting needed (SK already scopes it)
        knowledge_data = {
            'personal_project': personal_project,
            'professional_project': professional_project,
            'technical_skills': technical_skills,
            'custom_keywords': custom_keywords,
            'custom_framework_keywords': custom_framework_keywords,
        }

        # Check if this career-focus item already exists
        response = table.get_item(Key={'PK': cognito_sub, 'SK': sk})

        if 'Item' in response:
            logger.info(f"Updating established knowledge for user {cognito_sub} (career_focus={career_focus})")
            table.update_item(
                Key={'PK': cognito_sub, 'SK': sk},
                UpdateExpression='SET #data = :data, updatedAt = :updatedAt',
                ExpressionAttributeNames={'#data': 'data'},
                ExpressionAttributeValues={':data': knowledge_data, ':updatedAt': timestamp}
            )
            return {
                "status": "success",
                "message": "Knowledge updated successfully",
                "user_id": cognito_sub,
                "is_new_record": False
            }

        # Item does not exist — create it
        logger.info(f"Creating established knowledge for user {cognito_sub} (career_focus={career_focus})")
        table.put_item(Item={
            'PK': cognito_sub,
            'SK': sk,
            'data': knowledge_data,
            'createdAt': timestamp,
            'updatedAt': timestamp,
        })

        return {
            "status": "success",
            "message": "Knowledge created successfully",
            "user_id": cognito_sub,
            "is_new_record": True
        }

    except HTTPException as http_err:
        logger.error(f"HTTP error in knowledge_update: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in knowledge_update: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


# ambitology
@app.get("/get_knowledge/{cognito_sub}")
async def get_knowledge(cognito_sub: str, career_focus: str = 'software-engineering'):
    """
    Retrieve user established knowledge data from DynamoDB.

    Single-Table Design:
    - PK: Cognito subID (user identifier)
    - SK: KNOWLEDGE#ESTABLISHED#<career_focus>  (one item per career focus)

    Backward-compatibility: also checks the legacy SK 'KNOWLEDGE#ESTABLISHED' for users
    who had data before career-focus isolation was introduced.
    """
    try:
        logger.info(f"Received get knowledge request for user: {cognito_sub}, career_focus: {career_focus}")

        if not cognito_sub or not cognito_sub.strip():
            raise HTTPException(status_code=400, detail="cognito_sub is required")

        cognito_sub = cognito_sub.strip()
        career_focus = (career_focus or 'software-engineering').strip() or 'software-engineering'

        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        logger.info("Initializing DynamoDB connection for get knowledge")

        table_name = 'ambit-dashboard-application-data'
        table = dynamodb.Table(table_name)

        def _build_response(knowledge_data: dict, item: dict) -> dict:
            return {
                "status": "success",
                "knowledge_exists": True,
                "data": {
                    "personal_project": knowledge_data.get('personal_project', []),
                    "professional_project": knowledge_data.get('professional_project', []),
                    "technical_skills": knowledge_data.get('technical_skills', {}),
                    "custom_keywords": knowledge_data.get('custom_keywords', {}),
                    "custom_framework_keywords": knowledge_data.get('custom_framework_keywords', {}),
                },
                "createdAt": item.get('createdAt'),
                "updatedAt": item.get('updatedAt'),
            }

        # 1. Try the career-focus-specific item (new format)
        sk = f'KNOWLEDGE#ESTABLISHED#{career_focus}'
        response = table.get_item(Key={'PK': cognito_sub, 'SK': sk})
        if 'Item' in response:
            logger.info(f"Found career-focus-specific knowledge for user {cognito_sub} (career_focus={career_focus})")
            return _build_response(response['Item'].get('data', {}), response['Item'])

        # 2. Backward-compatibility: check legacy item SK = 'KNOWLEDGE#ESTABLISHED'
        if career_focus == 'software-engineering':
            legacy_response = table.get_item(Key={'PK': cognito_sub, 'SK': 'KNOWLEDGE#ESTABLISHED'})
            if 'Item' in legacy_response:
                raw_data = legacy_response['Item'].get('data', {})
                # Flat format (original legacy) — data is at the top level
                if 'personal_project' in raw_data or 'professional_project' in raw_data or 'technical_skills' in raw_data:
                    logger.info(f"Returning legacy flat-format knowledge for user {cognito_sub}")
                    return _build_response(raw_data, legacy_response['Item'])
                # Nested format (interim migration) — data is under career-focus key
                cf_data = raw_data.get('software-engineering', {})
                if cf_data:
                    logger.info(f"Returning legacy nested-format knowledge for user {cognito_sub}")
                    return _build_response(cf_data, legacy_response['Item'])

        logger.info(f"No knowledge found for user {cognito_sub} (career_focus={career_focus})")
        return {"status": "success", "knowledge_exists": False, "data": None}

    except HTTPException as http_err:
        logger.error(f"HTTP error in get_knowledge: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in get_knowledge: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


# ambitology
@app.post("/expanding_knowledge_update")
async def expanding_knowledge_update(request: Request):
    """
    Update user expanding knowledge data in DynamoDB.

    Single-Table Design:
    - PK: Cognito subID (user identifier)
    - SK: KNOWLEDGE#EXPANDING#<career_focus>  (one item per career focus)

    Data stored flat per item (fully isolated by career focus):
    - future_personal_project, future_professional_project, future_technical_skills
    - custom_future_keywords, custom_future_framework_keywords
    """
    try:
        logger.info("Received expanding knowledge update request")

        # Parse request body
        body = await request.json()
        cognito_sub = body.get('cognito_sub', '').strip()

        if not cognito_sub:
            raise HTTPException(status_code=400, detail="cognito_sub is required")

        # Extract expanding knowledge data
        career_focus = body.get('career_focus', 'software-engineering').strip() or 'software-engineering'
        future_personal_project = body.get('future_personal_project', [])
        future_professional_project = body.get('future_professional_project', [])
        future_technical_skills = body.get('future_technical_skills', {})
        custom_future_keywords = body.get('custom_future_keywords', {})
        custom_future_framework_keywords = body.get('custom_future_framework_keywords', {})

        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        logger.info("Initializing DynamoDB connection for expanding knowledge update")

        table_name = 'ambit-dashboard-application-data'
        table = dynamodb.Table(table_name)

        # Each career focus gets its own DynamoDB item (separate sort key)
        sk = f'KNOWLEDGE#EXPANDING#{career_focus}'
        timestamp = datetime.now().isoformat()

        # Flat data payload — no career-focus nesting needed (SK already scopes it)
        expanding_data = {
            'future_personal_project': future_personal_project,
            'future_professional_project': future_professional_project,
            'future_technical_skills': future_technical_skills,
            'custom_future_keywords': custom_future_keywords,
            'custom_future_framework_keywords': custom_future_framework_keywords,
        }

        # Check if this career-focus item already exists
        response = table.get_item(Key={'PK': cognito_sub, 'SK': sk})

        if 'Item' in response:
            logger.info(f"Updating expanding knowledge for user {cognito_sub} (career_focus={career_focus})")
            table.update_item(
                Key={'PK': cognito_sub, 'SK': sk},
                UpdateExpression='SET #data = :data, updatedAt = :updatedAt',
                ExpressionAttributeNames={'#data': 'data'},
                ExpressionAttributeValues={':data': expanding_data, ':updatedAt': timestamp}
            )
            return {
                "status": "success",
                "message": "Expanding knowledge updated successfully",
                "user_id": cognito_sub,
                "is_new_record": False
            }

        # Item does not exist — create it
        logger.info(f"Creating expanding knowledge for user {cognito_sub} (career_focus={career_focus})")
        table.put_item(Item={
            'PK': cognito_sub,
            'SK': sk,
            'data': expanding_data,
            'createdAt': timestamp,
            'updatedAt': timestamp,
        })

        return {
            "status": "success",
            "message": "Expanding knowledge created successfully",
            "user_id": cognito_sub,
            "is_new_record": True
        }

    except HTTPException as http_err:
        logger.error(f"HTTP error in expanding_knowledge_update: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in expanding_knowledge_update: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


# ambitology
@app.get("/get_expanding_knowledge/{cognito_sub}")
async def get_expanding_knowledge(cognito_sub: str, career_focus: str = 'software-engineering'):
    """
    Retrieve user expanding knowledge data from DynamoDB.

    Single-Table Design:
    - PK: Cognito subID (user identifier)
    - SK: KNOWLEDGE#EXPANDING#<career_focus>  (one item per career focus)

    Backward-compatibility: also checks the legacy SK 'KNOWLEDGE#EXPANDING' for users
    who had data before career-focus isolation was introduced.
    """
    try:
        logger.info(f"Received get expanding knowledge request for user: {cognito_sub}, career_focus: {career_focus}")

        if not cognito_sub or not cognito_sub.strip():
            raise HTTPException(status_code=400, detail="cognito_sub is required")

        cognito_sub = cognito_sub.strip()
        career_focus = (career_focus or 'software-engineering').strip() or 'software-engineering'

        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        logger.info("Initializing DynamoDB connection for get expanding knowledge")

        table_name = 'ambit-dashboard-application-data'
        table = dynamodb.Table(table_name)

        def _build_response(expanding_data: dict, item: dict) -> dict:
            return {
                "status": "success",
                "expanding_knowledge_exists": True,
                "data": {
                    "future_personal_project": expanding_data.get('future_personal_project', []),
                    "future_professional_project": expanding_data.get('future_professional_project', []),
                    "future_technical_skills": expanding_data.get('future_technical_skills', {}),
                    "custom_future_keywords": expanding_data.get('custom_future_keywords', {}),
                    "custom_future_framework_keywords": expanding_data.get('custom_future_framework_keywords', {}),
                },
                "createdAt": item.get('createdAt'),
                "updatedAt": item.get('updatedAt'),
            }

        # 1. Try the career-focus-specific item (new format)
        sk = f'KNOWLEDGE#EXPANDING#{career_focus}'
        response = table.get_item(Key={'PK': cognito_sub, 'SK': sk})
        if 'Item' in response:
            logger.info(f"Found career-focus-specific expanding knowledge for user {cognito_sub} (career_focus={career_focus})")
            return _build_response(response['Item'].get('data', {}), response['Item'])

        # 2. Backward-compatibility: check legacy item SK = 'KNOWLEDGE#EXPANDING'
        if career_focus == 'software-engineering':
            legacy_response = table.get_item(Key={'PK': cognito_sub, 'SK': 'KNOWLEDGE#EXPANDING'})
            if 'Item' in legacy_response:
                raw_data = legacy_response['Item'].get('data', {})
                # Flat format (original legacy)
                if 'future_personal_project' in raw_data or 'future_professional_project' in raw_data or 'future_technical_skills' in raw_data:
                    logger.info(f"Returning legacy flat-format expanding knowledge for user {cognito_sub}")
                    return _build_response(raw_data, legacy_response['Item'])
                # Nested format (interim migration)
                cf_data = raw_data.get('software-engineering', {})
                if cf_data:
                    logger.info(f"Returning legacy nested-format expanding knowledge for user {cognito_sub}")
                    return _build_response(cf_data, legacy_response['Item'])

        logger.info(f"No expanding knowledge found for user {cognito_sub} (career_focus={career_focus})")
        return {"status": "success", "expanding_knowledge_exists": False, "data": None}

    except HTTPException as http_err:
        logger.error(f"HTTP error in get_expanding_knowledge: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in get_expanding_knowledge: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


async def check_robots_txt(url: str) -> Tuple[bool, str]:
    """
    Check robots.txt to see if the URL path is blocked from scraping.
    
    Returns:
        (is_blocked, reason) - True if blocked, False otherwise, with reason string
    """
    try:
        from urllib.parse import urlparse, urljoin
        
        # Parse the URL to get base domain
        parsed = urlparse(url)
        base_url = f"{parsed.scheme}://{parsed.netloc}"
        robots_url = urljoin(base_url, '/robots.txt')
        
        logger.info(f"Checking robots.txt at: {robots_url}")
        
        # Fetch robots.txt
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
        }
        
        try:
            response = requests.get(robots_url, headers=headers, timeout=10)
            response.raise_for_status()
            robots_content = response.text.lower()
        except requests.exceptions.RequestException as e:
            # If robots.txt doesn't exist or can't be fetched, assume not blocked
            logger.info(f"Could not fetch robots.txt: {str(e)}")
            return (False, "")
        
        # Get the path from the original URL
        url_path = parsed.path or '/'
        
        # Check for captcha mentions in robots.txt
        captcha_keywords = ['captcha', 'recaptcha', 'hcaptcha', 'verification required']
        if any(keyword in robots_content for keyword in captcha_keywords):
            logger.warning(f"Captcha requirement detected in robots.txt for: {url}")
            return (True, "Website requires captcha verification")
        
        # Parse robots.txt rules
        # Look for "Disallow:" entries that match our path
        lines = robots_content.split('\n')
        current_user_agents = []
        is_blocked = False
        blocking_reason = ""
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            
            # Check for User-agent directive
            if line.lower().startswith('user-agent:'):
                agent = line.split(':', 1)[1].strip().lower()
                current_user_agents = [agent]
                if agent == '*':
                    current_user_agents = ['*']
            
            # Check for Disallow directive
            elif line.lower().startswith('disallow:'):
                disallow_path = line.split(':', 1)[1].strip()
                
                # Check if this disallow rule applies to our path
                if disallow_path:
                    # If disallow is empty, it means allow all (skip)
                    if disallow_path == '':
                        continue
                    
                    # Check if our path matches the disallow pattern
                    # Simple pattern matching: if disallow path is a prefix of our path
                    if url_path.startswith(disallow_path) or disallow_path == '/':
                        # Check if this applies to all user agents or common ones
                        if '*' in current_user_agents or len(current_user_agents) == 0:
                            is_blocked = True
                            blocking_reason = f"Path blocked by robots.txt (disallow: {disallow_path})"
                            logger.warning(f"URL path blocked by robots.txt: {url_path} (disallow: {disallow_path})")
                            break
        
        return (is_blocked, blocking_reason)
        
    except Exception as e:
        logger.error(f"Error checking robots.txt: {str(e)}")
        # On error, assume not blocked to avoid false positives
        return (False, "")


# ambitology
@app.post("/validate_and_fetch_job_url")
async def validate_and_fetch_job_url(request: Request):
    """
    Validate a job posting URL, fetch its content, and extract structured job data using OpenAI.
    
    Returns:
    - success: True/False
    - error_code: "INVALID_URL", "FETCH_FAILED", "PARSE_FAILED" (if failed)
    - data: { target_job_title, target_job_company, target_job_description, target_job_skill_keywords }
    """
    try:
        logger.info("Received validate_and_fetch_job_url request")

        # Parse request body
        body = await request.json()
        url = body.get('url', '').strip()
        fallback_title = body.get('fallback_title', '').strip()
        fallback_company = body.get('fallback_company', '').strip()

        # ------------------------------------------------------------------
        # Helper: OpenAI fallback using known job title + company
        # Called whenever URL fetch fails but we have panel-provided metadata.
        # ------------------------------------------------------------------
        async def openai_fallback(title: str, company: str) -> dict:
            logger.info(f"Using OpenAI fallback for '{title}' at '{company}'")
            fallback_prompt = f"""
You are a job market expert with deep knowledge of technology companies and their hiring practices.

Generate accurate, specific structured job information for the following position:

Job Title: {title}
Company: {company}

Based on your knowledge of {company}'s technology stack, culture, and typical role requirements, provide:
1. target_job_title: The exact standardized job title
2. target_job_company: The company name
3. target_job_description: A specific 2-3 sentence description of what this role involves at {company}, referencing their actual tech stack and work if known
4. target_job_skill_keywords: List of 6-10 key technical skills, tools, and qualifications typically required for this exact role at {company}
"""
            fallback_response = client.responses.parse(
                model="gpt-5-mini",
                input=[
                    {"role": "system", "content": "You are an expert job market analyst with knowledge of major tech companies and their hiring requirements."},
                    {"role": "user", "content": fallback_prompt}
                ],
                text_format=JobUrlExtraction
            )
            d = fallback_response.output_parsed
            return {
                "success": True,
                "fallback_used": True,
                "data": {
                    "target_job_title": d.target_job_title or title,
                    "target_job_company": d.target_job_company or company,
                    "target_job_description": d.target_job_description or "",
                    "target_job_skill_keywords": d.target_job_skill_keywords or []
                }
            }

        if not url:
            return {
                "success": False,
                "error_code": "INVALID_URL",
                "message": "URL is required"
            }

        # Validate URL format
        if not url.startswith(('http://', 'https://')):
            return {
                "success": False,
                "error_code": "INVALID_URL",
                "message": "Invalid URL format. URL must start with http:// or https://"
            }

        # Check robots.txt to see if scraping is blocked
        logger.info(f"Checking robots.txt for: {url}")
        is_blocked, blocking_reason = await check_robots_txt(url)

        if is_blocked:
            logger.warning(f"Website blocks scraping for: {url} - {blocking_reason}")
            if fallback_title and fallback_company:
                return await openai_fallback(fallback_title, fallback_company)
            return {
                "success": False,
                "error_code": "BLOCKED_OR_AUTH_REQUIRED",
                "message": blocking_reason or "The website blocks direct fetch or requires authentication/captcha"
            }

        # Try to fetch web content
        logger.info(f"Fetching job posting content from: {url}")
        web_content = await fetch_web_page_content(url)

        # Check if fetch was successful (fetch_web_page_content returns error messages as strings)
        error_indicators = [
            "Invalid URL format",
            "Non-HTML content",
            "Timeout fetching",
            "Connection error",
            "HTTP error",
            "Request error",
            "Missing dependency",
            "Unexpected error",
            "Very little content extracted"
        ]

        is_error = any(indicator in web_content for indicator in error_indicators)

        if is_error or len(web_content.strip()) < 100:
            logger.warning(f"Failed to fetch valid content from URL: {url}")
            if fallback_title and fallback_company:
                return await openai_fallback(fallback_title, fallback_company)
            return {
                "success": False,
                "error_code": "FETCH_FAILED",
                "message": web_content if is_error else "Unable to extract meaningful content from the URL"
            }

        # Use OpenAI to extract structured job data
        logger.info("Extracting structured job data using OpenAI...")

        prompt = f"""
        Extract structured job information from the following job posting content:

        JOB POSTING CONTENT:
        {web_content[:6000]}  # Limit to avoid token limits

        Extract the following information:
        1. target_job_title: The job title/position being advertised
        2. target_job_company: The company name offering the job
        3. target_job_description: A concise summary of the job description (2-3 sentences)
        4. target_job_skill_keywords: List of key skills, technologies, and requirements mentioned

        If any information is not clearly available, provide a reasonable inference or "Not specified".
        """

        try:
            response = client.responses.parse(
                model="gpt-5-mini",
                input=[
                    {
                        "role": "system",
                        "content": "You are an expert job posting analyzer. Extract key information from job postings accurately and concisely."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                text_format=JobUrlExtraction
            )

            extracted_data = response.output_parsed

            # If the extracted content looks empty/thin, fall back to OpenAI knowledge
            has_real_content = (
                extracted_data.target_job_title and
                extracted_data.target_job_description and
                len(extracted_data.target_job_description) > 30 and
                extracted_data.target_job_skill_keywords and
                len(extracted_data.target_job_skill_keywords) >= 2
            )

            if not has_real_content and fallback_title and fallback_company:
                logger.info("Extracted content is thin — switching to OpenAI fallback")
                return await openai_fallback(fallback_title, fallback_company)

            logger.info(f"Successfully extracted job data: {extracted_data.target_job_title} at {extracted_data.target_job_company}")

            return {
                "success": True,
                "data": {
                    "target_job_title": extracted_data.target_job_title,
                    "target_job_company": extracted_data.target_job_company,
                    "target_job_description": extracted_data.target_job_description,
                    "target_job_skill_keywords": extracted_data.target_job_skill_keywords
                }
            }

        except Exception as openai_error:
            logger.error(f"OpenAI parsing error: {str(openai_error)}")
            if fallback_title and fallback_company:
                try:
                    return await openai_fallback(fallback_title, fallback_company)
                except Exception as fallback_error:
                    logger.error(f"OpenAI fallback also failed: {str(fallback_error)}")
            return {
                "success": False,
                "error_code": "PARSE_FAILED",
                "message": f"Failed to parse job posting content: {str(openai_error)}"
            }

    except HTTPException as http_err:
        logger.error(f"HTTP error in validate_and_fetch_job_url: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in validate_and_fetch_job_url: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        return {
            "success": False,
            "error_code": "FETCH_FAILED",
            "message": f"Internal server error: {str(e)}"
        }


@app.post("/resolve_job_url")
async def resolve_job_url(request: Request):
    """
    Check whether a job posting URL is reachable.
    - If reachable (HTTP < 400), return it unchanged.
    - If not (4xx/5xx, timeout, connection error), use OpenAI web search to
      find the current live job listing for that title at that company and
      return the discovered URL instead.

    Input:  { url: str, job_title: str, company_name: str }
    Output: { url: str, fallback_used: bool }
    """
    body = {}
    try:
        body = await request.json()
        url          = (body.get("url")          or "").strip()
        job_title    = (body.get("job_title")    or "").strip()
        company_name = (body.get("company_name") or "").strip()

        if not url:
            return {"url": url, "fallback_used": False}

        # ------------------------------------------------------------------
        # Step 1: Quick reachability check — HEAD, then GET fallback
        # ------------------------------------------------------------------
        _headers = {"User-Agent": "Mozilla/5.0 (compatible; JobChecker/1.0)"}
        url_ok = False

        try:
            head_resp = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: requests.head(url, timeout=5, allow_redirects=True, headers=_headers),
            )
            url_ok = head_resp.status_code < 400
        except Exception:
            pass

        if not url_ok:
            # Some servers reject HEAD — try a streaming GET without downloading the body
            try:
                get_resp = await asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: requests.get(url, timeout=5, allow_redirects=True,
                                         headers=_headers, stream=True),
                )
                url_ok = get_resp.status_code < 400
                get_resp.close()
            except Exception:
                pass

        if url_ok:
            return {"url": url, "fallback_used": False}

        # ------------------------------------------------------------------
        # Step 2: URL is dead — use OpenAI web search to find the live posting
        # ------------------------------------------------------------------
        logger.info(f"resolve_job_url: '{url}' unreachable — OpenAI fallback for '{job_title}' at '{company_name}'")

        if not (job_title and company_name):
            # No metadata to search with — return original and let the browser handle it
            return {"url": url, "fallback_used": False}

        oa_client = openai.OpenAI(api_key=OPENAI_APIKEY)
        prompt = (
            f'Find the current, live job application page for the "{job_title}" position '
            f"at {company_name} on their official careers website.\n\n"
            "Return ONLY a single raw URL — no markdown, no explanation, nothing else. "
            "The URL must point to a real, currently open job posting or to the company's "
            f'careers search page filtered for "{job_title}". '
            f"If you cannot find a specific open posting, return {company_name}'s main careers page URL."
        )

        oa_response = oa_client.responses.create(
            model="gpt-4o-mini",
            tools=[{"type": "web_search_preview"}],
            input=prompt,
        )

        found_url = ""
        for item in oa_response.output:
            if getattr(item, "type", None) == "message":
                for block in getattr(item, "content", []):
                    if getattr(block, "type", None) == "output_text":
                        found_url += block.text

        # The model should return just a URL; take the first whitespace-separated token
        found_url = (found_url.strip().split()[0] if found_url.strip() else "")

        if found_url.startswith(("http://", "https://")):
            logger.info(f"resolve_job_url: found fallback URL '{found_url}'")
            return {"url": found_url, "fallback_used": True}

        # Model didn't return a usable URL — open original so user still gets something
        return {"url": url, "fallback_used": False}

    except Exception as e:
        logger.error(f"resolve_job_url unexpected error: {e}")
        return {"url": body.get("url", ""), "fallback_used": False}


@app.post("/fetch_with_job_title")
async def fetch_with_job_title(request: Request):
    """
    Validate a job title and generate a generic job description using OpenAI.
    
    Input: job_title (string, <150 chars)
    
    Returns:
    - success: True/False
    - data: { target_job_title, target_job_company, target_job_description, target_job_skill_keywords }
    - message: Error message if validation fails
    """
    try:
        logger.info("Received fetch_with_job_title request")
        
        # Parse request body
        body = await request.json()
        job_title = body.get('job_title', '').strip()
        
        if not job_title:
            return {
                "success": False,
                "message": "Job title is required"
            }
        
        if len(job_title) >= 150:
            return {
                "success": False,
                "message": "Input too long. Please enter a shorter job title or use the job description field."
            }
        
        # Use OpenAI to validate job title and generate generic job description
        logger.info(f"Validating and generating job description for: {job_title}")
        
        validation_prompt = f"""
        Analyze the following text and determine if it represents a valid job title or job position:
        
        INPUT: "{job_title}"
        
        A valid job title typically:
        - Describes a professional role or position (e.g., "Software Engineer", "Product Manager at Google", "Senior Data Scientist")
        - May include company name (e.g., "Software Engineer at Meta")
        - May include seniority level (e.g., "Senior", "Junior", "Lead")
        
        Invalid inputs include:
        - Random text or gibberish
        - Complete sentences that aren't job titles
        - Names of people
        - Locations without a job context
        
        If VALID: Extract the job title and company (if mentioned), then generate:
        1. target_job_title: The standardized job title
        2. target_job_company: The company name if mentioned, otherwise "Various Companies"
        3. target_job_description: A generic 2-3 sentence description of what this role typically involves
        4. target_job_skill_keywords: List of 5-10 key skills typically required for this role
        
        If INVALID: Set is_valid_job_title to false
        """
        
        class JobTitleValidation(BaseModel):
            is_valid_job_title: bool
            target_job_title: Optional[str] = None
            target_job_company: Optional[str] = None
            target_job_description: Optional[str] = None
            target_job_skill_keywords: Optional[List[str]] = None
        
        try:
            response = client.responses.parse(
                model="gpt-5-mini",
                input=[
                    {
                        "role": "system",
                        "content": "You are an expert job market analyst. Validate job titles and generate accurate, industry-standard job descriptions."
                    },
                    {
                        "role": "user",
                        "content": validation_prompt
                    }
                ],
                text_format=JobTitleValidation
            )
            
            result = response.output_parsed
            
            if not result.is_valid_job_title:
                logger.info(f"Invalid job title detected: {job_title}")
                return {
                    "success": False,
                    "message": "Please enter a valid job title"
                }
            
            logger.info(f"Successfully generated job data for: {result.target_job_title}")
            
            return {
                "success": True,
                "data": {
                    "target_job_title": result.target_job_title or job_title,
                    "target_job_company": result.target_job_company or "Various Companies",
                    "target_job_description": result.target_job_description or "",
                    "target_job_skill_keywords": result.target_job_skill_keywords or []
                }
            }
            
        except Exception as openai_error:
            logger.error(f"OpenAI parsing error: {str(openai_error)}")
            return {
                "success": False,
                "message": f"Failed to process job title: {str(openai_error)}"
            }
        
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in fetch_with_job_title: {str(e)}")
        logger.error(f"Full traceback: {error_traceback}")
        return {
            "success": False,
            "message": f"Internal server error: {str(e)}"
        }


@app.post("/parse_job_description")
async def parse_job_description(request: Request):
    """
    Parse a job description text and extract structured job data using OpenAI.
    
    Input: job_description (string, >150 chars)
    
    Returns:
    - success: True/False
    - data: { target_job_title, target_job_company, target_job_description, target_job_skill_keywords }
    - message: Error message if parsing fails
    """
    try:
        logger.info("Received parse_job_description request")
        
        # Parse request body
        body = await request.json()
        job_description = body.get('job_description', '').strip()
        
        if not job_description:
            return {
                "success": False,
                "message": "Job description is required"
            }
        
        if len(job_description) < 150:
            return {
                "success": False,
                "message": "Job description is too short. Please provide a complete job description."
            }
        
        # Use OpenAI to extract structured job data
        logger.info("Extracting structured job data from description using OpenAI...")
        
        prompt = f"""
        Extract structured job information from the following job description:

        JOB DESCRIPTION:
        {job_description[:6000]}  # Limit to avoid token limits

        Extract the following information:
        1. target_job_title: The job title/position being described
        2. target_job_company: The company name if mentioned, otherwise "Not specified"
        3. target_job_description: A concise summary of the job description (2-3 sentences)
        4. target_job_skill_keywords: List of key skills, technologies, and requirements mentioned

        If any information is not clearly available, provide a reasonable inference.
        """
        
        try:
            response = client.responses.parse(
                model="gpt-5-mini",
                input=[
                    {
                        "role": "system",
                        "content": "You are an expert job posting analyzer. Extract key information from job descriptions accurately and concisely."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                text_format=JobUrlExtraction
            )
            
            extracted_data = response.output_parsed
            
            logger.info(f"Successfully extracted job data: {extracted_data.target_job_title} at {extracted_data.target_job_company}")
            
            return {
                "success": True,
                "data": {
                    "target_job_title": extracted_data.target_job_title,
                    "target_job_company": extracted_data.target_job_company,
                    "target_job_description": extracted_data.target_job_description,
                    "target_job_skill_keywords": extracted_data.target_job_skill_keywords
                }
            }
            
        except Exception as openai_error:
            logger.error(f"OpenAI parsing error: {str(openai_error)}")
            return {
                "success": False,
                "message": f"Failed to parse job description: {str(openai_error)}"
            }
        
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in parse_job_description: {str(e)}")
        logger.error(f"Full traceback: {error_traceback}")
        return {
            "success": False,
            "message": f"Internal server error: {str(e)}"
        }


@app.post("/craft_resume_from_knowledge_base")
async def craft_resume_from_knowledge_base(request: Request):
    """
    Craft a resume from knowledge base data using OpenAI.
    
    Request body should contain:
    - basic_info: Profile information (firstName, lastName, email, phone, address, etc.)
    - target_company_type: Industry sector (if provided)
    - target_job_position: Job position details (if provided, includes job description)
    - personal_projects: List of selected personal projects
    - professional_projects: List of selected professional projects
    - future_personal_projects: List of selected future personal projects
    - future_professional_projects: List of selected future professional projects
    - skills: Combined technical skills (from established and future)
    
    Returns structured resume data ready for template population.
    """
    try:
        logger.info("Received craft_resume_from_knowledge_base request")
        
        body = await request.json()
        basic_info = body.get('basic_info', {})
        target_company_type = body.get('target_company_type')
        target_job_position = body.get('target_job_position')
        personal_projects = body.get('personal_projects', [])
        professional_projects = body.get('professional_projects', [])
        future_personal_projects = body.get('future_personal_projects', [])
        future_professional_projects = body.get('future_professional_projects', [])
        skills = body.get('skills', [])
        education = body.get('education', [])
        professional_history = body.get('professional_history', [])
        achievements = body.get('achievements', [])
        cognito_sub = body.get('cognito_sub')

        if cognito_sub:
            dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
            table = dynamodb.Table('ambit-dashboard-application-data')
            timestamp = datetime.utcnow().isoformat()

            # Check user plan
            sub_response = table.get_item(Key={'PK': cognito_sub, 'SK': 'SUBSCRIPTION'})
            plan = sub_response.get('Item', {}).get('plan', 'free')

            # Fetch or create USAGE item
            usage_response = table.get_item(Key={'PK': cognito_sub, 'SK': 'USAGE'})
            usage_item = usage_response.get('Item')

            if not usage_item:
                usage_item = {
                    'PK': cognito_sub,
                    'SK': 'USAGE',
                    'craft_count': 0,
                    'analysis_count': 0,
                    'download_count': 0,
                    'createdAt': timestamp,
                    'updatedAt': timestamp
                }
                table.put_item(Item=usage_item)

            craft_count = int(usage_item.get('craft_count', 0))

            if plan == 'free' and craft_count >= 3:
                return {"success": False, "error_code": "CRAFT_LIMIT_EXCEEDED", "message": "Free plan craft limit reached."}

            # Increment craft_count for all users
            table.update_item(
                Key={'PK': cognito_sub, 'SK': 'USAGE'},
                UpdateExpression='SET craft_count = craft_count + :inc, updatedAt = :ts',
                ExpressionAttributeValues={':inc': 1, ':ts': timestamp}
            )

        # Build the prompt for OpenAI
        job_context = ""
        if target_job_position:
            if isinstance(target_job_position, dict):
                job_description = target_job_position.get('description', '')
                job_title = target_job_position.get('title', '')
                job_company = target_job_position.get('company', '')
                job_context = f"""TARGET JOB POSITION:
- Job Title: {job_title}
- Company: {job_company}
- Job Description: {job_description}
"""
            else:
                job_context = f"TARGET JOB POSITION: {target_job_position}\n"
        if target_company_type:
            job_context += f"TARGET INDUSTRY SECTOR: {target_company_type}\n"

        # Build goal context sentence for the prompt header
        if target_job_position and target_company_type:
            goal_context = f'The resume will be tailored for the "{target_company_type}" industry sector and the specific job position described below.'
        elif target_company_type:
            goal_context = f'The resume will be tailored specifically for the "{target_company_type}" industry sector, emphasizing technical keywords, frameworks, and projects most valued in that field.'
        elif target_job_position:
            goal_context = 'The resume will be tailored for the specific job position described below.'
        else:
            goal_context = "The resume will professionally showcase the candidate's technical skills and projects."

        # Build optional industry-sector instruction block
        industry_instruction = ""
        if target_company_type:
            industry_instruction = f"""
INDUSTRY SECTOR FOCUS — {target_company_type}:
- Prioritize and front-load technical keywords, frameworks, tools, and architectural patterns that are most in-demand for the "{target_company_type}" industry.
- In tech_content for each project, highlight the technologies and design patterns that resonate most with {target_company_type} companies and hiring managers.
- In the technical_skills section, surface and group skills that are most valued by {target_company_type} employers; reorder within each category to lead with the most sector-relevant items.
- In overview_content and achievement_content, use terminology and framing familiar to {target_company_type} hiring managers to maximize ATS relevance.
"""

        # Combine all projects
        all_personal_projects = personal_projects + future_personal_projects
        all_professional_projects = professional_projects + future_professional_projects
        
        # Format additional links, which may be list of strings or list of {name, url} objects
        raw_links = basic_info.get('links', [])
        formatted_links: list[str] = []
        try:
            for link in raw_links:
                if isinstance(link, dict):
                    name = link.get('name') or link.get('label') or "Link"
                    url = link.get('url') or ""
                    if url:
                        formatted_links.append(f"{name}: {url}")
                else:
                    formatted_links.append(str(link))
        except Exception:
            # Fallback to simple string conversion
            formatted_links = [str(l) for l in raw_links]

        prompt = f"""
You are an expert resume writer. Your goal is to build a professional resume for a candidate using all the information provided below. {goal_context}

{job_context}{industry_instruction}
CANDIDATE BASIC INFORMATION:
- First Name: {basic_info.get('firstName', '')}
- Middle Name: {basic_info.get('middleName', '')}
- Last Name: {basic_info.get('lastName', '')}
- Email: {basic_info.get('email', '')}
- Phone: {basic_info.get('phone', '')}
- Address: {basic_info.get('addressStreet', '')}, {basic_info.get('addressState', '')} {basic_info.get('addressZip', '')}
- Personal Website: {basic_info.get('personalWebsite', '')}
- LinkedIn: {basic_info.get('linkedin', '')}
- Additional Links: {', '.join(formatted_links)}

EDUCATION HISTORY:
{json.dumps(education, indent=2)}

PROFESSIONAL WORK HISTORY:
{json.dumps(professional_history, indent=2)}

PROFESSIONAL ACHIEVEMENTS:
{json.dumps(achievements, indent=2)}

PERSONAL PROJECTS (from Established Expertise and Expanding Knowledge Base):
{json.dumps(all_personal_projects, indent=2)}

PROFESSIONAL PROJECTS (from Established Expertise and Expanding Knowledge Base):
{json.dumps(all_professional_projects, indent=2)}

TECHNICAL SKILLS (from Established Expertise and Expanding Knowledge Base):
{json.dumps(skills, indent=2)}

INSTRUCTIONS:
1. Extract full_name by combining firstName, middleName (if provided), and lastName from basic_info.
2. Extract email_address, phone_number, home_address, and links from basic_info.
3. Process education_history: Convert college data to list format, ordered from most recent to oldest. Each item should include college_name, degree, major (if available), coursework (if available), location (if available in input data), start_date, and end_date. IMPORTANT: Include the location field exactly as provided in the input data for each college.
4. Process professional_history: Convert work experience data to list format, ordered from most recent to oldest. Each item should include company_name, job_title, location (if available in input data), start_date, and end_date. IMPORTANT: Include the location field exactly as provided in the input data for each company.
5. Process professional_achievement: Include all achievement items from the achievements list, keeping type and value.
6. Process personal_projects: For each project in the combined list (personal_projects + future_personal_projects):
   - Keep project_name, location (if available in input data), start_date, and end_date from the input data
   - IMPORTANT: Include the location field exactly as provided in the input data for each project
   - Create overview_content: Transform the project description overview field into a high-level professional resume description (max 250 characters)
   - Create tech_content: Create 1-2 items (each max 250 characters) showcasing technologies and tools used for the most important functionalities. Prioritize:
     * Selected technologies and frameworks from the project input
     * Technologies that match target job description requirements
     * If content exceeds 250 chars, split into two items (max 2 items total)
   - Create achievement_content: Describe technical performance improvements first, then sales/functionality improvements if applicable. Match both user input achievements and job description expectations (max 250 characters)
   - Create technologies list: Include remaining technologies/frameworks/tools not mentioned in tech_content, prioritizing those matching target job requirements
7. Process professional_projects: Follow the same rules as personal_projects, but:
   - Include work_experience field from the input
   - IMPORTANT: Include the location field exactly as provided in the input data for each project
   - Use data from professional_projects and future_professional_projects
8. Process technical_skills: Organize skills into categories: "Languages", "Frameworks", "Tools", and any other relevant topics. Prioritize skills that match the target job description technical keywords.
9. IMPORTANT RULES:
   - All content fields (overview_content, tech_content items, achievement_content) must be limited to 250 characters including spaces
   - Use professional resume-style syntax and tone
   - No duplicate technical keywords should appear across personal_projects, professional_projects, and technical_skills
   - Prioritize matching target job description requirements throughout
   - If a TARGET INDUSTRY SECTOR is specified, surface and prioritize technical keywords, tools, and project descriptions most relevant to that industry throughout all sections — this overrides generic keyword selection
   - Location fields must be included exactly as provided in the input data for education, professional history, and all projects

Generate a structured resume data that can be directly used to populate a resume template.
"""
        
        logger.info("Calling OpenAI to craft resume...")
        
        try:
            response = client.responses.parse(
                model="gpt-4o-mini",
                input=[
                    {
                        "role": "system",
                        "content": "You are an expert resume writer specializing in creating ATS-friendly, professional resumes that highlight technical skills and project achievements. Always follow the instructions precisely and ensure all character limits are respected."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                text_format=CraftResumeResponse
            )
            
            crafted_resume = response.output_parsed
            
            logger.info(f"Successfully crafted resume for candidate: {crafted_resume.full_name}")
            
            # Convert Pydantic model to dict for JSON response
            return {
                "success": True,
                "data": crafted_resume.model_dump()
            }
            
        except Exception as openai_error:
            logger.error(f"OpenAI error in craft_resume_from_knowledge_base: {str(openai_error)}")
            return {
                "success": False,
                "error_code": "OPENAI_ERROR",
                "message": f"Failed to craft resume: {str(openai_error)}"
            }
        
    except HTTPException as http_err:
        logger.error(f"HTTP error in craft_resume_from_knowledge_base: {http_err.status_code} - {http_err.detail}")
        raise
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Unexpected error in craft_resume_from_knowledge_base: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        return {
            "success": False,
            "error_code": "INTERNAL_ERROR",
            "message": f"Internal server error: {str(e)}"
        }


@app.post("/craft_resume_from_existing_resume")
async def craft_resume_from_existing_resume(
    resume_file: UploadFile = File(...),
    form_data: str = Form(...)
):
    """
    Craft a resume from an uploaded existing resume file using OpenAI.

    Multipart form fields:
    - resume_file: uploaded PDF / DOC / DOCX resume
    - form_data: JSON string containing:
        - cognito_sub: user identifier (for usage tracking)
        - target_job_position: dict with keys title, company, description, skill_keywords
                               (or None if industry sector only)
        - target_company_type: industry sector string (or None)

    Returns the same CraftResumeResponse structure as craft_resume_from_knowledge_base
    so the frontend can reuse processCraftedResumeData unchanged.
    """
    try:
        logger.info("Received craft_resume_from_existing_resume request")

        # ── Parse JSON form_data ──────────────────────────────────────────────
        try:
            body = json.loads(form_data)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid form_data JSON")

        cognito_sub = body.get("cognito_sub")
        target_job_position = body.get("target_job_position")   # dict or None
        target_company_type = body.get("target_company_type")   # str or None

        # ── Usage / plan gate (same logic as knowledge-base endpoint) ─────────
        if cognito_sub:
            dynamodb = boto3.resource("dynamodb", region_name="us-east-1")
            table = dynamodb.Table("ambit-dashboard-application-data")
            timestamp = datetime.utcnow().isoformat()

            sub_response = table.get_item(Key={"PK": cognito_sub, "SK": "SUBSCRIPTION"})
            plan = sub_response.get("Item", {}).get("plan", "free")

            usage_response = table.get_item(Key={"PK": cognito_sub, "SK": "USAGE"})
            usage_item = usage_response.get("Item")

            if not usage_item:
                usage_item = {
                    "PK": cognito_sub,
                    "SK": "USAGE",
                    "craft_count": 0,
                    "analysis_count": 0,
                    "download_count": 0,
                    "createdAt": timestamp,
                    "updatedAt": timestamp,
                }
                table.put_item(Item=usage_item)

            craft_count = int(usage_item.get("craft_count", 0))
            if plan == "free" and craft_count >= 3:
                return {
                    "success": False,
                    "error_code": "CRAFT_LIMIT_EXCEEDED",
                    "message": "Free plan craft limit reached.",
                }

            table.update_item(
                Key={"PK": cognito_sub, "SK": "USAGE"},
                UpdateExpression="SET craft_count = craft_count + :inc, updatedAt = :ts",
                ExpressionAttributeValues={":inc": 1, ":ts": timestamp},
            )

        # ── Extract text from uploaded resume ─────────────────────────────────
        file_content = await resume_file.read()
        filename = resume_file.filename or "resume"
        logger.info(f"Extracting text from uploaded resume: {filename}")

        resume_text = await extract_text_from_resume(file_content, filename)

        if not resume_text or len(resume_text.strip()) < 50:
            return {
                "success": False,
                "error_code": "EXTRACTION_FAILED",
                "message": "Could not extract meaningful text from the uploaded resume. Please ensure the file is not scanned/image-only.",
            }

        # Truncate to avoid hitting token limits (keep first 6000 chars)
        if len(resume_text) > 6000:
            resume_text = resume_text[:6000] + "\n... [resume content truncated]"

        # ── Build job / industry context ──────────────────────────────────────
        job_context = ""
        if target_job_position and isinstance(target_job_position, dict):
            job_title = target_job_position.get("title", "")
            job_company = target_job_position.get("company", "")
            job_description = target_job_position.get("description", "")
            skill_keywords = target_job_position.get("skill_keywords", [])
            job_context = f"""TARGET JOB POSITION:
- Job Title: {job_title}
- Company: {job_company}
- Job Description: {job_description}
- Key Skill Keywords: {", ".join(skill_keywords) if skill_keywords else "N/A"}
"""
        elif target_job_position and isinstance(target_job_position, str):
            job_context = f"TARGET JOB POSITION: {target_job_position}\n"

        if target_company_type:
            job_context += f"TARGET INDUSTRY SECTOR: {target_company_type}\n"

        # ── Build OpenAI prompt ───────────────────────────────────────────────
        prompt = f"""You are an expert resume writer and career coach. Your goal is to:
1. Parse the candidate's existing resume to extract all structured information.
2. Enrich and enhance every section — especially projects, technical skills, and domain expertise — to better align with the target job position and/or industry sector described below.
3. Return the enriched data in the exact structured format specified.

{job_context}

EXISTING RESUME CONTENT:
{resume_text}

INSTRUCTIONS:
1. Extract full_name, email_address, phone_number, home_address, and any profile links (LinkedIn, GitHub, personal site, etc.) from the resume. Return links as a list of strings.
2. Extract education_history: all education entries ordered most-recent first. Each entry needs college_name, degree, major (if available), coursework (list, if mentioned), location (if available), start_date, end_date.
3. Extract professional_history: all work experience ordered most-recent first. Each entry needs company_name, job_title, location (if available), start_date, end_date.
4. Extract professional_achievement: any awards, certifications, publications, or notable achievements. Each item has a "type" (e.g. "Award", "Certification") and "value" (the description).
5. Process personal_projects: Extract all personal/side projects from the resume. For each project:
   - Keep project_name, location (if available), start_date, end_date.
   - Create overview_content: a high-level professional description ENHANCED to highlight relevance to the target job/industry (max 250 characters).
   - Create tech_content: 1–2 items (each max 250 characters) showcasing the most impactful technologies and methodologies used. Prioritize skills that match the target job's requirements. Split into two items if content exceeds 250 characters.
   - Create achievement_content: quantified or impact-focused achievement statement, ENHANCED toward the target job/industry expectations (max 250 characters).
   - Create technologies: remaining relevant tech keywords not already covered in tech_content, prioritizing those matching target job skills.
6. Process professional_projects: Extract all work-related projects from the professional experience section. Follow the same enrichment rules as personal_projects, and additionally set work_experience to "CompanyName - JobTitle" so it maps back to the correct professional history entry.
7. Process technical_skills: Organize ALL technical skills from the resume into categories: "Languages", "Frameworks", "Tools", and any other relevant topic groupings. Enrich with skills that are closely related to the target job description or industry sector (only include skills the candidate demonstrably has or has used — do not fabricate). If a TARGET INDUSTRY SECTOR is provided, reorder within each category to lead with the most sector-relevant skills.
8. IMPORTANT RULES:
   - All content fields (overview_content, tech_content items, achievement_content) must be ≤ 250 characters including spaces.
   - Use active, professional resume-style language.
   - No duplicate technical keywords across personal_projects, professional_projects, and technical_skills.
   - Prioritize keywords from the target job description throughout.
   - If a TARGET INDUSTRY SECTOR is specified, surface and prioritize technical keywords, tools, and project descriptions most relevant to that industry throughout all sections — this overrides generic keyword selection.
   - If the resume does not have personal projects or professional projects, return empty lists.
   - If a date is unknown or not mentioned, use an empty string.

Generate structured resume data ready for direct template population.
"""

        logger.info("Calling OpenAI to craft resume from existing resume...")

        try:
            response = client.responses.parse(
                model="gpt-4o-mini",
                input=[
                    {
                        "role": "system",
                        "content": (
                            "You are an expert resume writer specialising in creating ATS-friendly, "
                            "professional resumes that highlight technical skills and project achievements. "
                            "Always follow the instructions precisely and ensure all character limits are respected."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                text_format=CraftResumeResponse,
            )

            crafted_resume = response.output_parsed
            logger.info(f"Successfully crafted resume from existing file for: {crafted_resume.full_name}")

            return {"success": True, "data": crafted_resume.model_dump()}

        except Exception as openai_error:
            logger.error(f"OpenAI error in craft_resume_from_existing_resume: {str(openai_error)}")
            return {
                "success": False,
                "error_code": "OPENAI_ERROR",
                "message": f"Failed to craft resume: {str(openai_error)}",
            }

    except HTTPException as http_err:
        logger.error(
            f"HTTP error in craft_resume_from_existing_resume: {http_err.status_code} - {http_err.detail}"
        )
        raise
    except Exception as e:
        import traceback
        logger.error(f"Unexpected error in craft_resume_from_existing_resume: {str(e)}")
        logger.error(traceback.format_exc())
        return {
            "success": False,
            "error_code": "INTERNAL_ERROR",
            "message": f"Internal server error: {str(e)}",
        }


@app.post("/lead_sign_up")
async def lead_sign_up(request: Request):
    try:
        logger.info("Received lead sign up request")
        
        # Get request body
        body = await request.json()
        email = body.get('email', '').strip()
        
        if not email:
            raise HTTPException(status_code=400, detail="Email is required")
        
        # Validate email format
        import re
        email_pattern = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
        if not re.match(email_pattern, email):
            raise HTTPException(status_code=400, detail="Invalid email format")
        
        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        logger.info("Initializing DynamoDB connection")
        
        # Check if table exists, create if not
        table_name = 'lead_email'
        try:
            table = dynamodb.Table(table_name)
            # Try to load the table to see if it exists
            table.load()
            logger.info(f"Table {table_name} exists")
        except dynamodb.meta.client.exceptions.ResourceNotFoundException:
            logger.info(f"Table {table_name} does not exist, creating it...")
            # Create the table
            table = dynamodb.create_table(
                TableName=table_name,
                KeySchema=[
                    {
                        'AttributeName': 'id',
                        'KeyType': 'HASH'
                    }
                ],
                AttributeDefinitions=[
                    {
                        'AttributeName': 'id',
                        'AttributeType': 'S'
                    }
                ],
                BillingMode='PAY_PER_REQUEST'
            )
            # Wait for table to be created
            table.wait_until_exists()
            logger.info(f"Table {table_name} created successfully")
        except Exception as table_check_error:
            # If we can't describe the table due to permissions, assume it exists and try to use it
            logger.warning(f"Could not check table status (permissions issue): {str(table_check_error)}")
            table = dynamodb.Table(table_name)
        
        # Generate unique ID and timestamp
        lead_id = str(uuid.uuid4())
        timestamp = datetime.now().isoformat()
        
        logger.info(f"Attempting to save lead email: {email}")
        # Save to DynamoDB
        table.put_item(
            Item={
                'id': lead_id,
                'email': email,
                'timestamp': timestamp,
                'source': 'lead_page'
            },
            ConditionExpression="attribute_not_exists(id)",
        )
        logger.info(f"Successfully wrote to DynamoDB table: {table_name}")
        
        logger.info(f"Successfully saved lead email: {email}")
        
        return {
            "status": "success",
            "message": "Email successfully saved",
            "lead_id": lead_id
        }
        
    except HTTPException as http_err:
        # Log HTTP exceptions with details
        logger.error(f"HTTP error in lead_sign_up: {http_err.status_code} - {http_err.detail}")
        logger.error(f"Request details - Email: {email if 'email' in locals() else 'N/A'}")
        raise
    except Exception as e:
        # Comprehensive error logging for debugging
        import traceback
        error_traceback = traceback.format_exc()
        
        logger.error(f"Unexpected error in lead_sign_up: {str(e)}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Full traceback: {error_traceback}")
        
        # Log request context for debugging
        try:
            logger.error(f"Request details - Email: {email if 'email' in locals() else 'N/A'}")
            logger.error(f"Lead ID: {lead_id if 'lead_id' in locals() else 'N/A'}")
            logger.error(f"Timestamp: {timestamp if 'timestamp' in locals() else 'N/A'}")
        except:
            logger.error("Could not log request details")
        
        # Log system information for infrastructure debugging
        import os
        logger.error(f"Environment: {os.environ.get('AWS_LAMBDA_FUNCTION_NAME', 'local')}")
        logger.error(f"Region: {os.environ.get('AWS_REGION', 'unknown')}")
        
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")



@app.post("/job_application")
async def job_application(
    form_data: str = Form(...),
    resume: Optional[UploadFile] = File(None)
):
    try:
        logger.info(f"Received job application request")

        # Parse JSON text fields from the single form_data field
        try:
            data = json.loads(form_data)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid form_data JSON")

        firstName = data.get("firstName", "")
        lastName = data.get("lastName", "")
        email = data.get("email", "")
        phoneNumber = data.get("phoneNumber", "")
        isStudent = data.get("isStudent", "")
        currentEmployer = data.get("currentEmployer", "")
        linkedinUrl = data.get("linkedinUrl", "")
        githubUrl = data.get("githubUrl", "")
        portfolioUrl = data.get("portfolioUrl", "")
        websiteUrl = data.get("websiteUrl", "")
        selectedPosition = data.get("selectedPosition", "")

        logger.info(f"Resume filename: {resume.filename if resume else 'None'}")
        logger.info(f"Resume content_type: {resume.content_type if resume else 'None'}")

        full_name = f"{firstName} {lastName}"
        logger.info(f"Processing job application for {full_name} ({email})")
        
        # S3 bucket and date
        bucket_name = "career-landing-group"
        current_date = datetime.now().strftime("%Y-%m-%d")
        base_submission_id = f"{email}-{current_date}"
        client_uuid = str(uuid.uuid4())
        submission_id = f"{base_submission_id}-{client_uuid}"
        
        logger.info(f"Generated unique submission ID: {submission_id}")
        logger.info(f"UUID component: {client_uuid}")
        
        # Save file to S3 for later processing (non-blocking) - following resume_analysis_lab pattern
        file_content = None
        file_name = None
        object_key = None
        
        if resume is not None:
            logger.info(f"Resume file received: {resume.filename}, type: {resume.content_type}")
            file_content = await resume.read()
            logger.info(f"File content read, size: {len(file_content)} bytes")
            file_name = resume.filename
            object_key = f"job_application/{submission_id}/{file_name}"
            
            async def upload_file_to_s3():
                try:
                    s3.put_object(
                        Bucket=bucket_name, 
                        Key=object_key,
                        Body=file_content, 
                        ContentType=resume.content_type
                    )
                    logger.info(f"Resume file saved to S3: {object_key}")
                except Exception as e:
                    logger.error(f"Error uploading file to S3: {str(e)}")
                    raise
            
            asyncio.create_task(upload_file_to_s3())
        else:
            logger.info("No resume file provided")
        
        # Store client info as JSON (non-blocking)
        client_info = {
            "client_uuid": client_uuid,
            "email": email,
            "firstName": firstName,
            "lastName": lastName,
            "full_name": full_name,
            "phoneNumber": phoneNumber,
            "isStudent": isStudent,
            "currentEmployer": currentEmployer,
            "linkedinUrl": linkedinUrl,
            "githubUrl": githubUrl,
            "portfolioUrl": portfolioUrl,
            "websiteUrl": websiteUrl,
            "selectedPosition": selectedPosition,
            "submission_date": current_date,
            "submission_id": submission_id,
        }
        
        # Add resume info if file was uploaded
        if resume is not None and file_name and object_key:
            client_info["file_name"] = file_name
            client_info["resume_link"] = f"https://career-landing-group.s3.us-east-1.amazonaws.com/{object_key}"
        
        async def upload_client_info_to_s3():
            client_info_json = json.dumps(client_info, indent=2)
            info_key = f"job_application/{submission_id}/client_info.json"
            s3.put_object(
                Bucket=bucket_name,
                Key=info_key,
                Body=client_info_json,
                ContentType="application/json"
            )
            logger.info(f"Client information saved to S3: {info_key}")
        asyncio.create_task(upload_client_info_to_s3())
        
        logger.info(f"Job application processed successfully for {full_name} ({email})")
        
        return {
            "status": "success",
            "message": "Job application submitted successfully",
            "submission_id": submission_id,
            "client_uuid": client_uuid
        }
    except Exception as e:
        logger.error(f"Error in job application: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing job application: {str(e)}")


@app.post("/referral_application")
async def referral_application(file: UploadFile = File(...), form_data: str = Form(...)):
    try:
        logger.info(f"Received referral application request")
        
        # Parse form data
        form_data_json = json.loads(form_data)
        email = form_data_json.get('email', '')
        
        logger.info(f"Processing referral application for ({email})")
        
        # Save file to S3 for later processing (non-blocking)
        bucket_name = "career-landing-group"
        current_date = datetime.now().strftime("%Y-%m-%d")
        base_submission_id = f"{email}-{current_date}"
        client_uuid = str(uuid.uuid4())
        submission_id = f"{base_submission_id}-{client_uuid}"
        file_content = await file.read()
        file_name = file.filename
        object_key = f"referral/{submission_id}/{file_name}"
        
        async def upload_file_to_s3():
            s3.put_object(
                Bucket=bucket_name, 
                Key=object_key,
                Body=file_content, 
                ContentType=file.content_type
            )
            logger.info(f"Resume file saved to S3: {object_key}")
        asyncio.create_task(upload_file_to_s3())
        
        # Save client info JSON to S3 (non-blocking)
        info_key = f"referral/{submission_id}/client_info.json"
        client_info = {
            "client_uuid": client_uuid,
            "email": email,
            "submission_id": submission_id,
            "submission_date": current_date,
            "file_name": file_name,
            "resume_link": f"https://career-landing-group.s3.us-east-1.amazonaws.com/{object_key}"
        }
        client_info_json = json.dumps(client_info, indent=2)
        
        async def upload_client_info_to_s3():
            s3.put_object(
                Bucket=bucket_name,
                Key=info_key,
                Body=client_info_json,
                ContentType="application/json"
            )
            logger.info(f"Client information saved to S3: {info_key}")
        asyncio.create_task(upload_client_info_to_s3())
        
        logger.info(f"Referral application processed successfully for ({email})")
        
        return {
            "status": "success",
            "message": "Referral application submitted successfully",
            "submission_id": submission_id,
            "client_uuid": client_uuid
        }
    except Exception as e:
        logger.error(f"Error in referral application: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing referral application: {str(e)}")


@app.post("/job-recommendations")
async def get_job_recommendations(request: Request):
    """
    Load all cached job listings for a position from DynamoDB, then use OpenAI
    to rank the top 20 most relevant jobs based on the user's career focus and
    skills, with popular / well-known companies surfaced first.
    """
    try:
        body = await request.json()
        position_name = body.get("position_name", "").strip()
        career_focus  = body.get("career_focus", "").strip()
        user_skills   = body.get("user_skills", []) or []

        if not position_name:
            return {"success": False, "jobs": [], "error": "position_name is required"}

        normalized = (
            position_name.strip()
            .lower()
            .replace(" ", "_")
            .replace("/", "_")
            .replace("&", "and")
        )

        # ------------------------------------------------------------------
        # 1. Load ALL items for this position from DynamoDB (paginated)
        # ------------------------------------------------------------------
        job_cache_table = boto3.resource("dynamodb", region_name="us-east-1").Table("jobCache")
        all_jobs = []
        query_kwargs = dict(
            IndexName="position-index",
            KeyConditionExpression=Key("position_name").eq(normalized),
            ProjectionExpression="job_title, company_name, job_url",
        )
        while True:
            resp = job_cache_table.query(**query_kwargs)
            all_jobs.extend(resp.get("Items", []))
            last_key = resp.get("LastEvaluatedKey")
            if not last_key:
                break
            query_kwargs["ExclusiveStartKey"] = last_key

        if not all_jobs:
            return {"success": True, "jobs": []}

        # If 20 or fewer items there is nothing to rank — return as-is
        if len(all_jobs) <= 20:
            return {"success": True, "jobs": all_jobs}

        # ------------------------------------------------------------------
        # 2. Ask OpenAI to rank the top 20 most relevant + reputable listings
        # ------------------------------------------------------------------
        job_list_text = "\n".join(
            f"{i}: {j.get('job_title', '')} @ {j.get('company_name', '')}"
            for i, j in enumerate(all_jobs)
        )
        skills_text = ", ".join(user_skills) if user_skills else "not specified"

        ranking_prompt = f"""You are an expert career advisor and technical recruiter.

User profile:
- Career focus: {career_focus or 'not specified'}
- Technical skills: {skills_text}

Below is a numbered list of job listings for the position "{position_name}".
Select the 20 most suitable listings for this user, prioritising:
1. Relevance to the user's career focus and technical skills
2. Company reputation and prestige (FAANG, top-tier tech companies, well-funded startups, unicorns)
3. Variety — avoid returning duplicates from the same company

Return ONLY a JSON array of the selected indices (0-based integers), ordered with the best match first.
Example: [3, 0, 12, 7, ...]

Job listings:
{job_list_text}"""

        class JobRanking(BaseModel):
            ranked_indices: List[int]

        ranking_resp = client.responses.parse(
            model="gpt-5-mini",
            input=[
                {"role": "system", "content": "You are a concise JSON-only responder. Return only the JSON object requested."},
                {"role": "user",   "content": ranking_prompt},
            ],
            text_format=JobRanking,
        )

        ranked_indices = ranking_resp.output_parsed.ranked_indices or []

        # Validate and deduplicate indices, cap at 20
        seen = set()
        ordered = []
        for idx in ranked_indices:
            if isinstance(idx, int) and 0 <= idx < len(all_jobs) and idx not in seen:
                seen.add(idx)
                ordered.append(all_jobs[idx])
                if len(ordered) == 20:
                    break

        # Pad with unranked items if OpenAI returned fewer than 20
        if len(ordered) < 20:
            for i, job in enumerate(all_jobs):
                if i not in seen:
                    ordered.append(job)
                    if len(ordered) == 20:
                        break

        logger.info(f"Ranked {len(ordered)} jobs for position '{position_name}' (total pool: {len(all_jobs)})")
        return {"success": True, "jobs": ordered}

    except Exception as e:
        logger.error(f"Error fetching job recommendations for '{position_name}': {e}")
        return {"success": False, "jobs": [], "error": str(e)}


async def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text content from PDF file bytes
    """
    try:
        # Create a BytesIO object from the file content
        pdf_file = io.BytesIO(file_content)
        
        # Create PDF reader object
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Extract text from all pages
        text_content = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text_content += page.extract_text() + "\n"
        
        logger.info(f"Successfully extracted text from PDF with {len(pdf_reader.pages)} pages")
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        # Fallback: try to decode as text if it's not a PDF
        try:
            return file_content.decode('utf-8')
        except:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")


async def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from DOCX file bytes
    """
    try:
        if Document is None:
            raise Exception("python-docx library is not installed")
        
        # Create a BytesIO object from the file content
        docx_file = io.BytesIO(file_content)
        
        # Create Document object
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " "
                text_content += "\n"
        
        logger.info(f"Successfully extracted text from DOCX")
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


async def extract_text_from_resume(file_content: bytes, filename: str) -> str:
    """
    Extract text from resume file (PDF, DOCX, or DOC)
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return await extract_text_from_pdf(file_content)
    elif filename_lower.endswith('.docx'):
        return await extract_text_from_docx(file_content)
    elif filename_lower.endswith('.doc'):
        # For .doc files, we'll try to extract as text or raise an error
        # In production, you might want to use a library like python-docx2txt or convert DOC to DOCX
        try:
            return file_content.decode('utf-8', errors='ignore')
        except:
            raise HTTPException(status_code=400, detail="DOC file format is not fully supported. Please use PDF or DOCX format.")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please use PDF, DOC, or DOCX.")
        try:
            return file_content.decode('utf-8', errors='ignore')
        except:
            return "Unable to extract text from file"


async def fetch_web_page_content(url: str):
    """
    Fetch and extract text content from a web page URL
    """
    try:
        logger.info(f"Fetching web page content from: {url}")
        
        # Validate URL format
        if not url.startswith(('http://', 'https://')):
            logger.warning(f"Invalid URL format: {url}")
            return f"Invalid URL format: {url}"
        
        # Add headers to mimic a real browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
        }
        
        # Make the request with timeout
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Check if content is HTML
        content_type = response.headers.get('content-type', '').lower()
        if 'text/html' not in content_type:
            logger.warning(f"Non-HTML content type: {content_type}")
            return f"Non-HTML content detected: {content_type}"
        
        # Parse the HTML content
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Extract text content
        text_content = soup.get_text()
        
        # Clean up the text
        lines = (line.strip() for line in text_content.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        # Check if we got meaningful content
        if len(text.strip()) < 50:
            logger.warning(f"Very little content extracted: {len(text)} characters")
            return f"Very little content extracted from the page. This might not be a valid job posting."
        
        # Limit content length to avoid token limits (keep first 8000 characters)
        if len(text) > 4000:
            text = text[:4000] + "... [content truncated]"
        
        logger.info(f"Successfully extracted {len(text)} characters from web page")
        return text
        
    except requests.exceptions.Timeout:
        logger.error(f"Timeout fetching web page: {url}")
        return f"Timeout fetching web page content from: {url}"
    except requests.exceptions.ConnectionError:
        logger.error(f"Connection error fetching web page: {url}")
        return f"Connection error fetching web page content from: {url}"
    except requests.exceptions.HTTPError as e:
        logger.error(f"HTTP error fetching web page: {str(e)}")
        return f"HTTP error fetching web page content: {str(e)}"
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error fetching web page: {str(e)}")
        return f"Request error fetching web page content: {str(e)}"
    except ImportError as e:
        logger.error(f"Missing dependency: {str(e)}")
        return f"Missing dependency for web scraping: {str(e)}"
    except Exception as e:
        logger.error(f"Unexpected error parsing web page content: {str(e)}")
        return f"Unexpected error parsing web page content: {str(e)}"


async def analyze_target_job_with_openai(user_id: str, user_data: dict = None):
    """
    Analyze target job using OpenAI GPT-5-mini with stored job data and compare with user data
    """
    try:
        logger.info(f"Analyzing target job with OpenAI for user_id: {user_id}")
        
        # Get stored latest analysis item from DynamoDB
        stored_analysis = await get_saved_job_analysis(user_id)
        if not stored_analysis:
            raise Exception(f"No job analysis found for user_id: {user_id}")
        
        # Extract job requirements from stored data (career_analysis_data latest item)
        analysis_data = stored_analysis.get('job_analysis', {})
        job_requirements = f"""
        JOB REQUIREMENTS:
        Title: {analysis_data.get('standardized_title', 'Unknown')}
        Company: {analysis_data.get('company_name', 'Unknown')}
        Technical Skills: {', '.join(analysis_data.get('technical_skills', []))}
        Soft Skills: {', '.join(analysis_data.get('soft_skills', []))}
        Industry: {analysis_data.get('industry', 'Unknown')}
        Experience Level: {analysis_data.get('experience_level', 'Unknown')}
        Key Responsibilities: {', '.join(analysis_data.get('key_responsibilities', []))}
        Salary Range: {analysis_data.get('salary_range', 'Not specified')}
        """
        
        # Build user profile summary
        user_profile = ""
        if user_data:
            user_profile = f"""
            USER PROFILE:
            Name: {user_data.get('firstName', '')} {user_data.get('lastName', '')}
            Education: {user_data.get('degree', '')} in {user_data.get('major', '')} from {user_data.get('collegeName', '')} ({user_data.get('graduationYear', '')})
            Technical Skills:
            - Programming Languages: {user_data.get('programmingLanguages', '')}
            - Frameworks: {user_data.get('frameworks', '')}
            - Databases: {user_data.get('databases', '')}
            - Tools: {user_data.get('tools', '')}
            Work Experience: {len(user_data.get('workExperiences', []))} positions
            """

            # Add work experience details
            work_experiences = user_data.get('workExperiences', [])
            if work_experiences:
                user_profile += "\nWork Experience Details:\n"
                for i, exp in enumerate(work_experiences, 1):
                    user_profile += f"{i}. {exp.get('title', '')} at {exp.get('company', '')} ({exp.get('startDate', '')} - {exp.get('endDate', 'Present')})\n"
                    user_profile += f"   Description: {exp.get('description', '')}\n"

        # Create the comprehensive prompt for comparison and scoring
        prompt = f"""
        Compare the user's profile against the job requirements and provide detailed analysis:

        {job_requirements}

        {user_profile}

        ANALYSIS REQUIREMENTS:
        2. Compare user profile against job requirements
        3. Score each aspect (1-10) with specific strengths and 2-5 improvement recommendations:
           - Background: Overall professional background alignment
           - Education: Degree relevance, school prestige, academic achievements
           - Professional: Work experience relevance, career progression, industry experience
           - Technical Skills: Programming languages, frameworks, tools alignment
           - Teamwork: Collaboration skills, team leadership, communication
           - Job Match: Overall fit for this specific position

        4. Provide overall score (1-100) representing job match percentage

        For each improvement recommendation, be specific and actionable (e.g., "Learn React.js through online courses" instead of "Improve frontend skills").
        """
        
        # Call OpenAI GPT-5-mini with structured output using responses.parse
        response = client.responses.parse(
            model="gpt-5-mini",
            input=[
                {
                    "role": "system",
                    "content": "You are a career analyst with access to current job market data. Analyze job positions and provide detailed information about required skills, industry trends, and market insights. Compare user profiles against job requirements and provide specific, actionable improvement recommendations."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            text_format=JobMatchAnalysis
        )
        
        # Extract the parsed structured analysis
        analysis_data = response.output_parsed

        logger.info(f"OpenAI analysis completed for user_id: {user_id}")
        logger.info(f"Analysis result: {analysis_data}")

        return {
            "success": True,
            "analysis": analysis_data.dict(),
            "structured_output": True
        }
        
    except Exception as e:
        logger.error(f"Error in OpenAI job analysis: {str(e)}")
        # Create default AspectScore for error case
        default_aspect = {
            "score": 1,
            "strengths": ["Analysis could not be completed"],
            "improvements": ["Please try again later", "Ensure job URL is accessible"]
        }

        return {
            "success": False,
            "error": str(e),
            "analysis": {
                "standardized_title": "Unknown",
                "technical_skills": [],
                "soft_skills": [],
                "industry": "Unknown",
                "experience_level": "Unknown",
                "key_responsibilities": [],
                "salary_range": "Unknown",
                "company_name": "Unknown",
                "background_score": default_aspect,
                "education_score": default_aspect,
                "professional_score": default_aspect,
                "tech_skills_score": default_aspect,
                "teamwork_score": default_aspect,
                "job_match_score": default_aspect,
                "overall_score": 1,
                "analysis_confidence": "low"
            }
        }


async def parse_resume_with_openai(resume_content: str):
    """
    Parse resume content using OpenAI GPT-5-mini as a file parser, ATS checker, and job recruiter
    """
    try:
        logger.info("Starting resume parsing with OpenAI...")
        
        prompt = f"""
        You are an expert resume parser, ATS (Application Tracking System) checker, and job recruiter. 
        Parse the following resume content and extract structured information.

        RESUME CONTENT:
        {resume_content}

        EXTRACTION REQUIREMENTS:
        1. Background Information:
           - Full name
           - Email address
           - Home address
           - Personal website

        2. Education (extract ALL education entries):
           - College/university name
           - Degree type
           - Major/field of study
           - Location
           - Duration (start and end dates)
           - GPA (if mentioned)
           - Relevant coursework

        3. Professional Experience (extract ALL work entries):
           - Company name
           - Job title
           - Location
           - Duration (start and end dates)
           - Technologies/tools used in projects

        4. Technical Skills:
           - Domain knowledge areas
           - Technologies mentioned
           - Tools used
           - Software applications

        5. Teamwork:
           - Teamwork experiences mentioned
           - Leadership experiences mentioned

        6. ATS Review:
           - Formatting issues that could affect ATS parsing
           - Syntax issues in the document
           - ATS compatibility score (0-100)

        IMPORTANT:
        - Extract ALL instances of education and work experience (not just the first one)
        - If information is not available, leave as null/empty
        - Be thorough in identifying all technologies, tools, and skills mentioned
        - Focus on ATS compatibility issues that could prevent resume parsing
        - Provide specific, actionable feedback for ATS optimization
        """
        
        # Call OpenAI with structured output for resume parsing
        response = client.responses.parse(
            model="gpt-5-mini",
            input=[
                {
                    "role": "system",
                    "content": "You are an expert resume parser, ATS specialist, and technical recruiter with deep knowledge of resume formats, applicant tracking systems, and technical hiring. Extract comprehensive information from resumes with high accuracy and attention to ATS compatibility."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            text_format=ResumeParsedData
        )
        
        # Extract the parsed structured data
        parsed_data = response.output_parsed
        
        logger.info(f"Resume parsing completed successfully")
        logger.info(f"Resume parsed data: {parsed_data}")
        
        return {
            "success": True,
            "parsed_data": parsed_data.dict(),
            "structured_output": True
        }
        
    except Exception as e:
        logger.error(f"Error in resume parsing: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "parsed_data": {
                "background": {
                    "full_name": None,
                    "email_address": None,
                    "home_address": None,
                    "personal_website": None
                },
                "education": [],
                "professional": [],
                "technical_skills": {
                    "domain_knowledge": [],
                    "technologies": [],
                    "tools": [],
                    "software": []
                },
                "teamwork": {
                    "teamwork_experience": [],
                    "leadership_experience": []
                },
                "ats_review": {
                    "formatting_issues": ["Resume parsing failed"],
                    "syntax_issues": ["Unable to parse resume"],
                    "ats_compatibility_score": 0
                }
            }
        }


async def analyze_resume_with_openai(resume_parsed_data: dict, job_requirements: str):
    """
    Analyze parsed resume data using OpenAI GPT-5-mini as a technical recruiter/HR
    """
    try:
        logger.info("Starting resume analysis with OpenAI...")
        
        prompt = f"""
        You are an experienced technical recruiter and HR professional writing in a friendly, supportive tone. 
        Analyze the following parsed resume data against the job requirements and provide comprehensive feedback with constructive, encouraging language.

        JOB REQUIREMENTS:
        {job_requirements}

        PARSED RESUME DATA:
        {resume_parsed_data}

        ANALYSIS REQUIREMENTS:
        For each section (Background, Education, Professional, Technical Skills, Teamwork, ATS), provide:
        1. Score (1-10): How well does this section align with job requirements?
        2. Improvement suggestions: Specific, actionable advice to improve this section, phrased positively and encouragingly.

        Focus on:
        - How well the resume content matches the job requirements
        - Missing information that would strengthen the application
        - Areas where the candidate could better highlight relevant experience
        - ATS optimization opportunities
        - Specific improvements to make the resume more competitive

        Provide specific, actionable feedback that would help the candidate improve their resume for this specific role.
        Use a friendly tone. Keep suggestions concise, clear, and encouraging.
        Only include non-empty, relevant improvement suggestions.
        """
        
        # Call OpenAI with structured output for resume analysis
        response = client.responses.parse(
            model="gpt-5-mini",
            input=[
                {
                    "role": "system",
                    "content": "You are an expert technical recruiter and HR professional with deep knowledge of resume optimization and technical hiring. Analyze parsed resume data with the perspective of both human recruiters and ATS systems to provide comprehensive, actionable feedback."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            text_format=ResumeAnalysis
        )
        
        # Extract the parsed structured analysis
        analysis_data = response.output_parsed
        
        logger.info(f"Resume analysis completed successfully")
        logger.info(f"Resume analysis result: {analysis_data}")
        
        return {
            "success": True,
            "analysis": analysis_data.dict(),
            "structured_output": True
        }
        
    except Exception as e:
        logger.error(f"Error in resume analysis: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "analysis": {
                "background_score": 1,
                "education_score": 1,
                "professional_score": 1,
                "technical_skills_score": 1,
                "teamwork_score": 1,
                "ats_score": 1,
                "overall_score": 1,
                "background_improvements": ["Analysis could not be completed"],
                "education_improvements": ["Unable to analyze education"],
                "professional_improvements": ["Unable to analyze professional experience"],
                "technical_skills_improvements": ["Unable to analyze technical skills"],
                "teamwork_improvements": ["Unable to analyze teamwork"],
                "ats_improvements": ["Unable to analyze ATS compatibility"],
                "general_improvements": ["Please try again later"]
            }
        }


# The 1st alph aAPI
@app.post("/alpha_target_job_analysis")
async def alpha_target_job_analysis(
    target_job: str = Form(...),
    form_data: str = Form(...),
    user_id: str = Form(...)
):
    try:
        logger.info("Received target job analysis request")
        
        # Parse form data
        import json
        form_data_parsed = json.loads(form_data) if form_data else {}
        
        # Check if target_job is a URL or direct job description
        is_url = target_job.startswith(('http://', 'https://'))
        
        if is_url:
            # Fetch web page content if it's a URL
            logger.info(f"Detected URL, fetching web page content: {target_job}")
            web_content = await fetch_web_page_content(target_job)
            job_content = web_content
        else:
            # Use target_job as direct job description
            logger.info("Detected direct job description")
            job_content = target_job
        
        # Build user profile summary
        user_profile = ""
        if form_data_parsed:
            user_profile = f"""
            USER PROFILE:
            Name: {form_data_parsed.get('firstName', '')} {form_data_parsed.get('lastName', '')}
            Education: {form_data_parsed.get('degree', '')} in {form_data_parsed.get('major', '')} from {form_data_parsed.get('collegeName', '')} ({form_data_parsed.get('graduationYear', '')})
            Technical Skills:
            - Programming Languages: {form_data_parsed.get('programmingLanguages', '')}
            - Frameworks: {form_data_parsed.get('frameworks', '')}
            - Databases: {form_data_parsed.get('databases', '')}
            - Tools: {form_data_parsed.get('tools', '')}
            Work Experience: {len(form_data_parsed.get('workExperiences', []))} positions
            """

            # Add work experience details
            work_experiences = form_data_parsed.get('workExperiences', [])
            if work_experiences:
                user_profile += "\nWork Experience Details:\n"
                for i, exp in enumerate(work_experiences, 1):
                    user_profile += f"{i}. {exp.get('title', '')} at {exp.get('company', '')} ({exp.get('startDate', '')} - {exp.get('endDate', 'Present')})\n"
                    user_profile += f"   Description: {exp.get('description', '')}\n"

        # Create the comprehensive prompt
        prompt = f"""
        Analyze the following job posting and extract key information:

        JOB POSTING:
        URL: "{target_job}"
        Content: {job_content}

        {user_profile}

        ANALYSIS REQUIREMENTS:
        Extract job details from the posting content:
           - standardized_title: Clean, professional job title
           - company_name: Extract the company/organization name from the job posting
           - technical_skills: Required programming languages, frameworks, tools
           - soft_skills: Communication, leadership, teamwork requirements
           - industry: Business sector (e.g., Technology, Finance, Healthcare)
           - experience_level: Entry, Mid, Senior level
           - key_responsibilities: Main job duties and tasks
           - salary_range: Compensation if mentioned

        IMPORTANT: Pay special attention to extracting the actual company name from the job posting content.
        """
        
        # Call OpenAI GPT-5-mini with structured output
        response = client.responses.parse(
            model="gpt-5-mini",
            input=[
                {
                    "role": "system",
                    "content": "You are a career analyst with access to current job market data. Analyze job positions and extract key information including company details, required skills, industry classification, and experience levels."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            text_format=JobExtraction
        )
        
        # Extract the parsed structured analysis
        analysis_data = response.output_parsed
        
        # Use the provided user_id from frontend
        email = form_data_parsed.get('email', '')
        first_name = form_data_parsed.get('firstName', '')
        last_name = form_data_parsed.get('lastName', '')
        
        logger.info(f"Using provided user_id: {user_id}")
        
        # Save the analyzed data to DynamoDB (career_analysis_data)
        current_date = datetime.now().strftime("%Y-%m-%d")
        timestamp = datetime.now().isoformat()
        
        async def save_job_analysis_to_career_table():
            try:
                dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
                table_name = 'career_analysis_data'
                
                # Try to get the table - if it doesn't exist, create it
                try:
                    table = dynamodb.Table(table_name)
                    # Try to describe the table to see if it exists
                    table.load()
                except dynamodb.meta.client.exceptions.ResourceNotFoundException:
                    logger.info(f"Table {table_name} does not exist, creating it...")
                    table = dynamodb.create_table(
                        TableName=table_name,
                        KeySchema=[
                            {'AttributeName': 'user_id', 'KeyType': 'HASH'},
                            {'AttributeName': 'timestamp', 'KeyType': 'RANGE'}
                        ],
                        AttributeDefinitions=[
                            {'AttributeName': 'user_id', 'AttributeType': 'S'},
                            {'AttributeName': 'timestamp', 'AttributeType': 'S'}
                        ],
                        BillingMode='PAY_PER_REQUEST'
                    )
                    table.wait_until_exists()
                    logger.info(f"Table {table_name} created successfully")
                except Exception as table_error:
                    # If we can't describe the table due to permissions, assume it exists and try to use it
                    logger.warning(f"Could not check table status (permissions issue): {str(table_error)}")
                    table = dynamodb.Table(table_name)
                # Insert a new item for each new job_analysis
                item = {
                    'user_id': user_id,
                    'timestamp': timestamp,
                    'date': current_date,
                    'target_job': target_job,
                    'is_url': is_url,
                    'first_name': first_name,
                    'last_name': last_name,
                    'email_address': email,
                    'phone_number': form_data_parsed.get('phoneNumber', ''),
                    # store object fields
                    'job_analysis': analysis_data.dict(),
                    'capability_analysis': None,
                    'resume_analysis': None
                }
                table.put_item(Item=item)
                logger.info(f"Saved job_analysis item for user_id={user_id} ts={timestamp}")
            except Exception as e:
                logger.error(f"Error saving career_analysis_data job_analysis: {str(e)}")
        
        asyncio.create_task(save_job_analysis_to_career_table())
        
        logger.info(f"Target job analysis completed successfully: {user_id}")
        
        return {
            "status": "success",
            "user_id": user_id,
            "is_url": is_url,
            "analysis_data": analysis_data.dict(),
            "message": "Job analysis completed and saved successfully"
        }
        
    except Exception as e:
        logger.error(f"Error in alpha_target_job_analysis: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error analyzing target job: {str(e)}")


# ==================== Overall Analysis Endpoints ====================

async def personal_capability_analysis(user_id: str, target_job_data: dict, knowledge_scope_tags: List[str]) -> PersonalCapabilityAnalysis:
    """
    Analyze personal capability against target job requirements.
    Fetches user profile and knowledge scope from DynamoDB, then uses OpenAI for analysis.
    """
    try:
        logger.info(f"Starting personal capability analysis for user_id: {user_id}")
        
        # Initialize DynamoDB
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        table_name = 'ambit-dashboard-application-data'
        table = dynamodb.Table(table_name)
        
        # Fetch user profile
        profile_response = table.get_item(
            Key={'PK': user_id, 'SK': 'PROFILE#MAIN'}
        )
        profile_data = profile_response.get('Item', {}).get('data', {}) if 'Item' in profile_response else {}
        
        # Derive career focus from profile (defaults to software-engineering)
        career_focus = (profile_data.get('careerFocus', '') or 'software-engineering').strip() or 'software-engineering'

        def _get_knowledge_data(sk_prefix: str, legacy_flat_key: str) -> dict:
            """Fetch knowledge data for the user's career focus with legacy fallback."""
            # 1. Try career-focus-specific item (new format)
            cf_response = table.get_item(Key={'PK': user_id, 'SK': f'{sk_prefix}#{career_focus}'})
            if 'Item' in cf_response:
                return cf_response['Item'].get('data', {})
            # 2. Legacy fallback (only for software-engineering)
            if career_focus == 'software-engineering':
                legacy_response = table.get_item(Key={'PK': user_id, 'SK': sk_prefix})
                if 'Item' in legacy_response:
                    raw = legacy_response['Item'].get('data', {})
                    # Flat legacy format
                    if legacy_flat_key in raw:
                        return raw
                    # Nested interim format
                    return raw.get('software-engineering', {})
            return {}

        # Fetch knowledge scope data based on tags
        knowledge_content = []
        if 'Established Expertise' in knowledge_scope_tags:
            established_data = _get_knowledge_data('KNOWLEDGE#ESTABLISHED', 'personal_project')
            if established_data:
                knowledge_content.append(('Established Expertise', established_data))

        if 'Expanding Knowledge Base' in knowledge_scope_tags:
            expanding_data = _get_knowledge_data('KNOWLEDGE#EXPANDING', 'future_personal_project')
            if expanding_data:
                knowledge_content.append(('Expanding Knowledge Base', expanding_data))
        
        # Extract plain text content from knowledge scope
        knowledge_text = ""
        for scope_name, scope_data in knowledge_content:
            knowledge_text += f"\n{scope_name}:\n"
            
            # Personal projects
            personal_projects = scope_data.get('personal_project', []) if scope_name == 'Established Expertise' else scope_data.get('future_personal_project', [])
            for project in personal_projects:
                desc = project.get('projectDescription', {})
                if isinstance(desc, dict):
                    knowledge_text += f"Personal Project: {project.get('projectName', '')}\n"
                    knowledge_text += f"  Overview: {desc.get('overview', '')}\n"
                    knowledge_text += f"  Tech & Teamwork: {desc.get('techAndTeamwork', '')}\n"
                    knowledge_text += f"  Achievement: {desc.get('achievement', '')}\n"
                    knowledge_text += f"  Technologies: {', '.join(project.get('selectedTechnologies', []))}\n"
                    knowledge_text += f"  Frameworks: {', '.join(project.get('selectedFrameworks', []))}\n"
            
            # Professional projects
            professional_projects = scope_data.get('professional_project', []) if scope_name == 'Established Expertise' else scope_data.get('future_professional_project', [])
            for project in professional_projects:
                desc = project.get('projectDescription', {})
                if isinstance(desc, dict):
                    knowledge_text += f"Professional Project: {project.get('projectName', '')}\n"
                    knowledge_text += f"  Overview: {desc.get('overview', '')}\n"
                    knowledge_text += f"  Tech & Teamwork: {desc.get('techAndTeamwork', '')}\n"
                    knowledge_text += f"  Achievement: {desc.get('achievement', '')}\n"
                    knowledge_text += f"  Technologies: {', '.join(project.get('selectedTechnologies', []))}\n"
                    knowledge_text += f"  Frameworks: {', '.join(project.get('selectedFrameworks', []))}\n"
            
            # Technical skills
            tech_skills = scope_data.get('technical_skills', {}) if scope_name == 'Established Expertise' else scope_data.get('future_technical_skills', {})
            selected_skills = tech_skills.get('selectedSkills', [])
            custom_keywords = tech_skills.get('customKeywords', {})
            knowledge_text += f"Technical Skills: {', '.join(selected_skills)}\n"
            for topic, keywords in custom_keywords.items():
                knowledge_text += f"  {topic}: {', '.join(keywords)}\n"
        
        # Build user profile summary
        profile_summary = f"""
USER PROFILE:
Career Focus: {profile_data.get('careerFocus', 'Not specified')}
Education: {json.dumps(profile_data.get('education', []), indent=2)}
Professional Experience: {json.dumps(profile_data.get('professional', {}), indent=2)}
"""
        
        # Build job requirements summary
        job_summary = f"""
TARGET JOB REQUIREMENTS:
Title: {target_job_data.get('target_job_title', 'Unknown')}
Company: {target_job_data.get('target_job_company', 'Unknown')}
Description: {target_job_data.get('target_job_description', '')}
Required Skills: {', '.join(target_job_data.get('target_job_skill_keywords', []))}
"""
        
        # Build OpenAI prompt
        prompt = f"""
You are a professional career analyst. Analyze the user's personal capability against the target job requirements.

{profile_summary}

KNOWLEDGE SCOPE CONTENT:
{knowledge_text}

{job_summary}

ANALYSIS REQUIREMENTS:
Evaluate the user's fit across 6 dimensions (score each 1-10):
1. Background: Check if user's career focus, education level, and working experience match basic job requirements
2. Education: Check if user's education level, degree, and major relate to job requirements. Also evaluate how much personal projects and technical skills under selected knowledge scope match job required experience and skills
3. Professional: Check if user's work experience and achievements match basic job requirements. Also evaluate how much professional projects and technical skills under selected knowledge scope match job required experience and technical skills
4. Tech Skills: Check if user's technical skills under knowledge scope (from all personal projects, professional projects, and technical skills) match job required skills
5. Teamwork: Check if user's personal and professional projects under selected knowledge scope mention enough teamwork experience, including teamwork and leadership, and how much it matches target job requirements
6. Job Match: Calculate the average of the above 5 scores

For each dimension, provide:
- A score from 1-10
- At least 3 and at most 6 professional career advice items (each 15-30 words, actionable and specific)

Return the analysis in the following JSON structure:
{{
  "background_score": <int 1-10>,
  "education_score": <int 1-10>,
  "professional_score": <int 1-10>,
  "tech_skills_score": <int 1-10>,
  "teamwork_score": <int 1-10>,
  "job_match_score": <int 1-10 (average of above 5)>,
  "background_advice": [<list of 4+ strings, each 15-30 words>],
  "education_advice": [<list of 4+ strings, each 15-30 words>],
  "professional_advice": [<list of 4+ strings, each 15-30 words>],
  "tech_skills_advice": [<list of 4+ strings, each 15-30 words>],
  "teamwork_advice": [<list of 4+ strings, each 15-30 words>]
}}
"""
        
        # Call OpenAI with structured output
        try:
            response = client.beta.chat.completions.parse(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a professional career analyst. Provide detailed, actionable analysis and advice."},
                    {"role": "user", "content": prompt}
                ],
                response_format=PersonalCapabilityAnalysis
            )
            
            analysis_result = response.choices[0].message.parsed
            if not analysis_result:
                # Fallback: parse from JSON if parsed is None
                parsed_json = json.loads(response.choices[0].message.content)
                analysis_result = PersonalCapabilityAnalysis(**parsed_json)
        except Exception as e:
            logger.error(f"OpenAI API error: {str(e)}")
            # Fallback: use regular chat completion and parse JSON
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a professional career analyst. Provide detailed, actionable analysis and advice. Return only valid JSON."},
                    {"role": "user", "content": prompt + "\n\nReturn the result as valid JSON only."}
                ],
                response_format={"type": "json_object"}
            )
            parsed_json = json.loads(response.choices[0].message.content)
            analysis_result = PersonalCapabilityAnalysis(**parsed_json)
        
        logger.info(f"Personal capability analysis completed for user_id: {user_id}")
        return analysis_result
        
    except Exception as e:
        logger.error(f"Error in personal_capability_analysis: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise


async def resume_power_analysis(resume_file: UploadFile, target_job_data: dict) -> ResumePowerAnalysis:
    """
    Analyze resume power against target job requirements.
    Parses resume file, extracts structured data, performs ATS check, then uses OpenAI for analysis.
    """
    try:
        logger.info(f"Starting resume power analysis")
        
        # Read and parse resume file
        file_content = await resume_file.read()
        resume_text = await extract_text_from_resume(file_content, resume_file.filename or 'resume.pdf')
        
        # Extract structured data from resume using OpenAI
        extraction_prompt = f"""
Extract structured information from the following resume text:

{resume_text}

Extract and return:
1. Background: full_name, email_address, home_address, personal_website
2. Education: List of education entries with college_name, degree, major, location, duration, gpa, coursework
3. Professional: List of professional experiences with company_name, job_title, location, duration, technologies_used
4. Technical Skills: Extract all technical skills keywords from projects and skills sections, categorize into domain_knowledge, technologies, tools, software
5. Teamwork: Extract teamwork_experience and leadership_experience mentions

Also perform ATS compatibility check:
- Check for formatting issues
- Check for syntax issues
- Check if basic information is present: email, address, phone, education
- Provide ATS compatibility score (0-100)

Return structured JSON with all extracted information.
"""
        
        # First extraction to get structured data
        try:
            extraction_response = client.beta.chat.completions.parse(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a resume parser. Extract structured information from resumes accurately."},
                    {"role": "user", "content": extraction_prompt}
                ],
                response_format=ResumeParsedData
            )
            
            parsed_resume = extraction_response.choices[0].message.parsed
            if not parsed_resume:
                # Fallback: parse from JSON if parsed is None
                parsed_json = json.loads(extraction_response.choices[0].message.content)
                parsed_resume = ResumeParsedData(**parsed_json)
        except Exception as e:
            logger.error(f"OpenAI extraction API error: {str(e)}")
            # Fallback: use regular chat completion and parse JSON
            extraction_response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a resume parser. Extract structured information from resumes accurately. Return only valid JSON."},
                    {"role": "user", "content": extraction_prompt + "\n\nReturn the result as valid JSON only."}
                ],
                response_format={"type": "json_object"}
            )
            parsed_json = json.loads(extraction_response.choices[0].message.content)
            parsed_resume = ResumeParsedData(**parsed_json)
        
        # Build analysis prompt
        analysis_prompt = f"""
You are a professional resume analyst. Analyze the extracted resume data against the target job requirements.

EXTRACTED RESUME DATA:
Background: {parsed_resume.background.dict()}
Education: {json.dumps([e.dict() for e in parsed_resume.education], indent=2)}
Professional: {json.dumps([p.dict() for p in parsed_resume.professional], indent=2)}
Technical Skills: {parsed_resume.technical_skills.dict()}
Teamwork: {parsed_resume.teamwork.dict()}
ATS Review: {parsed_resume.ats_review.dict()}

TARGET JOB REQUIREMENTS:
Title: {target_job_data.get('target_job_title', 'Unknown')}
Company: {target_job_data.get('target_job_company', 'Unknown')}
Description: {target_job_data.get('target_job_description', '')}
Required Skills: {', '.join(target_job_data.get('target_job_skill_keywords', []))}

ANALYSIS REQUIREMENTS:
Evaluate the resume's fit across 6 dimensions (score each 1-10):
1. Background: Check if resume's education level and working experience match basic job requirements
2. Education: Check if resume's education level, degree, and major relate to job requirements
3. Professional: Check if resume's work experience and achievements match basic job requirements. Also evaluate how much professional projects and technical skills match job required experience and technical skills
4. Tech Skills: Check if extracted resume's technical skill keywords match job required skills
5. Teamwork: Check if extracted resume teamwork and leadership experience matches target job requirements
6. Job Match: Calculate the average of the above 5 scores

For each dimension, provide:
- A score from 1-10
- At least 3 and at most 6 professional resume improvement advice items (each 15-30 words, actionable and specific)

Also include:
- ATS compatibility score (0-100) from the ATS review
- List of ATS issues (formatting, syntax, missing info)

Return the analysis in the following JSON structure:
{{
  "background_score": <int 1-10>,
  "education_score": <int 1-10>,
  "professional_score": <int 1-10>,
  "tech_skills_score": <int 1-10>,
  "teamwork_score": <int 1-10>,
  "job_match_score": <int 1-10 (average of above 5)>,
  "ats_compatibility_score": <int 0-100>,
  "background_advice": [<list of 4+ strings, each 15-30 words>],
  "education_advice": [<list of 4+ strings, each 15-30 words>],
  "professional_advice": [<list of 4+ strings, each 15-30 words>],
  "tech_skills_advice": [<list of 4+ strings, each 15-30 words>],
  "teamwork_advice": [<list of 4+ strings, each 15-30 words>],
  "ats_issues": [<list of strings describing ATS issues>]
}}
"""
        
        # Call OpenAI for analysis
        try:
            analysis_response = client.beta.chat.completions.parse(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a professional resume analyst. Provide detailed, actionable analysis and improvement advice."},
                    {"role": "user", "content": analysis_prompt}
                ],
                response_format=ResumePowerAnalysis
            )
            
            analysis_result = analysis_response.choices[0].message.parsed
            if not analysis_result:
                # Fallback: parse from JSON if parsed is None
                parsed_json = json.loads(analysis_response.choices[0].message.content)
                analysis_result = ResumePowerAnalysis(**parsed_json)
        except Exception as e:
            logger.error(f"OpenAI analysis API error: {str(e)}")
            # Fallback: use regular chat completion and parse JSON
            analysis_response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a professional resume analyst. Provide detailed, actionable analysis and improvement advice. Return only valid JSON."},
                    {"role": "user", "content": analysis_prompt + "\n\nReturn the result as valid JSON only."}
                ],
                response_format={"type": "json_object"}
            )
            parsed_json = json.loads(analysis_response.choices[0].message.content)
            analysis_result = ResumePowerAnalysis(**parsed_json)
        
        # Merge ATS issues from parsed resume
        if parsed_resume.ats_review.formatting_issues:
            analysis_result.ats_issues.extend(parsed_resume.ats_review.formatting_issues)
        if parsed_resume.ats_review.syntax_issues:
            analysis_result.ats_issues.extend(parsed_resume.ats_review.syntax_issues)
        
        logger.info(f"Resume power analysis completed")
        return analysis_result
        
    except Exception as e:
        logger.error(f"Error in resume_power_analysis: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise


@app.post("/overall_analysis")
async def overall_analysis(
    target_job_data: str = Form(...),
    knowledge_scope: str = Form(...),
    resume_file: UploadFile = File(...),
    user_id: str = Form(...)
):
    """
    Main endpoint for overall analysis.
    Launches two parallel analyses: personal capability and resume power.
    """
    try:
        logger.info(f"Received overall_analysis request for user_id: {user_id}")
        
        # Parse form data
        target_job_dict = json.loads(target_job_data)
        knowledge_scope_tags = json.loads(knowledge_scope)

        # Check plan and usage
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        table = dynamodb.Table('ambit-dashboard-application-data')
        timestamp = datetime.utcnow().isoformat()

        sub_response = table.get_item(Key={'PK': user_id, 'SK': 'SUBSCRIPTION'})
        plan = sub_response.get('Item', {}).get('plan', 'free')

        usage_response = table.get_item(Key={'PK': user_id, 'SK': 'USAGE'})
        usage_item = usage_response.get('Item')

        if not usage_item:
            usage_item = {
                'PK': user_id,
                'SK': 'USAGE',
                'craft_count': 0,
                'analysis_count': 0,
                'download_count': 0,
                'createdAt': timestamp,
                'updatedAt': timestamp
            }
            table.put_item(Item=usage_item)

        analysis_count = int(usage_item.get('analysis_count', 0))

        if plan == 'free' and analysis_count >= 3:
            return {"status": "error", "error_code": "ANALYSIS_LIMIT_EXCEEDED", "message": "Free plan analysis limit reached."}

        table.update_item(
            Key={'PK': user_id, 'SK': 'USAGE'},
            UpdateExpression='SET analysis_count = analysis_count + :inc, updatedAt = :ts',
            ExpressionAttributeValues={':inc': 1, ':ts': timestamp}
        )

        # Reset file pointer for resume file (it may have been read already)
        await resume_file.seek(0)

        # Launch both analyses in parallel
        personal_task = personal_capability_analysis(user_id, target_job_dict, knowledge_scope_tags)
        resume_task = resume_power_analysis(resume_file, target_job_dict)
        
        # Wait for both to complete
        personal_result, resume_result = await asyncio.gather(personal_task, resume_task)
        
        logger.info(f"Overall analysis completed successfully for user_id: {user_id}")
        
        return {
            "status": "success",
            "personal_capability": personal_result.dict(),
            "resume_power": resume_result.dict()
        }
        
    except Exception as e:
        logger.error(f"Error in overall_analysis: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error performing overall analysis: {str(e)}")


# Helper to convert all floats in nested structures to Decimal for DynamoDB
def convert_floats_to_decimal(value):
    if isinstance(value, float):
        # Convert Python float to Decimal via string to preserve precision
        return Decimal(str(value))
    if isinstance(value, list):
        return [convert_floats_to_decimal(v) for v in value]
    if isinstance(value, dict):
        return {k: convert_floats_to_decimal(v) for k, v in value.items()}
    return value


async def get_saved_job_analysis(user_id: str):
    """
    Retrieve saved job analysis data from DynamoDB
    """
    try:
        # Initialize DynamoDB client
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        table_name = 'career_analysis_data'
        table = dynamodb.Table(table_name)
        # Query latest item for this user_id
        response = table.query(
            KeyConditionExpression=Key('user_id').eq(user_id),
            ScanIndexForward=False,
            Limit=1
        )
        items = response.get('Items', [])
        if items:
            logger.info(f"Retrieved latest analysis item for user_id: {user_id}")
            return items[0]
        logger.warning(f"No analysis items found for user_id: {user_id}")
        return None
        
    except Exception as e:
        logger.error(f"Error retrieving saved job analysis {user_id}: {str(e)}")
        return None


@app.get("/get_job_analysis/{user_id}")
async def get_job_analysis(user_id: str):
    """
    Retrieve saved job analysis data by user_id
    """
    try:
        logger.info(f"Retrieving job analysis: {user_id}")
        
        saved_analysis = await get_saved_job_analysis(user_id)
        
        if saved_analysis:
            return {
                "status": "success",
                "analysis_data": saved_analysis
            }
        else:
            return {
                "status": "error",
                "message": "Analysis not found"
            }
            
    except Exception as e:
        logger.error(f"Error retrieving job analysis: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error retrieving job analysis: {str(e)}")


# The 3rd alph aAPI
@app.post("/alpha_resume_analysis")
async def alpha_resume_analysis(
    form_data: str = Form(...),
    user_id: str = Form(...),
    resume_file: UploadFile = File(...)
):
    """
    Analyze resume against job requirements using stored job analysis data
    """
    try:
        logger.info("Received Alpha Resume Analysis request")
        
        # Parse form data
        import json
        form_data_parsed = json.loads(form_data) if form_data else {}
        
        # Get stored job analysis data from DynamoDB
        stored_analysis = await get_saved_job_analysis(user_id)
        if not stored_analysis:
            raise HTTPException(status_code=404, detail=f"No job analysis found for user_id: {user_id}")
        
        # Extract job requirements from stored data (career_analysis_data latest item)
        analysis_data = stored_analysis.get('job_analysis', {})
        job_requirements = f"""
        Job Title: {analysis_data.get('standardized_title', 'Unknown')}
        Technical Skills Required: {', '.join(analysis_data.get('technical_skills', []))}
        Soft Skills Required: {', '.join(analysis_data.get('soft_skills', []))}
        Industry: {analysis_data.get('industry', 'Unknown')}
        Experience Level: {analysis_data.get('experience_level', 'Unknown')}
        Key Responsibilities: {', '.join(analysis_data.get('key_responsibilities', []))}
        """
        
        # Log resume file information
        resume_info = f"Resume file: {resume_file.filename} ({resume_file.content_type}, {resume_file.size} bytes)"
        logger.info(f"Resume file details: {resume_info}")
        
        # Analyze resume
        try:
            logger.info("Starting resume parsing and analysis...")
            resume_content = await resume_file.read()
            
            # Check if it's a PDF file and extract text properly
            if resume_file.content_type == 'application/pdf' or resume_file.filename.lower().endswith('.pdf'):
                logger.info("Detected PDF file, extracting text...")
                resume_text = await extract_text_from_pdf(resume_content)
            else:
                # For non-PDF files, try to decode as text
                resume_text = resume_content.decode('utf-8', errors='ignore')
            
            logger.info("Step 0: Extracted resume text content (first 200 chars): %s", resume_text[:200] if resume_text else "No content")
            
            # Step 1: Parse resume content
            logger.info("Step 1: Parsing resume content...")
            resume_parsing_result = await parse_resume_with_openai(resume_text)
            
            if not resume_parsing_result.get("success"):
                logger.error(f"Resume parsing failed: {resume_parsing_result.get('error')}")
                raise Exception(f"Resume parsing failed: {resume_parsing_result.get('error')}")
            
            # Step 2: Analyze parsed resume data
            logger.info("Step 2: Analyzing parsed resume data...")
            
            resume_analysis_result = await analyze_resume_with_openai(
                resume_parsing_result.get('parsed_data', {}), 
                job_requirements
            )
            
            # Combine parsing and analysis results
            resume_analysis_result['parsed_data'] = resume_parsing_result.get('parsed_data', {})
            
            # Calculate overall score and rating as average of all individual scores
            if resume_analysis_result.get('success') and resume_analysis_result.get('analysis'):
                analysis = resume_analysis_result['analysis']
                scores = [
                    analysis.get('background_score', 1),
                    analysis.get('education_score', 1),
                    analysis.get('professional_score', 1),
                    analysis.get('technical_skills_score', 1),
                    analysis.get('teamwork_score', 1),
                    analysis.get('ats_score', 1)
                ]
                
                # Calculate average score (1-10 scale)
                overall_score = round(sum(scores) / len(scores), 1)
                analysis['overall_score'] = overall_score
                
                # Calculate overall rating (1-10 scale, same as score)
                analysis['overall_resume_rating'] = overall_score
                
                logger.info(f"Calculated overall resume score: {overall_score} (average of: {scores})")
            
            logger.info("Resume parsing and analysis completed successfully")
            
        except Exception as e:
            logger.error(f"Error in resume analysis: {str(e)}")
            resume_analysis_result = {
                "success": False,
                "error": str(e),
                "analysis": {
                    "ats_score": 0,
                    "resume_strengths": ["Resume analysis failed"],
                    "resume_weaknesses": ["Unable to analyze resume"],
                    "missing_keywords": ["Analysis error"],
                    "formatting_issues": ["Unable to assess"],
                    "improvement_suggestions": ["Please try again"],
                    "overall_resume_rating": 1,
                    "recruiter_feedback": "Resume analysis encountered an error. Please try again."
                }
            }
        
        # Update resume_analysis into career_analysis_data latest item (ensure non-null)
        try:
            async def update_resume_analysis_to_career_table():
                try:
                    dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
                    table = dynamodb.Table('career_analysis_data')
                    # Find latest item for user
                    response = table.query(
                        KeyConditionExpression=Key('user_id').eq(user_id),
                        ScanIndexForward=False,
                        Limit=1
                    )
                    items = response.get('Items', [])
                    if not items:
                        logger.warning(f"No base item to update resume_analysis for user_id={user_id}")
                        return
                    ts = items[0]['timestamp']
                    # Convert floats to Decimal before saving to DynamoDB
                    safe_resume = convert_floats_to_decimal(resume_analysis_result or {"success": False, "analysis": None})
                    table.update_item(
                        Key={'user_id': user_id, 'timestamp': ts},
                        UpdateExpression='SET resume_analysis = :ra',
                        ExpressionAttributeValues={':ra': safe_resume}
                    )
                    logger.info(f"Updated resume_analysis for user_id={user_id} ts={ts}")
                except Exception as e:
                    logger.error(f"Error updating resume_analysis in career_analysis_data: {str(e)}")
            asyncio.create_task(update_resume_analysis_to_career_table())
        except Exception as e:
            logger.error(f"Error preparing resume analysis career update: {str(e)}")

        # Query DynamoDB for latest job_analysis data with retry logic
        job_analysis_data = None
        max_retries = 10  # Maximum 20 seconds wait
        retry_count = 0
        while retry_count < max_retries:
            try:
                dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
                table = dynamodb.Table('career_analysis_data')
                response = table.query(
                    KeyConditionExpression=Key('user_id').eq(user_id),
                    ScanIndexForward=False,
                    Limit=1
                )
                items = response.get('Items', [])
                if items and items[0].get('job_analysis'):
                    job_analysis_data = items[0]
                    logger.info(f"Found job analysis data for user_id: {user_id}")
                    break
                logger.info(f"No job analysis data found for user_id: {user_id}, retrying in 2 seconds...")
            except Exception as e:
                logger.error(f"Error querying job analysis data: {str(e)}")
            retry_count += 1
            if retry_count < max_retries:
                await asyncio.sleep(2)
        
        if not job_analysis_data:
            logger.warning(f"Could not retrieve job analysis data for user_id: {user_id} after {max_retries} retries")
            # Create default job analysis data
            job_analysis_data = {
                "analysis_data": {
                    "standardized_title": "Unknown",
                    "company_name": "Unknown",
                    "technical_skills": [],
                    "soft_skills": [],
                    "industry": "Unknown",
                    "experience_level": "Unknown",
                    "key_responsibilities": [],
                    "salary_range": "Unknown"
                }
            }
        
        # Extract job analysis information
        analysis_data = job_analysis_data.get('job_analysis', {})
        job_title = analysis_data.get('standardized_title', 'Unknown')
        company_name = analysis_data.get('company_name', 'Unknown')
        
        # Extract capability analysis information
        capability_analysis_data = job_analysis_data.get('capability_analysis', {})
        
        # Create job analysis result in the format expected by frontend
        job_analysis_result = {
            "success": True,
            "analysis": analysis_data,
            "structured_output": True
        }
        
        # Add resume analysis to the job analysis result
        if resume_analysis_result and resume_analysis_result.get('analysis'):
            job_analysis_result["analysis"]["resume_analysis"] = resume_analysis_result.get('analysis')
        
        return {
            "status": "success",
            "message": "Resume analysis completed successfully",
            "analysis_id": f"resume_analysis_{int(time.time())}",
            "completion_percentage": 100.0,  # Resume analysis is complete
            "resume_uploaded": True,
            "job_analysis": job_analysis_result,
            "resume_analysis": resume_analysis_result,
            "capability_analysis": capability_analysis_data,
            "data_summary": {
                "target_job": job_title,
                "job_title": job_title,
                "company_name": company_name,
                "resume_file": resume_info
            }
        }
        
    except Exception as e:
        logger.error(f"Error in alpha_resume_analysis: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error in resume analysis: {str(e)}")


# The 2nd alph aAPI
@app.post("/alpha_capability_analysis")
async def alpha_capability_analysis(
    target_job: str = Form(...),
    form_data: str = Form(...),
    user_id: str = Form(...)
):
    try:
        logger.info("Received Ambit Alpha analysis request")
        
        # Parse form data
        import json
        form_data_parsed = json.loads(form_data) if form_data else {}
        
        # Analyze target job with OpenAI using stored job data and user data
        logger.info("Starting OpenAI job analysis with stored job data and user data comparison...")
        try:
            job_analysis_result = await analyze_target_job_with_openai(user_id, form_data_parsed)
            logger.info("OpenAI analysis completed successfully")
        except Exception as e:
            logger.error(f"Error in OpenAI analysis: {str(e)}")
            # Create default AspectScore for error case
            default_aspect = {
                "score": 1,
                "strengths": ["Analysis could not be completed"],
                "improvements": ["Please try again later", "Ensure job analysis data is available"]
            }
            job_analysis_result = {
                "success": False,
                "error": str(e),
                "analysis": {
                    "standardized_title": "Unknown",
                    "technical_skills": [],
                    "soft_skills": [],
                    "industry": "Unknown",
                    "experience_level": "Unknown",
                    "key_responsibilities": [],
                    "salary_range": "Unknown",
                    "company_name": "Unknown",
                    "background_score": default_aspect,
                    "education_score": default_aspect,
                    "professional_score": default_aspect,
                    "tech_skills_score": default_aspect,
                    "teamwork_score": default_aspect,
                    "job_match_score": default_aspect,
                    "overall_score": 1,
                    "analysis_confidence": "low"
                }
            }

        if job_analysis_result.get("success"):
            analysis_data = job_analysis_result.get("analysis", {})
        else:
            logger.warning(f"OpenAI analysis failed: {job_analysis_result.get('error', 'Unknown error')}")
        
        # Extract specific sections for detailed logging
        basic_info = {
            'firstName': form_data_parsed.get('firstName', ''),
            'lastName': form_data_parsed.get('lastName', ''),
            'email': form_data_parsed.get('email', ''),
            'phoneNumber': form_data_parsed.get('phoneNumber', '')
        }
        
        education_info = {
            'collegeName': form_data_parsed.get('collegeName', ''),
            'degree': form_data_parsed.get('degree', ''),
            'major': form_data_parsed.get('major', ''),
            'graduationYear': form_data_parsed.get('graduationYear', '')
        }
        
        skills_info = {
            'programmingLanguages': form_data_parsed.get('programmingLanguages', ''),
            'frameworks': form_data_parsed.get('frameworks', ''),
            'databases': form_data_parsed.get('databases', ''),
            'tools': form_data_parsed.get('tools', '')
        }
        
        work_experience = form_data_parsed.get('workExperiences', [])
        
        # Log detailed breakdown
        logger.info("=== DETAILED BREAKDOWN ===")
        logger.info(f"Basic Info: {json.dumps(basic_info, indent=2)}")
        logger.info(f"Education: {json.dumps(education_info, indent=2)}")
        logger.info(f"Skills: {json.dumps(skills_info, indent=2)}")
        logger.info(f"Work Experience ({len(work_experience)} entries): {json.dumps(work_experience, indent=2)}")
        
        # Calculate some basic metrics
        filled_basic_fields = sum(1 for v in basic_info.values() if v.strip())
        filled_education_fields = sum(1 for v in education_info.values() if v.strip())
        filled_skills_fields = sum(1 for v in skills_info.values() if v.strip())
        total_work_entries = len(work_experience)
        
        logger.info("=== ANALYSIS METRICS ===")
        logger.info(f"Filled Basic Info Fields: {filled_basic_fields}/4")
        logger.info(f"Filled Education Fields: {filled_education_fields}/4")
        logger.info(f"Filled Skills Fields: {filled_skills_fields}/4")
        logger.info(f"Work Experience Entries: {total_work_entries}")
        
        # Log completion status
        completion_percentage = ((filled_basic_fields + filled_education_fields + filled_skills_fields) / 12) * 100
        logger.info(f"Form Completion: {completion_percentage:.1f}%")
        
        logger.info("=== AMBIT ALPHA ANALYSIS COMPLETE ===")
        
        # Store capability_analysis in career_analysis_data (update latest item for this user_id)
        if job_analysis_result.get("success") and job_analysis_result.get("analysis"):
            try:
                async def update_capability_analysis():
                    try:
                        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
                        table = dynamodb.Table('career_analysis_data')
                        # Get latest item
                        response = table.query(
                            KeyConditionExpression=Key('user_id').eq(user_id),
                            ScanIndexForward=False,
                            Limit=1
                        )
                        items = response.get('Items', [])
                        if not items:
                            logger.warning(f"No base item to update capability_analysis for user_id={user_id}")
                            return
                        ts = items[0]['timestamp']
                        # Update capability_analysis field only
                        table.update_item(
                            Key={'user_id': user_id, 'timestamp': ts},
                            UpdateExpression='SET capability_analysis = :ca',
                            ExpressionAttributeValues={
                                ':ca': job_analysis_result.get('analysis')
                            }
                        )
                        logger.info(f"Updated capability_analysis for user_id={user_id} ts={ts}")
                    except Exception as e:
                        logger.error(f"Error updating capability_analysis: {str(e)}")
                asyncio.create_task(update_capability_analysis())
                
            except Exception as e:
                logger.error(f"Error preparing improvement recommendations: {str(e)}")

        return {
            "status": "success",
            "message": "Ambit Alpha analysis data received and logged successfully",
            "analysis_id": f"ambit_alpha_{int(time.time())}",
            "completion_percentage": round(completion_percentage, 1),
            "job_analysis": job_analysis_result,
            "data_summary": {
                "target_job": target_job,
                "basic_info_complete": filled_basic_fields,
                "education_complete": filled_education_fields,
                "skills_complete": filled_skills_fields,
                "work_experience_count": total_work_entries
            }
        }
        
    except Exception as e:
        logger.error(f"Error in alpha_capability_analysis: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing Ambit Alpha analysis: {str(e)}")


# ==================== Resume PDF Generation Helper Functions ====================

def escape_latex(text: str) -> str:
    """
    Escape special LaTeX characters in user input.

    Args:
        text: Raw text string from user input

    Returns:
        LaTeX-safe string with special characters escaped

    Escapes: & % $ # _ { } ~ ^ \ |
    """
    if not text:
        return ""

    # Order matters! Escape backslash first
    replacements = [
        ('\\', r'\textbackslash{}'),
        ('&', r'\&'),
        ('%', r'\%'),
        ('$', r'\$'),
        ('#', r'\#'),
        ('_', r'\_'),
        ('{', r'\{'),
        ('}', r'\}'),
        ('~', r'\textasciitilde{}'),
        ('^', r'\textasciicircum{}'),
        ('|', r'\textbar{}'),
    ]

    result = text
    for char, replacement in replacements:
        result = result.replace(char, replacement)

    return result


def format_location_date(date_str: str) -> str:
    if not date_str:
        return ""
    
    # Just escape LaTeX special characters and return
    return escape_latex(date_str)

def build_section_professional_experience(experiences: List[ResumeProfessionalExperience]) -> str:
    """
    Build LaTeX for professional section.

    Args:
        experiences: List of professional experiences with nested job titles

    Returns:
        LaTeX string for complete section or empty string if no data
    """
    if not experiences:
        return ""

    latex = r"""\section*{Professional}
\vspace{1pt}
\hrule
\vspace{3pt}

"""

    for i, exp in enumerate(experiences):
        company = escape_latex(exp.company)

        for job_title in exp.jobTitles:
            title = escape_latex(job_title.title)
            # Format location and date: "Location | Date1 - Date2"
            date = format_location_date(job_title.date)

            # Add space gap above company name if it's not the first company
            if i > 0:
                latex += "\\vspace{4pt}\n"
            
            # Company name on its own row (always start on new line)
            latex += f"\n\\textbf{{{company}}} \\\\\n"
            # Job title with location/date on the right side
            latex += f"\\textbf{{\\textit{{{title}}}}} \\hfill \\textit{{{date}}}\n"

            if job_title.bullets:
                # Track if we're inside an itemize environment and current project name
                in_itemize = False
                current_project_name = None
                is_first_project = True
                
                # Get project technologies mapping if available
                project_tech_map = job_title.projectTechnologies or {}
                
                for bullet in job_title.bullets:
                    escaped_bullet = escape_latex(bullet)
                    # Check if bullet starts with project name prefix (handle both escaped and unescaped)
                    if bullet.startswith("__PROJECT_NAME__:") or escaped_bullet.startswith("__PROJECT\\_NAME__:"):
                        # Close any open itemize and add technologies for previous project
                        if in_itemize:
                            latex += r"\end{itemize}" + "\n"
                            # Add technologies for the previous project if available
                            if current_project_name and current_project_name in project_tech_map:
                                tech_list = project_tech_map[current_project_name]
                                if tech_list:
                                    if isinstance(tech_list, list):
                                        tech_str = ", ".join([escape_latex(str(t)) for t in tech_list])
                                    else:
                                        tech_str = escape_latex(str(tech_list))
                                    # Technologies with space gap below (with 2-character indentation)
                                    latex += f"\\hspace{{1em}}\\textit{{Technologies: {tech_str}}}\n"
                                    latex += "\\vspace{3pt}\n"
                            in_itemize = False
                        # Remove prefix and render as non-bold project name on its own row (always start on new line)
                        current_project_name = bullet.replace("__PROJECT_NAME__:", "").strip()
                        escaped_project_name = escape_latex(current_project_name)
                        # Add 2-character indentation for project name
                        latex += f"\n\\hspace{{1em}}{escaped_project_name}\n"
                        is_first_project = False
                    else:
                        # Start itemize if not already in one
                        if not in_itemize:
                            # Increase leftmargin by 2 characters (0.2in) for indentation
                            latex += r"\begin{itemize}[leftmargin=0.3in, itemsep=0pt, parsep=0pt, topsep=0pt]" + "\n"
                            in_itemize = True
                        latex += f"    \\item {escaped_bullet}\n"
                
                # Close itemize if still open and add technologies for last project
                if in_itemize:
                    latex += r"\end{itemize}" + "\n"
                    # Add technologies for the last project if available
                    if current_project_name and current_project_name in project_tech_map:
                        tech_list = project_tech_map[current_project_name]
                        if tech_list:
                            if isinstance(tech_list, list):
                                tech_str = ", ".join([escape_latex(str(t)) for t in tech_list])
                            else:
                                tech_str = escape_latex(str(tech_list))
                            # Technologies with space gap below (with 2-character indentation)
                            latex += f"\\hspace{{1em}}\\textit{{Technologies: {tech_str}}}\n"
                            latex += "\\vspace{3pt}\n"

    return latex


def build_section_education(education_list: List[ResumeEducation]) -> str:
    """
    Build LaTeX for education section.

    Args:
        education_list: List of education entries with nested degrees

    Returns:
        LaTeX string for complete section or empty string if no data
    """
    if not education_list:
        return ""

    latex = r"""\section*{Education}
\vspace{1pt}
\hrule
\vspace{3pt}

"""

    for i, edu in enumerate(education_list):
        university = escape_latex(edu.university)
        # Format location and date: "Location | Date1 - Date2"
        date = format_location_date(edu.date)

        # Add spacing between universities (except first)
        if i > 0:
            latex += "\\vspace{2pt}\n"

        # University name with location/duration on its own row (always start from beginning of row)
        latex += f"\n\\textbf{{{university}}} \\hfill \\textit{{{date}}} \\\\\n"

        for j, degree in enumerate(edu.degrees):
            degree_name = escape_latex(degree.degree)
            description = escape_latex(degree.description)

            # Degree name on its own row
            latex += f"\\textit{{{degree_name}}} \\\\\n"
            # Description on a new separate row from the beginning
            if description:
                latex += f"{description} \\\\\n"

    return latex


def build_section_projects(projects: List[ResumeProject]) -> str:
    """
    Build LaTeX for project experience section.

    Args:
        projects: List of projects (already filtered on frontend)

    Returns:
        LaTeX string for complete section or empty string if no data
    """
    if not projects:
        return ""

    latex = r"""\section*{Project Experience}
\vspace{1pt}
\hrule
\vspace{3pt}

"""

    for i, project in enumerate(projects):
        name = escape_latex(project.name)
        # Format location and date: "Location | Date1 - Date2"
        date = format_location_date(project.date)

        # Add small spacing above each project name (except first)
        if i > 0:
            latex += "\\vspace{4pt}\n"
        
        # Project name (bold, same style as university) with location/duration on its own row (always start on new line)
        latex += f"\n\\textbf{{{name}}} \\hfill \\textit{{{date}}}\n"

        # Bullet points
        if project.bullets:
            latex += r"\begin{itemize}[leftmargin=0.2in, itemsep=0pt, parsep=0pt, topsep=0pt]" + "\n"
            for bullet in project.bullets:
                latex += f"    \\item {escape_latex(bullet)}\n"
            latex += r"\end{itemize}" + "\n"

        # Technologies at the bottom for each project with space gap below
        if project.technologies:
            tech_str = ", ".join([escape_latex(t) for t in project.technologies])
            latex += f"\\textit{{Technologies: {tech_str}}}\n"
            latex += "\\vspace{3pt}\n"

    return latex


def build_section_skills(skills: List[ResumeSkill]) -> str:
    """
    Build LaTeX for technical skills section.

    Args:
        skills: List of skill categories with keywords

    Returns:
        LaTeX string for complete section or empty string if no data
    """
    if not skills:
        return ""

    latex = r"""\section*{Technical Skills}
\vspace{1pt}
\hrule
\vspace{3pt}

"""

    for skill in skills:
        topic = escape_latex(skill.topic)
        keywords = escape_latex(skill.keywords)

        latex += f"\\textbf{{{topic}}}: {keywords} \\\\\n"

    return latex


def build_section_achievements(achievements: List[ResumeAchievement]) -> str:
    """
    Build LaTeX for achievements section.

    Args:
        achievements: List of achievements/awards

    Returns:
        LaTeX string for complete section or empty string if no data
    """
    if not achievements:
        return ""

    latex = r"""\section*{Achievements}
\vspace{1pt}
\hrule
\vspace{3pt}

"""

    latex += r"\begin{itemize}[leftmargin=0.2in, itemsep=0pt, parsep=0pt, topsep=0pt]" + "\n"
    for achievement in achievements:
        value = escape_latex(achievement.value)
        latex += f"    \\item {value}\n"
    latex += r"\end{itemize}" + "\n"

    return latex


def generate_latex_content(resume_data: GenerateResumePDFRequest) -> str:
    """
    Generate complete LaTeX document from resume data.

    Args:
        resume_data: Validated Pydantic model with resume information

    Returns:
        Complete LaTeX document as string

    Features:
        - Conditionally includes sections only if data exists
        - Escapes all user input
        - Applies ATS-optimized formatting
        - Dynamic spacing based on content volume
    """
    # Calculate content density for spacing adjustments
    total_items = (
        len(resume_data.professional_experiences) +
        len(resume_data.education) +
        len(resume_data.projects) +
        len(resume_data.skills) +
        len(resume_data.achievements)
    )

    # Adjust spacing based on content density - optimized for single page
    if total_items > 25:  # Dense resume
        margin = "0.4in"
        line_spacing = "0.9"
    elif total_items > 15:  # Normal resume
        margin = "0.45in"
        line_spacing = "0.92"
    else:  # Sparse resume
        margin = "0.5in"
        line_spacing = "0.95"

    # Build contact information
    # Process each contact field with exact field name (label) and field data (value)
    # Use \mbox{} to keep each field name and field data on the same row
    contact_parts = []
    for contact in resume_data.contact:
        label = escape_latex(contact.label)
        value = escape_latex(contact.value)

        # Check if value is a URL (for LinkedIn, GitHub, etc.)
        if contact.value.startswith(('http://', 'https://', 'www.')):
            # Use href for clickable links (show label: URL)
            # Wrap in \mbox{} to keep field name and URL on same row
            contact_parts.append(f"\\mbox{{{label}: \\href{{{contact.value}}}{{{value}}}}}")
        else:
            # For non-URL fields, also include the label (field name) with the value (field data)
            # Format: "Field Name: Field Data"
            # Wrap in \mbox{} to keep field name and field data on same row
            contact_parts.append(f"\\mbox{{{label}: {value}}}")

    # Join contact fields with LaTeX \quad spacing (about 3 spaces)
    # LaTeX collapses regular spaces, so we use \quad for explicit spacing
    contact_line = " \\quad ".join(contact_parts)

    # Start building the LaTeX document
    latex = f"""\\documentclass[10pt,a4paper]{{article}}

% ATS-Optimized packages
\\usepackage[margin={margin}, top=0.3in, bottom=0.3in]{{geometry}}
\\usepackage{{enumitem}}
\\usepackage{{setspace}}
\\usepackage[hidelinks]{{hyperref}}
\\usepackage{{titlesec}}

% Disable page numbers
\\pagestyle{{empty}}

% Reduce spacing for compact layout
\\setlength{{\\parindent}}{{0pt}}
\\setlength{{\\parskip}}{{0pt}}
\\setlength{{\\topsep}}{{0pt}}
\\setlength{{\\partopsep}}{{0pt}}
\\setlist{{nosep, leftmargin=0.2in, topsep=0pt, partopsep=0pt, itemsep=0pt}}
\\setstretch{{{line_spacing}}}

% Reduce section spacing (before and after sections)
\\titlespacing*{{\\section}}{{0pt}}{{4pt}}{{1pt}}

% Hyperref setup - black links for ATS
\\hypersetup{{
    colorlinks=true,
    linkcolor=black,
    urlcolor=black,
    citecolor=black
}}

\\begin{{document}}

\\vspace{{-8pt}}

% NAME (centered, large)
\\begin{{center}}
    {{\\Large\\textbf{{{escape_latex(resume_data.name)}}}}}
\\end{{center}}

\\vspace{{4pt}}

% CONTACT INFO (centered, max 80% width)
\\begin{{center}}
    \\begin{{minipage}}{{0.85\\textwidth}}
        \\centering
        \\small {contact_line}
    \\end{{minipage}}
\\end{{center}}

\\vspace{{4pt}}

"""

    # Add sections (only if they have content)
    latex += build_section_professional_experience(resume_data.professional_experiences)
    latex += build_section_education(resume_data.education)
    latex += build_section_projects(resume_data.projects)
    latex += build_section_skills(resume_data.skills)
    latex += build_section_achievements(resume_data.achievements)

    # Close document
    latex += "\\end{document}\n"

    return latex


def compile_latex_to_pdf(latex_content: str, output_dir: str) -> str:
    """
    Compile LaTeX content to PDF using pdflatex.

    Args:
        latex_content: Complete LaTeX document string
        output_dir: Directory for temporary files (e.g., /tmp/resume_xyz)

    Returns:
        Path to generated PDF file

    Raises:
        RuntimeError: If pdflatex compilation fails
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    # Write .tex file
    tex_file = os.path.join(output_dir, "resume.tex")
    with open(tex_file, 'w', encoding='utf-8') as f:
        f.write(latex_content)

    # Run pdflatex (twice for proper formatting)
    pdflatex_cmd = [
        '/usr/local/bin/pdflatex',
        '-interaction=nonstopmode',
        '-output-directory', output_dir,
        tex_file
    ]

    # First compilation
    result = subprocess.run(
        pdflatex_cmd,
        capture_output=True,
        timeout=30,
        text=True
    )

    # Second compilation for cross-references
    result = subprocess.run(
        pdflatex_cmd,
        capture_output=True,
        timeout=30,
        text=True
    )

    # Check if PDF was generated
    pdf_file = os.path.join(output_dir, "resume.pdf")

    if result.returncode != 0:
        logger.error(f"pdflatex compilation failed with return code {result.returncode}")
        logger.error(f"pdflatex stdout: {result.stdout}")
        logger.error(f"pdflatex stderr: {result.stderr}")
        raise RuntimeError(f"PDF compilation failed: {result.stderr}")

    if not os.path.exists(pdf_file):
        raise RuntimeError("PDF file was not generated")

    logger.info(f"PDF generated successfully: {pdf_file}")
    return pdf_file


def cleanup_temp_files(directory: str) -> None:
    """
    Clean up temporary files created during PDF generation.

    Args:
        directory: Path to temporary directory to remove

    Purpose:
        Critical for Lambda to avoid /tmp/ space exhaustion
        Removes .tex, .aux, .log, .out files and directory
    """
    try:
        if os.path.exists(directory):
            shutil.rmtree(directory)
            logger.info(f"Cleaned up temp directory: {directory}")
    except Exception as e:
        # Log but don't raise - cleanup failure shouldn't break response
        logger.warning(f"Failed to cleanup {directory}: {str(e)}")


@app.post("/generate_resume_pdf")
async def generate_resume_pdf(
    request: GenerateResumePDFRequest,
    background_tasks: BackgroundTasks
):
    """
    Generate ATS-optimized resume PDF from structured data.

    Accepts JSON with resume information and returns PDF file.
    Frontend sends only selected projects based on knowledge scope.

    Returns:
        PDF file as FileResponse with automatic cleanup
    """
    request_id = str(uuid.uuid4())
    temp_dir = f"/tmp/resume_{request_id}"

    try:
        # Check plan and download usage quota
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        table = dynamodb.Table('ambit-dashboard-application-data')
        timestamp = datetime.utcnow().isoformat()

        sub_response = table.get_item(Key={'PK': request.user_id, 'SK': 'SUBSCRIPTION'})
        plan = sub_response.get('Item', {}).get('plan', 'free')

        usage_response = table.get_item(Key={'PK': request.user_id, 'SK': 'USAGE'})
        usage_item = usage_response.get('Item')

        if not usage_item:
            usage_item = {
                'PK': request.user_id,
                'SK': 'USAGE',
                'craft_count': 0,
                'analysis_count': 0,
                'download_count': 0,
                'createdAt': timestamp,
                'updatedAt': timestamp
            }
            table.put_item(Item=usage_item)

        download_count = int(usage_item.get('download_count', 0))

        if plan == 'free' and download_count >= 3:
            return JSONResponse(content={
                "status": "error",
                "error_code": "DOWNLOAD_LIMIT_EXCEEDED",
                "message": "Free plan download limit reached."
            })

        table.update_item(
            Key={'PK': request.user_id, 'SK': 'USAGE'},
            UpdateExpression='SET download_count = if_not_exists(download_count, :zero) + :inc, updatedAt = :ts',
            ExpressionAttributeValues={':inc': 1, ':zero': 0, ':ts': timestamp}
        )

        logger.info(f"PDF generation request for: {request.name}")
        logger.info(f"Sections: exp={len(request.professional_experiences)}, "
                   f"edu={len(request.education)}, proj={len(request.projects)}, "
                   f"skills={len(request.skills)}, achievements={len(request.achievements)}")

        # 1. Generate LaTeX content
        latex_content = generate_latex_content(request)
        logger.info("LaTeX content generated successfully")

        # 2. Compile to PDF
        pdf_path = compile_latex_to_pdf(latex_content, temp_dir)

        # 3. Return PDF with background cleanup
        safe_filename = request.name.replace(' ', '_').replace('/', '_')
        logger.info(f"Returning PDF: {pdf_path}")

        return FileResponse(
            path=pdf_path,
            media_type="application/pdf",
            filename=f"{safe_filename}_Resume.pdf",
            background=background_tasks.add_task(cleanup_temp_files, temp_dir)
        )

    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        cleanup_temp_files(temp_dir)
        raise HTTPException(
            status_code=500,
            detail=f"Error generating resume PDF: {str(e)}"
        )


@app.post("/generate_resume_tex")
async def generate_resume_tex(request: GenerateResumePDFRequest):
    """Return the LaTeX .tex source file used to generate the resume PDF."""
    try:
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        table = dynamodb.Table('ambit-dashboard-application-data')
        timestamp = datetime.utcnow().isoformat()

        sub_response = table.get_item(Key={'PK': request.user_id, 'SK': 'SUBSCRIPTION'})
        plan = sub_response.get('Item', {}).get('plan', 'free')

        usage_response = table.get_item(Key={'PK': request.user_id, 'SK': 'USAGE'})
        usage_item = usage_response.get('Item')

        if not usage_item:
            usage_item = {'PK': request.user_id, 'SK': 'USAGE', 'craft_count': 0,
                          'analysis_count': 0, 'download_count': 0,
                          'createdAt': timestamp, 'updatedAt': timestamp}
            table.put_item(Item=usage_item)

        download_count = int(usage_item.get('download_count', 0))
        if plan == 'free' and download_count >= 3:
            return JSONResponse(content={"status": "error",
                                         "error_code": "DOWNLOAD_LIMIT_EXCEEDED",
                                         "message": "Free plan download limit reached."})

        table.update_item(
            Key={'PK': request.user_id, 'SK': 'USAGE'},
            UpdateExpression='SET download_count = if_not_exists(download_count, :zero) + :inc, updatedAt = :ts',
            ExpressionAttributeValues={':inc': 1, ':zero': 0, ':ts': timestamp}
        )

        latex_content = generate_latex_content(request)
        safe_filename = request.name.replace(' ', '_').replace('/', '_')

        return Response(
            content=latex_content.encode('utf-8'),
            media_type="application/x-tex",
            headers={"Content-Disposition": f'attachment; filename="{safe_filename}_Resume.tex"'}
        )
    except Exception as e:
        logger.error(f"Error generating resume .tex: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating resume .tex: {str(e)}")


@app.post("/get_bullet_points_from_project_source")
async def get_bullet_points_from_project_source(request: Request):
    body = await request.json()
    source_input = body.get('source_input', '').strip()

    if not source_input:
        return {"success": False, "error_code": "EMPTY_INPUT",
                "message": "Please enter a project URL or description."}
    if len(source_input) > 20000:
        return {"success": False, "error_code": "INPUT_TOO_LONG",
                "message": "Input is too long. Please use a URL or a description under 20,000 characters."}

    is_url = source_input.startswith('http://') or source_input.startswith('https://')
    if is_url:
        content = await fetch_web_page_content(source_input)
        error_indicators = ["Invalid URL", "Non-HTML", "Timeout", "Connection error",
                            "HTTP error", "Request error", "Unexpected error"]
        if any(ind in content for ind in error_indicators) or len(content.strip()) < 50:
            return {"success": False, "error_code": "URL_FETCH_FAILED",
                    "message": f"Could not fetch URL content. Ensure it is publicly accessible. Detail: {content[:200]}"}
    else:
        content = source_input

    system_prompt = (
        "You are a professional technical resume writer for software engineering roles. "
        "Generate concise, impactful resume bullet points (~13 words each). "
        "Use strong past-tense verbs. "
        "overview: what the project is and its primary purpose. "
        "showcase1: key tech stack, tools, frameworks, or architectural choices used. "
        "showcase2: (optional) a second distinct feature or collaboration highlight. Only include if genuinely different. "
        "achievement: quantitative outcome (latency reduction, user count, accuracy, etc.)."
    )
    user_prompt = f"Generate resume bullet points for this project:\n\n{content[:4000]}"

    try:
        response = client.responses.parse(
            model="gpt-4o-mini",
            input=[{"role": "system", "content": system_prompt},
                   {"role": "user", "content": user_prompt}],
            text_format=ProjectBulletPoints
        )
        bp = response.output_parsed
        tech = bp.showcase1.strip()
        if bp.showcase2 and bp.showcase2.strip():
            tech = tech + '\n' + bp.showcase2.strip()
        return {"success": True, "data": {
            "overview": bp.overview.strip(),
            "techAndTeamwork": tech,
            "achievement": bp.achievement.strip()
        }}
    except Exception as e:
        logger.error(f"OpenAI error in get_bullet_points_from_project_source: {str(e)}")
        return {"success": False, "error_code": "AI_FAILED",
                "message": "Failed to generate bullet points. Please try again."}


@app.post("/auto_fill_info_from_resume")
async def auto_fill_info_from_resume(
    resume_file: UploadFile = File(...),
    cognito_sub: str = Form(...)
):
    file_content = await resume_file.read()
    resume_text = await extract_text_from_resume(file_content, resume_file.filename or 'resume.pdf')

    PREDEFINED_TECHNOLOGIES = """
Software Engineering:
  Languages/Build: TypeScript, JavaScript, Python, Java, C#, C++, Go, Rust, Kotlin, SQL, Bash, Node.js, JVM, .NET, PHP
  Client/UI: HTML5, CSS3, DOM, HTTP/HTTPS, PWA, iOS, Android, Swift, WebSockets
  Backend/Distributed: REST APIs, gRPC, GraphQL, Microservices, Event-driven architecture, Containers, Kubernetes, JWT, Caching patterns, Load balancing
  Data/Messaging: Relational data modeling, Key-value data modeling, ACID, Replication, Partitioning, Event streaming, Message queuing
  OS/Networks: TCP/UDP, Lock-free concurrency, Zero-copy I/O

Frameworks (SE):
  Languages/Build: Git, GitHub, Docker, GitHub Actions, Jenkins, Maven, Gradle, npm, pip, Poetry, Conda
  Client/UI: React, Next.js, Vue.js, Angular, Svelte, Tailwind CSS, Vite, Webpack, React Native, SwiftUI
  Backend: Node.js, Express, Spring Boot, FastAPI, Django, Flask, NestJS, Kubernetes, Terraform, AWS, GCP, Azure
  Data: PostgreSQL, MySQL, Redis, MongoDB, Elasticsearch, Apache Kafka, RabbitMQ, Apache Spark, Snowflake, BigQuery, dbt

AI/ML Technologies:
  AI App: LLM APIs, RAG, prompt engineering, vector databases, AI agents, function calling, fine-tuning, RLHF
  Foundation/Research: Transformers, attention mechanisms, diffusion models, reinforcement learning, computer vision, NLP
  MLOps: experiment tracking, model serving, model monitoring, feature stores, CI/CD for ML, A/B testing
  Data Engineering: data pipelines, ETL, stream processing, batch processing, data warehousing, feature engineering
  Infrastructure: GPU clusters, distributed training, CUDA, model optimization, quantization, inference optimization

Frameworks (AI/ML):
  AI App: LangChain, LlamaIndex, OpenAI API, Claude API, Hugging Face, LangGraph, AutoGen, ChromaDB, Pinecone, Weaviate
  Foundation: PyTorch, TensorFlow, JAX, Transformers, PEFT, DeepSpeed, vLLM, llama.cpp, Axolotl, Unsloth
  MLOps: MLflow, Kubeflow, Apache Airflow, Ray Serve, AWS SageMaker, Vertex AI, Weights & Biases, DVC
  Data: Spark, Kafka, dbt, Snowflake, Databricks, BigQuery, Pandas, Polars, DuckDB
  Infrastructure: CUDA, cuDNN, NCCL, TensorRT, Slurm, Triton
"""

    system_prompt = f"""You are an expert resume parser and career advisor. Extract ALL information from the resume into a structured format. Follow these rules strictly:

CAREER FOCUS: Determine the best fit from exactly these four values:
- "software-engineering" — backend, frontend, fullstack, systems, infrastructure
- "ai-ml" — machine learning, deep learning, LLMs, NLP, computer vision, AI research
- "data-science" — data analysis, statistics, BI, analytics, research science
- "data-engineering" — data pipelines, ETL, data platforms, data infrastructure

PROJECTS: Classify as:
- personal_projects: personal side projects, college/university projects, open-source contributions, hackathon projects
- professional_projects: work projects done while employed (link each to the company/job via selectedWorkExperience as "CompanyName – JobTitle")

TECHNOLOGIES & FRAMEWORKS: For selectedTechnologies and selectedFrameworks in projects, ONLY include items that exactly match from these predefined lists:
{PREDEFINED_TECHNOLOGIES}
Put items NOT in the predefined lists into customTechnologies / customFrameworks.

TECHNICAL SKILLS: For technical_skill_items, pick individual items from the predefined lists above that match skills in the resume. Put non-matching items in custom_technical_items.

DATES: Use month as 3-letter abbreviation (Jan, Feb, Mar, Apr, May, Jun, Jul, Aug, Sep, Oct, Nov, Dec) and year as 4-digit string. Use empty strings for unknown values.

ACHIEVEMENTS: Include awards, publications, certifications, patents, honors. Use descriptive type strings.

PROJECT DESCRIPTIONS:
- overview: what the project does and its purpose
- techAndTeamwork: technologies used and team collaboration details
- achievement: measurable outcomes, impact, or results

Fill all fields as completely as possible. Use empty strings for missing data, never null."""

    user_prompt = f"Extract all information from this resume:\n\n{resume_text}"

    response = client.responses.parse(
        model="gpt-4o",
        input=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        text_format=AutoFillResumeResult
    )

    result = response.output_parsed
    return {"success": True, "data": result.model_dump()}


RESUME_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "ask_resume_intent",
            "description": (
                "Use this when the user has a general or ambiguous resume request that could mean "
                "either building a brand-new resume from their profile OR improving an existing one. "
                "Trigger examples: 'Make my resume', 'Build my resume', 'Help me with my resume', "
                "'How do I write a resume?', 'I need a resume', 'Create my resume', 'Write my resume', "
                "'Can you make me a resume?', 'I want to make a resume', 'How to build a resume'. "
                "The UI will display two clickable option cards side by side so the user can choose their path. "
                "Always call this tool for general or ambiguous resume requests — never present options as plain text."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": (
                            "A short intro line before the option cards appear "
                            "(e.g. 'Here are two ways I can help you build your resume.'). "
                            "Keep it under 20 words."
                        )
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "navigate_to_existing_resume",
            "description": (
                "Use this ONLY when the user explicitly and specifically asks to improve, enhance, "
                "polish, tailor, fix, update, strengthen, or rewrite an existing resume they already have. "
                "Use navigate_to_existing_resume when intent is clearly about an existing resume they want to improve. "
                "Ambitology's 'Craft from Existing Resume' service handles this."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": (
                            "A warm, concise reply (2-3 sentences) telling the user that "
                            "Ambitology's 'Craft from Existing Resume' service can help. "
                            "Briefly explain it analyzes their resume and tailors it for any job. "
                            "Ask if they'd like to go there now."
                        )
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "navigate_to_knowledge_base_resume",
            "description": (
                "Use this when the user asks how to craft, create, generate, or build a resume "
                "from their knowledge base, experience, background, projects, technical skills, "
                "work history, or professional profile stored in Ambitology. "
                "Ambitology's 'Craft from Knowledge Base' service generates a tailored resume "
                "from the user's saved skills, projects, and experience for a target job or industry."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": (
                            "A warm, concise reply (2-3 sentences) telling the user that "
                            "Ambitology's 'Craft from Knowledge Base' service can generate a "
                            "tailored resume directly from their saved profile, skills, and projects. "
                            "Ask if they'd like to go there now."
                        )
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "show_pricing",
            "description": (
                "Use this when the user asks about plans, pricing, cost, subscription, upgrade, or feature access. "
                "Trigger examples: 'What plans do you have?', 'What's the price for pro?', 'What does the free plan include?', "
                "'What does the pro plan include?', 'How can I upgrade?', 'My usage is blocked', "
                "'How many free uses do I get?', 'Do I have unlimited usage?'. "
                "Always call this tool — never describe pricing in plain text."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": "A short intro line before the pricing card appears (e.g. 'Here's an overview of our plans.'). Keep it under 15 words."
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "ask_project_type",
            "description": (
                "Use this in STEP 1 of the project analysis flow to ask the user what kind of project "
                "they want to analyze. Always call this tool — never ask as plain text. "
                "The UI will present four clickable cards: Personal Project, Side Project, Work Project, Research Project."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": "A short intro line before the choices appear (e.g. 'First, what kind of project is this?'). Keep it under 15 words."
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "ask_project_status",
            "description": (
                "Use this in STEP 2 of the project analysis flow to ask the user whether the project "
                "is currently in progress, already completed, or a planned future project. "
                "Always call this tool — never ask as plain text. "
                "The UI will present three clickable cards: In Progress, Completed, Planning Ahead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": "A short line before the choices appear (e.g. 'And what's the status of this project?'). Keep it under 15 words."
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "navigate_to_established_personal_project",
            "description": (
                "Use this when the user has confirmed they want to analyze a personal project "
                "that is ongoing or already completed — to generate bullet points, find technologies, "
                "tools, or technical keywords for that project. Navigate them to "
                "Knowledge Base > Established Expertise > Personal Project and create a new project page."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": (
                            "A warm, concise reply (1-2 sentences) confirming you're taking them to "
                            "the Established Expertise > Personal Project page to add a new project. "
                            "Then ask them to provide the project URL or paste the project details "
                            "so you can generate the best description."
                        )
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "navigate_to_expanding_personal_project",
            "description": (
                "Use this when the user has confirmed they want to analyze a personal project "
                "that is planned for the future (not yet started) — to generate bullet points, "
                "find technologies, tools, or technical keywords for that project. Navigate them to "
                "Knowledge Base > Expanding Knowledge Base > Planned Personal Project and create a new project page."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": (
                            "A warm, concise reply (1-2 sentences) confirming you're taking them to "
                            "the Expanding Knowledge Base > Planned Personal Project page to add a new project. "
                            "Then ask them to provide the project URL or paste the project details "
                            "so you can generate the best description."
                        )
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "navigate_to_established_professional_project",
            "description": (
                "Use this when the user has confirmed they want to analyze a professional or work-related "
                "project that is ongoing or already completed — to generate bullet points, find technologies, "
                "tools, or technical keywords for that project. Navigate them to "
                "Knowledge Base > Established Expertise > Professional Project and create a new project page."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": (
                            "A warm, concise reply (1-2 sentences) confirming you're taking them to "
                            "the Established Expertise > Professional Project page to add a new project. "
                            "Then ask them to provide the project URL or paste the project details "
                            "so you can generate the best description."
                        )
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "navigate_to_expanding_professional_project",
            "description": (
                "Use this when the user has confirmed they want to analyze a professional or work-related "
                "project that is planned for the future (not yet started) — to generate bullet points, "
                "find technologies, tools, or technical keywords for that project. Navigate them to "
                "Knowledge Base > Expanding Knowledge Base > Planned Professional Project and create a new project page."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": (
                            "A warm, concise reply (1-2 sentences) confirming you're taking them to "
                            "the Expanding Knowledge Base > Planned Professional Project page to add a new project. "
                            "Then ask them to provide the project URL or paste the project details "
                            "so you can generate the best description."
                        )
                    }
                },
                "required": ["reply"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "analyze_project",
            "description": (
                "Use this when the user has provided project details (a URL or text description) "
                "and you need to analyze the project to generate professional resume bullet points, project details, "
                "and technical keyword or skills or tools suggestions. Only call this after the user has provided "
                "actual project content (URL or pasted details)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "reply": {
                        "type": "string",
                        "description": "A brief message (1 sentence) saying you're analyzing the project now."
                    },
                    "project_details": {
                        "type": "string",
                        "description": "The project URL or description text provided by the user."
                    },
                    "project_type": {
                        "type": "string",
                        "enum": [
                            "personal_established",
                            "personal_expanding",
                            "professional_established",
                            "professional_expanding"
                        ],
                        "description": (
                            "The project type inferred from the earlier conversation. "
                            "personal_established = personal + ongoing/done, "
                            "personal_expanding = personal + future planned, "
                            "professional_established = professional + ongoing/done, "
                            "professional_expanding = professional + future planned."
                        )
                    }
                },
                "required": ["reply", "project_details", "project_type"]
            }
        }
    }
]


@app.post("/ai-chat")
async def ai_chat(request: Request):
    try:
        data = await request.json()
        message = data.get("message", "")
        email   = data.get("email", "")
        name    = data.get("name", "")
        history = data.get("history", [])   # list of {role, content} dicts
        career_focus = data.get("career_focus", "software engineering")

        logger.info(f"AI chat question from {name} ({email}): {message}")

        messages = [
            {
                "role": "system",
                "content": (
                    "You are a career coach AI for Ambitology. Be warm, direct, and easy to read.\n\n"
                    "FORMATTING RULES (always follow):\n"
                    "- Keep every reply under 50 words total.\n"
                    "- Never write a wall of text. Use short sentences.\n"
                    "- When listing steps or items, put each on its own line starting with ① ② ③ or • — never run them together.\n"
                    "- Separate distinct thoughts with a blank line.\n"
                    "- No unnecessary filler phrases like 'Great question!' or 'Of course!'.\n\n"
                    "RESUME CRAFTING:\n"
                    "- User has a general or ambiguous resume request ('make my resume', 'build my resume', 'write a resume', 'help with my resume', 'I need a resume', 'create my resume') → use ask_resume_intent. Show option cards so user can choose their path.\n"
                    "- User explicitly wants to improve/polish/fix/enhance an existing resume they already have → use navigate_to_existing_resume.\n"
                    "- User explicitly wants to craft/create/generate a resume from their knowledge base or saved profile → use navigate_to_knowledge_base_resume.\n"
                    "- When in doubt between the two resume paths, always use ask_resume_intent — never guess.\n\n"
                    "PLAN PRICING:\n"
                    "- User asks about plans, pricing, cost, subscription, upgrade, feature access, usage limits, or how to unblock a feature → use show_pricing.\n"
                    "- Never answer pricing questions in plain text — always use the show_pricing tool.\n\n"
                    "PROJECT ANALYSIS FLOW (follow steps in order, never skip):\n"
                    "Trigger this flow when the user asks about ANY of the following:\n"
                    "• How to write project bullet points or project descriptions\n"
                    "• How to write or find technical skills for a project\n"
                    "• What tools to use or list for a project\n"
                    "• Popular skills, technologies, or technical keywords for a project\n"
                    "• What technologies or tools to write in the resume for a project\n"
                    "• Any question about resume content specific to a project (technologies, frameworks, skills, tools)\n"
                    "STEP 1 — Use ask_project_type tool (never ask as plain text).\n"
                    "STEP 2 — Use ask_project_status tool (never ask as plain text).\n"
                    "STEP 3 — Use matching navigation tool based on the user's card selections:\n"
                    "  Personal Project or Side Project + In Progress/Completed → navigate_to_established_personal_project\n"
                    "  Personal Project or Side Project + Planning Ahead → navigate_to_expanding_personal_project\n"
                    "  Work Project, Work Experience, or Research Project + In Progress/Completed → navigate_to_established_professional_project\n"
                    "  Work Project, Work Experience, or Research Project + Planning Ahead → navigate_to_expanding_professional_project\n"
                    "STEP 4 — Ask: 'Share a project URL or paste the project details.'\n"
                    "STEP 5 — User provides URL or details → use analyze_project tool.\n\n"
                    "Do NOT call a navigation tool before STEP 1 and STEP 2 are both answered."
                )
            },
            *[{"role": m["role"], "content": m["content"]} for m in history],
            {"role": "user", "content": message},
        ]

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            tools=RESUME_TOOLS,
            tool_choice="auto",
        )

        choice = response.choices[0]

        if choice.finish_reason == "tool_calls":
            tool_name = choice.message.tool_calls[0].function.name
            args = json.loads(choice.message.tool_calls[0].function.arguments)
            reply = args.get("reply", "I can help you with that!")
            logger.info(f"AI chat tool call {tool_name} for {email}: {reply}")

            # Handle project analysis inline
            if tool_name == "analyze_project":
                project_details = args.get("project_details", "")
                project_type = args.get("project_type", "personal_established")

                # If it looks like a URL, try to fetch it
                content = project_details
                if project_details.strip().startswith("http"):
                    try:
                        page_res = requests.get(project_details.strip(), timeout=10, headers={"User-Agent": "Mozilla/5.0"})
                        soup = BeautifulSoup(page_res.text, 'html.parser')
                        content = soup.get_text(separator=' ', strip=True)[:3000]
                    except Exception as fetch_err:
                        logger.warning(f"Failed to fetch project URL: {fetch_err}")
                        content = project_details

                career_focus_label = career_focus.replace("-", " ").title() if career_focus else "software engineering"

                if career_focus == 'ai-ml':
                    kw_sections = [
                        'AI Application & Product',
                        'Foundation Model & AI Research',
                        'MLOps & ML Platform',
                        'Data Engineering & Analytics',
                        'Infrastructure & Compute',
                    ]
                else:
                    kw_sections = [
                        'Languages, Runtimes & Build Tooling',
                        'Client, UI & Product Experience',
                        'Backend Services & Distributed Systems',
                        'Data & Messaging Layer',
                        'Operating Systems & Networks',
                    ]
                sections_str = ", ".join(f'"{s}"' for s in kw_sections)

                industry_options = (
                    "AI & Machine Learning, Blockchain & Web3, Cloud Computing, "
                    "SaaS / Enterprise Software, Big Tech / Consumer Internet, FinTech, "
                    "Trading & Quant Finance, E-commerce & Marketplace, Cybersecurity, "
                    "Data & Analytics, Developer Tools, Healthcare & Insurance, Gaming, "
                    "Autonomous Vehicles & Robotics, Generative AI Platforms"
                )
                analysis_prompt = (
                    f"You are a technical resume expert helping a {career_focus_label} professional.\n\n"
                    f"Project information:\n{content[:2500]}\n\n"
                    "Analyze this project and generate professional resume content. "
                    "Think about how the project should be designed and implemented in the best case scenario.\n\n"
                    "Return a JSON object with exactly these keys:\n"
                    "- projectName: A concise, professional project name (3-6 words) that clearly "
                    "describes what this project does. No special characters.\n"
                    "- industrySector: The single best-fit industry sector from this exact list: "
                    f"{industry_options}. Choose the closest match.\n"
                    "- overview: One clear sentence describing what the project does and its main purpose.\n"
                    "- techAndTeamwork: One to two sentences about implementation approach, "
                    "technologies used, architecture decisions, and technical challenges solved. "
                    "If multiple sentences, separate with a newline character (\\n).\n"
                    "- achievement: One sentence with quantified outcomes — use estimated percentages "
                    "or realistic numbers if exact metrics are not available (e.g. 'reduced latency by ~40%', "
                    "'supports 10k+ concurrent users').\n"
                    f"- technologies: An object with exactly these section keys: {sections_str}. "
                    "Each section must contain 5–8 technology keywords (languages, platforms, cloud services, "
                    "databases) most relevant to this project AND popular for this career focus. "
                    "Total roughly 30 keywords across all sections.\n"
                    f"- frameworks: An object with the same section keys: {sections_str}. "
                    "Each section must contain 5–8 framework/tool keywords (frameworks, libraries, dev tools, "
                    "testing tools, CI/CD tools) most relevant to this project AND popular for this career focus. "
                    "Total roughly 30 keywords across all sections. Must not duplicate technologies entries.\n\n"
                    "Return only valid JSON with no extra text."
                )

                try:
                    analysis_response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": "You are a technical resume expert. Return only valid JSON."},
                            {"role": "user", "content": analysis_prompt}
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.4,
                    )
                    analysis_data = json.loads(analysis_response.choices[0].message.content)
                    logger.info(f"Project analysis complete for {email}, project_type={project_type}")
                    return {
                        "status": "success",
                        "reply": "I've analyzed your project and prepared updates for each field. Which fields would you like to fill in?",
                        "action": {
                            "type": "update_project_description",
                            "data": {
                                "projectType": project_type,
                                "projectSource": project_details[:2000],
                                "projectName": analysis_data.get("projectName", ""),
                                "industrySector": analysis_data.get("industrySector", ""),
                                "overview": analysis_data.get("overview", ""),
                                "techAndTeamwork": analysis_data.get("techAndTeamwork", ""),
                                "achievement": analysis_data.get("achievement", ""),
                                "technologies": analysis_data.get("technologies", {}),
                                "frameworks": analysis_data.get("frameworks", {})
                            }
                        }
                    }
                except Exception as analysis_err:
                    logger.error(f"Project analysis failed: {analysis_err}")
                    return {
                        "status": "success",
                        "reply": "I had trouble analyzing the project details. Could you provide more information or try a different URL?",
                    }

            return {"status": "success", "reply": reply, "action": {"type": tool_name}}
        else:
            reply = choice.message.content
            logger.info(f"AI chat response to {email}: {reply}")
            import re as _re
            rl = reply.lower()

            # ── Semantic: project TYPE question ──
            type_patterns = [
                r'what (kind|type) of project',
                r'is (this|it) (a )?(personal|professional)',
                r'personal.{0,20}(or|vs|versus).{0,20}(professional|work experience)',
                r'personal project.{0,60}work experience',
                r'personal.{0,40}work.{0,40}research',
            ]
            if any(_re.search(p, rl) for p in type_patterns):
                return {"status": "success", "reply": reply, "action": {"type": "ask_project_type"}}

            # ── Semantic: project STATUS question ──
            status_patterns = [
                r'status of (this |the )?project',
                r'(currently in progress|already completed|planned for the future)',
                r'(in progress|ongoing).{0,60}(completed|finished).{0,60}(plan|future)',
                r'(completed|finished).{0,60}(in progress|ongoing).{0,60}(plan|future)',
                r'is (this |the )?project (currently|already|still|done)',
                r"what.{0,20}(status|stage|state).{0,20}(project|it)",
            ]
            if any(_re.search(p, rl) for p in status_patterns):
                return {"status": "success", "reply": reply, "action": {"type": "ask_project_status"}}

            # ── Auto-detect numbered/bulleted choice lists ──
            lines = [l.strip() for l in reply.split('\n') if l.strip()]
            numbered = []
            for line in lines:
                m = _re.match(r'^(?:\d+[.)]\s*|[①②③④⑤⑥⑦⑧⑨⑩]\s*)(.+)$', line)
                if m:
                    numbered.append(m.group(1).strip())
            if 2 <= len(numbered) <= 5 and all(len(item) <= 45 for item in numbered):
                return {"status": "success", "reply": reply, "action": {"type": "text_choices", "choices": numbered}}

            return {"status": "success", "reply": reply}

    except Exception as e:
        logger.error(f"Error in ai_chat: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing chat: {str(e)}")


# ── Resume Sanity Check ──────────────────────────────────────────────────────

class SanityContactField(BaseModel):
    label: str
    value: str

class SanityJobTitle(BaseModel):
    title: str
    date: str
    bullets: List[str]

class SanityExperience(BaseModel):
    company: str
    job_titles: List[SanityJobTitle]

class SanityDegree(BaseModel):
    degree: str
    description: str

class SanityEducation(BaseModel):
    university: str
    date: str
    degrees: List[SanityDegree]

class SanityProject(BaseModel):
    name: str
    date: str
    description: str
    bullets: List[str]
    technologies: List[str] = []

class SanitySkill(BaseModel):
    topic: str
    keywords: str

class ResumeSanityCheckRequest(BaseModel):
    name: str
    contact_fields: List[SanityContactField]
    professional_experiences: List[SanityExperience]
    education: List[SanityEducation]
    projects_established: List[SanityProject]
    projects_expanding: List[SanityProject]
    skills: List[SanitySkill]


def _format_resume_for_sanity(req: ResumeSanityCheckRequest) -> str:
    lines = []
    lines.append(f"NAME: {req.name}")
    lines.append("")
    lines.append("CONTACT FIELDS:")
    for f in req.contact_fields:
        lines.append(f"  {f.label}: {f.value}")
    lines.append("")

    if req.professional_experiences:
        lines.append("PROFESSIONAL EXPERIENCE:")
        for exp in req.professional_experiences:
            lines.append(f"  Company: {exp.company}")
            for jt in exp.job_titles:
                lines.append(f"    Job Title: {jt.title}")
                lines.append(f"    Date/Location Field: {jt.date}")
                for b in jt.bullets:
                    if b.strip():
                        lines.append(f"      - {b}")
        lines.append("")

    if req.education:
        lines.append("EDUCATION:")
        for edu in req.education:
            lines.append(f"  University: {edu.university}")
            lines.append(f"  Date/Location Field: {edu.date}")
            for deg in edu.degrees:
                lines.append(f"    Degree: {deg.degree}")
                if deg.description.strip():
                    lines.append(f"    Description: {deg.description}")
        lines.append("")

    all_projects = req.projects_established + req.projects_expanding
    if all_projects:
        lines.append("PROJECTS:")
        for proj in all_projects:
            lines.append(f"  Project Name: {proj.name}")
            lines.append(f"  Date/Location Field: {proj.date}")
            if proj.description.strip():
                lines.append(f"  Description: {proj.description}")
            if proj.technologies:
                lines.append(f"  Technologies Array: {', '.join(proj.technologies)}")
            for b in proj.bullets:
                if b.strip():
                    lines.append(f"    - {b}")
        lines.append("")

    if req.skills:
        lines.append("SKILLS:")
        for sk in req.skills:
            lines.append(f"  {sk.topic}: {sk.keywords}")
        lines.append("")

    return "\n".join(lines)


_SANITY_RULES_PROMPT = """
Evaluate the resume against ALL of the following rules. For each rule that IS triggered (the issue exists), include it in the output JSON.
Only include rules that are actually triggered. Be specific — mention the exact name, company, text, or section involved.

HIGH SEVERITY RULES:
H1. NAME_MISSING: Name is blank, 'Your Name', or any obvious placeholder. → "You are missing name at the top of the resume, please kindly add your name."
H2. EMAIL_MISSING: Email contact field is blank or placeholder (e.g. 'your.email@example.com'). → "You are missing email address in the resume, please kindly add your email address."
H3. EMAIL_INVALID: Email exists but lacks @ or domain (invalid format). → "Your email address format is invalid, please check and correct."
H4. PHONE_MISSING: Phone field is blank or placeholder like '+1 (555) 123-4567'. → "You are missing phone number in the resume, please kindly add your phone number."
H5. LOCATION_MISSING: Location/address field is blank or placeholder like 'City, State, Country'. → "You are missing home address in the resume, please kindly add your home address."
H6. RESUME_TOO_LONG: Estimate if content would exceed one page (many sections, 10+ bullet points across all work/projects). → "Your resume may be over one page, consider trimming lower-priority details to keep it concise."
H7. WORK_EXP_MISSING_LOCATION: For each job title whose Date/Location field does NOT contain any character string text. → "The work experience '{title}' at '{company}' is missing a location city, please add it." (one issue per affected job)
H8. EDUCATION_MISSING_LOCATION: Education Date/Location field has no text string. → "The education at '{university}' is missing a location city, please add it."
H9. PROJECT_MISSING_LOCATION: Project Date/Location field has no text string. → "The project '{name}' is missing a location city, please add it."
H10. WORK_EXP_MISSING_DURATION: Job title Date/Location field has no year (no 4-digit year). → "The work experience '{title}' at '{company}' is missing a duration, please add it."
H11. EDUCATION_MISSING_DURATION: Education Date/Location has no year (no 4-digit year). → "The education at '{university}' is missing a duration, please add it."
H12. PROJECT_MISSING_DURATION: Project Date/Location has no year (no 4-digit year). → "The project '{name}' is missing a duration, please add it."
H13. NO_EDUCATION: No non-placeholder education entries exist. → "Your resume doesn't include any education experience, please add at least one education experience."
H14. NO_PROJECTS_AND_NO_WORK_BULLETS: No real projects exist AND no work experience bullet points exist. → "Your resume doesn't include any project experience, please add at least one project experience to demonstrate your professional skills."
H15. DOUBLE_SPACES: Any bullet or description has 2+ consecutive spaces. → "The following text has multiple continuous spaces, please remove the duplicated ones: '{snippet}'"
H16. INCOMPLETE_SENTENCE: A bullet ends with a preposition (of, in, at, for, to, by, on, from), conjunction (and, but, or), comma, semicolon, or colon. → "The following sentence appears to be incomplete, please complete it: '{sentence}'"
H17. SYNTAX_ERROR: A bullet has double punctuation (.., !!, ??), lowercase standalone 'i', incompleted sentence, or mismatched parentheses. → "There's a syntax error in the following sentence, please correct it. Consider revising: '{sentence}'"
H19. EDUCATION_MISSING_COLLEGE_NAME: A university field is blank or a placeholder ('University Name'). → "One of the education experiences is missing a college name, please add it."
H20. EDUCATION_MISSING_DEGREE: A degree field is blank or placeholder ('Degree Name'). → "The education at '{university}' is missing a college degree, please add it."

MID SEVERITY RULES:
M1. NO_COURSEWORK: Real education exists but no degree description mentions 'coursework', 'course work', or 'courses'. → "Recommend to have some coursework listed for the college major, which doesn't have the coursework."
M2. LINK_TOO_LONG: A link-type contact field (label contains link/website/github/portfolio/linkedin) has value longer than 35 characters. → "The link '{label}: {value}' is too long (over 35 characters), recommend to use a shorter URL."
M3. PROJECT_FEW_TECHNOLOGIES: A project has some technologies listed (in Technologies Array or bullets) but fewer than 4 total. → "The project '{name}' has fewer than 4 technologies listed ({tech_list}), recommend adding more technical keywords to demonstrate your skills in the project."
M4. PROJECT_BULLET_TOO_LONG: A project bullet point has more than 30 words. → "The project '{name}' has a bullet point with {N} words, recommend shortening it or splitting it into multiple: '{bullet}'"
M5. DUPLICATE_SKILLS: A keyword appears more than once across all skill entries. → "The following keywords appear multiple times in your technical skill section, recommend removing duplicates: {keywords}."
M6. PROJECT_DUPLICATE_TECHNOLOGIES: A project's Technologies Array has repeated items. → "The project '{name}' has duplicate technical keywords in its bullet points. Recommend listing unique technical keywords: {dupes}."
M7. PROJECT_TOO_MANY_BULLETS: A project has more than 5 non-empty bullets. → "The project '{name}' has {N} bullet points. Recommend having less than 4 bullet points for a project."

LOW SEVERITY RULES:
L1. NO_GPA: Real education exists but no degree description mentions 'GPA'. → "Recommend to have GPA if it's over 3.5."

Return ONLY valid JSON in this exact format (no markdown, no explanation):
{"issues": [{"severity": "High", "message": "..."}, ...]}
"""


def _to_ordinal(n: int) -> str:
    if 11 <= (n % 100) <= 13:
        return f"{n}th"
    return f"{n}" + {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")


@app.post("/resume-sanity-check")
async def resume_sanity_check(request: ResumeSanityCheckRequest):
    try:
        resume_text = _format_resume_for_sanity(request)

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": _SANITY_RULES_PROMPT},
                {"role": "user", "content": f"Resume data:\n{resume_text}"},
            ],
            response_format={"type": "json_object"},
            temperature=0,
        )

        raw = json.loads(response.choices[0].message.content)
        raw_issues = raw.get("issues", [])

        formatted = [
            {
                "severity": issue.get("severity", "Mid"),
                "ordinal": _to_ordinal(i + 1),
                "message": issue.get("message", ""),
            }
            for i, issue in enumerate(raw_issues)
            if issue.get("message", "").strip()
        ]

        return {"issues": formatted, "matched_count": len(formatted)}

    except Exception as e:
        logger.error(f"Error in resume_sanity_check: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error processing sanity check: {str(e)}")


# ambitology
@app.get("/get_usage/{cognito_sub}")
async def get_usage(cognito_sub: str):
    try:
        dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
        table = dynamodb.Table('ambit-dashboard-application-data')
        response = table.get_item(Key={'PK': cognito_sub, 'SK': 'USAGE'})
        item = response.get('Item', {})
        return {
            "craft_count": int(item.get('craft_count', 0)),
            "analysis_count": int(item.get('analysis_count', 0)),
            "download_count": int(item.get('download_count', 0)),
        }
    except Exception as e:
        logger.error(f"Error in get_usage: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error fetching usage: {str(e)}")


# ambitology
@app.post("/get_payment_history")
async def get_payment_history(request: Request):
    try:
        data = await request.json()
        email = data.get('email', '').strip().lower()
        cognito_sub = data.get('cognito_sub', '').strip()
        if not email and not cognito_sub:
            raise HTTPException(status_code=400, detail="email or cognito_sub is required")

        original_key = stripe.api_key
        stripe.api_key = STRIPE_SECRET_KEY_AMBITOLOGY
        try:
            customer_id = None

            # Strategy 1: look up customer by exact email
            if email:
                customers = stripe.Customer.list(email=email, limit=10)
                if customers.data:
                    customer_id = customers.data[0].id
                    logger.info(f"Found Stripe customer {customer_id} by email")

            # Strategy 2: search recent checkout sessions by cognito_sub metadata
            if not customer_id and cognito_sub:
                sessions = stripe.checkout.Session.list(limit=100)
                for session in sessions.auto_paging_iter():
                    meta = session.get('metadata') or {}
                    if meta.get('cognito_sub') == cognito_sub and session.get('customer'):
                        customer_id = session['customer']
                        logger.info(f"Found Stripe customer {customer_id} via checkout session metadata")
                        break

            if not customer_id:
                logger.info(f"No Stripe customer found for email={email} cognito_sub={cognito_sub}")
                return {"payments": [], "subscription": None}

            # Fetch invoices (main payment history source for subscriptions)
            invoices = stripe.Invoice.list(customer=customer_id, limit=50, expand=['data.subscription'])
            payments = []
            for inv in invoices.data:
                # Build a human-readable description with fallbacks
                description = None
                if inv.get('description'):
                    description = inv['description']
                elif inv.lines and inv.lines.data:
                    line = inv.lines.data[0]
                    description = line.get('description') or None
                if not description:
                    description = "Ambitology Subscription"

                amount = inv.get('amount_paid') or inv.get('amount_due') or 0
                payments.append({
                    "id": inv.id,
                    "amount": amount,
                    "currency": inv.get('currency', 'usd'),
                    "status": inv.get('status', 'unknown'),
                    "date": inv.get('created', 0),
                    "description": description,
                    "invoice_url": inv.get('hosted_invoice_url'),
                    "invoice_pdf": inv.get('invoice_pdf'),
                })

            # Sort by date descending (most recent first)
            payments.sort(key=lambda p: p['date'], reverse=True)

            # Get subscription info — check both active and trialing; also include
            # subscriptions with cancel_at_period_end set (still technically active)
            subscription_info = None
            for status in ('active', 'trialing', 'past_due'):
                subs = stripe.Subscription.list(customer=customer_id, status=status, limit=1)
                if subs.data:
                    sub = subs.data[0]
                    sub_dict = sub if isinstance(sub, dict) else sub.to_dict_recursive() if hasattr(sub, 'to_dict_recursive') else dict(sub)
                    subscription_info = {
                        "id": sub_dict.get("id"),
                        "status": sub_dict.get("status"),
                        "current_period_end": sub_dict.get("current_period_end"),
                        "cancel_at_period_end": sub_dict.get("cancel_at_period_end", False),
                    }
                    break

            return {"payments": payments, "subscription": subscription_info}
        finally:
            stripe.api_key = original_key

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in get_payment_history: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error fetching payment history: {str(e)}")


# ambitology
@app.post("/cancel_subscription")
async def cancel_subscription(request: Request):
    try:
        data = await request.json()
        email = data.get('email', '')
        cognito_sub = data.get('cognito_sub', '')
        if not email:
            raise HTTPException(status_code=400, detail="email is required")

        original_key = stripe.api_key
        stripe.api_key = STRIPE_SECRET_KEY_AMBITOLOGY
        stripe_cancelled = False
        try:
            customers = stripe.Customer.list(email=email, limit=5)
            if customers.data:
                customer_id = customers.data[0].id
                subscriptions = stripe.Subscription.list(customer=customer_id, status='active', limit=1)
                if subscriptions.data:
                    sub_id = subscriptions.data[0].id
                    stripe.Subscription.modify(sub_id, cancel_at_period_end=True)
                    stripe_cancelled = True
        finally:
            stripe.api_key = original_key

        # Always update DynamoDB plan to 'free' and mark cancellation
        if cognito_sub:
            dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
            table = dynamodb.Table('ambit-dashboard-application-data')
            table.update_item(
                Key={'PK': cognito_sub, 'SK': 'SUBSCRIPTION'},
                UpdateExpression='SET #p = :free, cancel_at_period_end = :v',
                ExpressionAttributeNames={'#p': 'plan'},
                ExpressionAttributeValues={':free': 'free', ':v': True}
            )

        msg = "Subscription cancelled at period end" if stripe_cancelled else "Plan downgraded to free"
        return {"status": "success", "message": msg}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in cancel_subscription: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error cancelling subscription: {str(e)}")
